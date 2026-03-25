# ajustes_finos_plantilla.py
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

print("=" * 60)
print("🔧 APLICANDO AJUSTES FINOS A LA PLANTILLA")
print("=" * 60)

# Ruta de la plantilla
BASE_DIR = "G:/Mi unidad/RRHH DIGITAL OFICIAL/BASE_DATOS"
plantilla_path = os.path.join(BASE_DIR, "plantillas", "PLANTILLA_DDJJ_OFICIAL.docx")

print(f"📁 Ajustando: {plantilla_path}")

# Abrir la plantilla existente
doc = Document(plantilla_path)

# ========================================
# AJUSTE 1: Encabezado en una sola línea
# ========================================
# Ya está correcto

# ========================================
# AJUSTE 2: Eliminar "0" antes de los textos
# ========================================
# Esto se soluciona en el código Python, no en la plantilla

# ========================================
# AJUSTE 3: Compactar el texto (eliminar espacios extras)
# ========================================
for paragraph in doc.paragraphs:
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

# ========================================
# AJUSTE 4: Ajustar la tabla de familiares
# ========================================
if len(doc.tables) > 0:
    table = doc.tables[0]  # Primera tabla (familiares)
    table.autofit = False
    table.allow_autofit = False
    
    # Anchos de columna específicos (en cm)
    widths = [4.5, 2.2, 1.2, 1.5, 1.0, 1.8, 1.0, 1.0, 1.0, 1.8]
    for i, width in enumerate(widths):
        for row in table.rows:
            cell = row.cells[i]
            cell.width = Cm(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Texto más pequeño para que entre
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)

# ========================================
# AJUSTE 5: Ajustar la tabla de padres y hermanos
# ========================================
if len(doc.tables) > 1:
    table2 = doc.tables[1]  # Segunda tabla (padres)
    table2.autofit = False
    table2.allow_autofit = False
    
    # Anchos de columna
    widths2 = [5.0, 2.5, 1.5, 2.0, 2.0, 4.0]
    for i, width in enumerate(widths2):
        for row in table2.rows:
            cell = row.cells[i]
            cell.width = Cm(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)

# ========================================
# AJUSTE 6: Tamaño de fuente general
# ========================================
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.size is None or run.font.size.pt > 10:
            run.font.size = Pt(9)

# Guardar
doc.save(plantilla_path)

print("✅ Ajustes finos aplicados correctamente")
print(f"📊 Tamaño final: {os.path.getsize(plantilla_path)} bytes")
print("=" * 60)