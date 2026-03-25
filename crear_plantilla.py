# crear_plantilla.py
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

print("🚀 Creando plantilla oficial de DDJJ...")

# Crear documento
doc = Document()

# Configurar página tamaño oficio
section = doc.sections[0]
section.page_height = Cm(35.56)
section.page_width = Cm(25.59)
section.top_margin = Cm(0.94)
section.bottom_margin = Cm(0.49)
section.left_margin = Cm(3)
section.right_margin = Cm(0)

# ========================================
# ENCABEZADO
# ========================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HONORABLE CONCEJO DELIBERANTE")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DE LA CIUDAD DE POSADAS")
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DIRECCIÓN DE PERSONAL")
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AÑO: {{ anio }}")
run.underline = True
run.font.size = Pt(12)

doc.add_paragraph()

# ========================================
# TÍTULO
# ========================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DECLARACIÓN JURADA DEL AGENTE DE PLANTA PERMANENTE,")
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONTRATADOS CON RELACIÓN DE DEPENDENCIA Y NOMINADOS")
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph()

# ========================================
# DATOS DEL AGENTE
# ========================================
p = doc.add_paragraph()
p.add_run("Legajo: {{ legajo }} ").bold = True
p.add_run("Apellido y Nombre del Agente: {{ apellido_nombre }}")

p = doc.add_paragraph()
p.add_run("Fecha de Nacimiento: {{ fecha_nacimiento }} ")
p.add_run("Lugar: {{ lugar_nacimiento }} ")
p.add_run("C.I./L.E/L.C/DNI Nº: {{ dni }}")

p = doc.add_paragraph()
p.add_run("Estado Civil: {{ estado_civil }} ")
p.add_run("Estudios cursados: (Remarcar lo que corresponde) -{{ estudios_completos }}0Completos-{{ estudios_incompletos }} 0Incompletos")

p = doc.add_paragraph()
p.add_run("{{ primario }} 0PRIMARIO - {{ secundario }} 0SECUNDARIO - {{ terciario }} 0TERCIARIO – {{ universitario }} 0UNIVERSITARIO.")

p = doc.add_paragraph()
p.add_run("Título de: {{ titulo }} ")
p.add_run("Años cursados: {{ anios_cursados }}")

doc.add_paragraph()

# ========================================
# DOMICILIO
# ========================================
p = doc.add_paragraph()
p.add_run("Domicilio particular: Barrio: {{ domicilio_barrio }} ")
p.add_run("Chacra: {{ domicilio_chacra }} ")
p.add_run("(calle o Av.): {{ domicilio_calle }} ")
p.add_run("Número: {{ domicilio_numero }} ")
p.add_run("Piso: {{ domicilio_piso }} ")
p.add_run("Departamento: {{ domicilio_depto }} ")
p.add_run("Localidad: {{ domicilio_localidad }}")

# ========================================
# CONTACTO
# ========================================
p = doc.add_paragraph()
p.add_run("Teléfono fijo: {{ telefono_fijo }} ")
p.add_run("Teléfono celular: {{ telefono_celular }} ")
p.add_run("Correo Electrónico: {{ email }}")

p = doc.add_paragraph()
p.add_run("Teléfono de referencia: {{ telefono_referencia }} ")
p.add_run("Nombre y Apellido: {{ nombre_referencia }} ")
p.add_run("parentesco: {{ parentesco_referencia }}")

# ========================================
# LABORAL
# ========================================
p = doc.add_paragraph()
p.add_run("Fecha de Ingreso: {{ fecha_ingr }} ")
p.add_run("Categoría: {{ categ }} ")
p.add_run("Antigüedad: {{ antig }} ")
p.add_run("0Planta Permanente {{ p_permanente }} ")
p.add_run("0Planta temporaria {{ p_temporaria }}")

p = doc.add_paragraph()
p.add_run("Lugar donde se Desempeña: {{ lugar_dependencia }}")

p = doc.add_paragraph()
p.add_run("Actividad que realiza: {{ actividad_realiza_dependencia }}")

# ========================================
# SALUD
# ========================================
p = doc.add_paragraph()
p.add_run("Enfermedad preexistente: {{ enfermedad_preexistente }}")

p = doc.add_paragraph()
p.add_run("Vacuna Covid-19, cantidad de dosis: {{ vacuna_covid_dosis }} ")
p.add_run("Vacuna Gripe: {{ vacuna_gripe }} ")
p.add_run("Vacuna Neumonía: {{ vacuna_neumonia }}")

# ========================================
# ASIGNACIONES
# ========================================
p = doc.add_paragraph()
p.add_run("¿Va a percibir Asig. Fliares en el HCD?: {{ percibe_asig_fliares_hcd }} ")
p.add_run("Servicios Anteriores: {{ servicios_anteriores }} ")
p.add_run("Antigüedad: {{ antiguedad_servicios }} ")
p.add_run("Trabaja/o en otra repartición: {{ trabaja_otra_reparticion }}")

p = doc.add_paragraph()
p.add_run("Donde: {{ otra_reparticion_donde }} ")
p.add_run("Percibe Asig. Fliares: {{ percibe_asig_fliares_otro }}")

doc.add_paragraph()

# ========================================
# ESTADO CIVIL Y CÓNYUGE
# ========================================
p = doc.add_paragraph()
p.add_run("PARA LOS AGENTES PARA EL COBRO DE ASIGNACIONES PERSONALES Y DEL GRUPO FAMILIAR").bold = True

p = doc.add_paragraph()
p.add_run("(Remarcar lo que corresponde) {{ soltero }} 0Soltero - {{ casado }}0 Casado - {{ divorciado }} 0Divorciado - {{ separado }} 0Separado - {{ viudo }} 0Viudo - {{ conviviente }} 0Conviviente.-")

p = doc.add_paragraph()
p.add_run("Apellido y Nombre del Cónyuge: {{ conyuge_apellido_nombre }}")
p.add_run("C.I./L.E./L.C./DNI Nº: {{ conyuge_dni }} ")
p.add_run("Enlace Contraído el: {{ conyuge_fecha_enlace }} ")
p.add_run("En: {{ conyuge_lugar_enlace }} ")
p.add_run("Trabaja: {{ conyuge_trabaja }} ")
p.add_run("Razón Social: {{ conyuge_razon_social }}")

p = doc.add_paragraph()
p.add_run("¿Percibe el cónyuge Asignaciones Familiares? {{ conyuge_percibe_asignaciones }} ")
p.add_run("Observaciones: {{ conyuge_observaciones }}")

doc.add_paragraph()

# ========================================
# FAMILIARES A CARGO
# ========================================
p = doc.add_paragraph()
p.add_run("Declaro Bajo Juramento que los Mencionados a Continuación están a Mi Cargo").bold = True

# Tabla de familiares
table = doc.add_table(rows=8, cols=10)
table.style = 'Table Grid'

# Encabezados
headers = ["APELLIDO Y NOMBRE", "Parentesco", "Convive", "F.NAC", "Edad", "DNI", "N.I.", "PRI", "SEC", "TER/UNIV"]
for i, header in enumerate(headers):
    table.cell(0, i).text = header

# Filas de datos
for i in range(1, 8):
    table.cell(i, 0).text = f"{{{{ f{i}.apellido }}}}"
    table.cell(i, 1).text = f"{{{{ f{i}.par }}}}"
    table.cell(i, 2).text = f"{{{{ f{i}.convive }}}}"
    table.cell(i, 3).text = f"{{{{ f{i}.fe_nac }}}}"
    table.cell(i, 4).text = f"{{{{ f{i}.ed }}}}"
    table.cell(i, 5).text = f"{{{{ f{i}.dni }}}}"
    table.cell(i, 6).text = f"{{{{ f{i}.ni }}}}"
    table.cell(i, 7).text = f"{{{{ f{i}.pri }}}}"
    table.cell(i, 8).text = f"{{{{ f{i}.sec }}}}"
    table.cell(i, 9).text = f"{{{{ f{i}.ter }}}}"

doc.add_paragraph()

# ========================================
# NOTIFICACIÓN
# ========================================
p = doc.add_paragraph()
p.add_run("ME NOTIFICO QUE EN EL CASO DE CONCURRIR MIS HIJOS A ESTABLECIMIENTOS EDUCATIVOS QUE ORIGINEN EL PAGO DE ASIGNACIONES POR ESCOLARIDAD, PRESENTARÉ AL INICIAR Y AL FINALIZAR EL AÑO LECTIVO LOS CERTIFICADOS CORRSPONDIENTES: ME COMPROMETO A COMUNICAR DE INMEDIATO A LA DIRECCIÓN DE PERSONAL DEL HCD LAS VARIANTES QUE SE PRODUZCAN Y MANTENDRÉ PERMANENTEMENTE ACTUALIZADA LA PRESENTE DECLARACIÓN JURADA.-").italic = True

doc.add_paragraph()

# ========================================
# LUGAR Y FECHA
# ========================================
p = doc.add_paragraph()
p.add_run("LUGAR Y FECHA:  {{ lugar_declaracion }};{{ fecha_declaracion }}")

doc.add_paragraph()

# ========================================
# PRIMERA FIRMA
# ========================================
p = doc.add_paragraph()
p.add_run("      {{ sello_firma_director }}                              …{{ sello_organismo }}…      .                          {{ firma_agente }}")

p = doc.add_paragraph()
p.add_run("SELLO Y FIRMA DEL ORGANISMO                      SELLO DEL ORGANISMO                                     FIRMA DEL AGENTE")

doc.add_paragraph()

# ========================================
# NOTA
# ========================================
p = doc.add_paragraph()
p.add_run("NOTA: la demostración de la existencia de los cargos denunciados debe efectuarse exhibiendo original de: TITULOS, PARTIDAS DE MATRIMONIO y NACIMIENTO DE LOS HIJOS y entregando fotocopias de las mismas que quedaran en poder de la DIRECCIÓN DE PERSONAL del H.C.D. Posadas. -").italic = True

doc.add_page_break()

# ========================================
# SEGUNDA PÁGINA - PADRES Y HERMANOS
# ========================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HONORABLE CONCEJO DELIBERANTE")
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DE LA CIUDAD DE POSADAS")
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DIRECCIÓN DE PERSONAL")
run.bold = True

doc.add_paragraph()

# Tabla de padres y hermanos
table2 = doc.add_table(rows=7, cols=6)
table2.style = 'Table Grid'

# Encabezados
headers2 = ["Nombre y Apellido de Padres y Hermanos", "Parentesco", "Convive", "F. Nac.", "DNI", "Observación"]
for i, header in enumerate(headers2):
    table2.cell(0, i).text = header

# Filas de datos
for i in range(1, 7):
    table2.cell(i, 0).text = f"{{{{ ph_{i}_apellido }}}}"
    table2.cell(i, 1).text = f"{{{{ ph_{i}_parentesco }}}}"
    table2.cell(i, 2).text = f"{{{{ ph_{i}_convive }}}}"
    table2.cell(i, 3).text = f"{{{{ ph_{i}_fecha_nac }}}}"
    table2.cell(i, 4).text = f"{{{{ ph_{i}_dni }}}}"
    table2.cell(i, 5).text = f"{{{{ ph_{i}_observacion }}}}"

doc.add_paragraph()

# ========================================
# ANEXO LEGAL
# ========================================
p = doc.add_paragraph()
p.add_run("ANEXO A DECLARACIÓN JURADA").bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Declaro bajo juramento de ley, que expresamente reconozco y no me comprende las limitaciones y prohibiciones contenidas en el Artículo 11 Inciso m) Apartado 1) del Estatuto del Empleado Municipal Ordenanza XI- N° 15 (antes Ordenanza Nº: 1059/03) y el Artículo 1º de la Ley Provincial Nº: 3.200 que textualmente dicen:")

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("ESTATUTO DEL EMPLEADO MUNICIPAL – CAPITULO III – INCISO m) – APARTADO 1):").bold = True

p = doc.add_paragraph()
p.add_run("Los Empleados Públicos están obligados a declarar bajo Juramento: “Su situación patrimonial y modificaciones posteriores en el tiempo y forma que prescriben la Carta Orgánica y su reglamentación, proporcionando los informes y documentos que al respecto lo requieran. -")
p.add_run("La nómina de familiares a cargo, y comunicar dentro del plazo de quince (15) días de producido el cambio de estado civil o variantes de carácter familiar, acompañando la documentación correspondiente. -")
p.add_run("Sus actividades fuera y dentro de la Administración Pública Provincial y Nacional, si posee otros ingresos a través de empleo de la Administración Pública, a fin de establecer si son compatibles con sus funciones en los términos del presente Estatuto”.-")

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("LEY PROVINCIAL Nº 3.200 Artículo 1º:").bold = True

p = doc.add_paragraph()
p.add_run("“No podrá acumularse en una misma persona, dos o más empleos así sean de la Administración Nacional, Provincial o Municipal, estando comprendidos en la prohibición:")
p.add_run("a) Los empleados y funcionarios de todas las categorías y niveles;")
p.add_run("b) Los que se desempeñan en cualquiera de los Poderes del Estado, Organismos de la Constitución, Organismos Descentralizados o Autárquicos, Sociedades con Participación Accionaría Mayoritaria del Estado, Economía Mixta, o Sociedades con Participación Accionaria Mayoritaria del Estado Nacional y/o Provincial;")
p.add_run("c) El personal permanente o temporario, con relación de dependencia, contratados o nombrado, cualquiera que fuere la modalidad jurídica que adopte;")
p.add_run("d) Los agentes que acumulen un beneficio provisional o retiro de cualquier índole, Nacional, Provincial o Municipal con algún empleo provincial definido en los incisos anteriores, salvo que estuviesen expresamente autorizados por la ley;")
p.add_run("e) Los casos que no estuvieren establecidos expresamente por ley como comprendidos en la excepción del Artículo 78º de la Constitución Provincial, o cuando estuviere expresamente autorizados por ley y hubiere incompatibilidad horaria total o parcial, para el desempeño de ambos cargos”.-")

doc.add_paragraph()

# ========================================
# SEGUNDA FIRMA
# ========================================
p = doc.add_paragraph()
p.add_run("{{ sello_firma_director2 }}               …	{{ sello_organismo2 }}…                             {{ firma_agente2 }}")

p = doc.add_paragraph()
p.add_run("SELLO Y FIRMA DEL ORGANISMO                                       SELLO DEL ORGANISMO                                               FIRMA DEL AGENTE")

doc.add_paragraph()

# ========================================
# QR Y CÓDIGO
# ========================================
p = doc.add_paragraph()
p.add_run("{{qr_validacion}}")

p = doc.add_paragraph()
p.add_run("Código de verificación: {{hash_documento}}")

# Guardar el documento
plantilla_path = "G:/Mi unidad/RRHH DIGITAL OFICIAL/BASE_DATOS/plantillas/PLANTILLA_DDJJ_OFICIAL.docx"
os.makedirs(os.path.dirname(plantilla_path), exist_ok=True)
doc.save(plantilla_path)

print(f"✅ Plantilla creada exitosamente en: {plantilla_path}")
print("✅ Listo para usar con el sistema RRHH DIGITAL")