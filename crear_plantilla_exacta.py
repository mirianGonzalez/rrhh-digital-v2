# crear_plantilla_exacta.py
import os
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

print("=" * 60)
print("🚀 CREANDO PLANTILLA EXACTA AL MODELO FÍSICO (2 PÁGINAS)")
print("=" * 60)

# Ruta de la plantilla
BASE_DIR = "G:/Mi unidad/RRHH DIGITAL OFICIAL/BASE_DATOS"
plantilla_path = os.path.join(BASE_DIR, "plantillas", "PLANTILLA_DDJJ_OFICIAL.docx")

print(f"📁 Creando en: {plantilla_path}")

# Crear carpeta si no existe
os.makedirs(os.path.dirname(plantilla_path), exist_ok=True)

# Crear documento
doc = Document()

# Configurar página tamaño oficio
section = doc.sections[0]
section.page_height = Cm(35.56)
section.page_width = Cm(25.59)
section.top_margin = Cm(0.94)
section.bottom_margin = Cm(0.49)
section.left_margin = Cm(3)
section.right_margin = Cm(0)

# ========================================
# PÁGINA 1 - ENCABEZADO (todo en una línea como el modelo)
# ========================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HONORABLE CONCEJO DELIBERANTE  DE LA CIUDAD DE POSADAS  DIRECCIÓN DE PERSONAL  ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("AÑO: {{ anio }}")
run.underline = True
run.font.size = Pt(10)

doc.add_paragraph()

# TÍTULO
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DECLARACIÓN JURADA DEL AGENTE DE PLANTA PERMANENTE,")
run.bold = True
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONTRATADOS CON RELACIÓN DE DEPENDENCIA Y NOMINADOS")
run.bold = True
run.font.size = Pt(9)

doc.add_paragraph()

# ========================================
# DATOS DEL AGENTE (formato compacto)
# ========================================
p = doc.add_paragraph()
p.add_run("Legajo: {{ legajo }} ").bold = True
p.add_run("Apellido y Nombre del Agente: {{ apellido_nombre }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Fecha de Nacimiento: {{ fecha_nacimiento }} ")
p.add_run("Lugar: {{ lugar_nacimiento }} ")
p.add_run("C.I./L.E/L.C/DNI Nº: {{ dni }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Estado Civil: {{ estado_civil }} ")
p.add_run("Estudios cursados:(Remarcar lo que corresponde) -{{ estudios_completos }}Completos-{{ estudios_incompletos }} Incompletos")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("{{ primario }} PRIMARIO - {{ secundario }} SECUNDARIO - {{ terciario }} TERCIARIO – {{ universitario }} UNIVERSITARIO.")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Título de: {{ titulo }} ")
p.add_run("Años cursados: {{ anios_cursados }}")
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# DOMICILIO
p = doc.add_paragraph()
p.add_run("Domicilio particular: Barrio: {{ domicilio_barrio }} ")
p.add_run("Chacra: {{ domicilio_chacra }} ")
p.add_run("(calle o Av.): {{ domicilio_calle }} ")
p.add_run("Número: {{ domicilio_numero }} ")
p.add_run("Piso: {{ domicilio_piso }} ")
p.add_run("Departamento: {{ domicilio_depto }} ")
p.add_run("Localidad: {{ domicilio_localidad }}")
p.paragraph_format.space_after = Pt(0)

# CONTACTO
p = doc.add_paragraph()
p.add_run("Teléfono fijo: {{ telefono_fijo }} ")
p.add_run("Teléfono celular: {{ telefono_celular }} ")
p.add_run("Correo Electrónico: {{ email }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Teléfono de referencia: {{ telefono_referencia }} ")
p.add_run("Nombre y Apellido: {{ nombre_referencia }} ")
p.add_run("parentesco: {{ parentesco_referencia }}")
p.paragraph_format.space_after = Pt(0)

# LABORAL
p = doc.add_paragraph()
p.add_run("Fecha de Ingreso: {{ fecha_ingr }} ")
p.add_run("Categoría: {{ categ }} ")
p.add_run("Antigüedad: {{ antig }} ")
p.add_run("Planta Permanente {{ p_permanente }} ")
p.add_run("Planta temporaria {{ p_temporaria }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Lugar donde se Desempeña: {{ lugar_dependencia }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Actividad que realiza: {{ actividad_realiza_dependencia }}")
p.paragraph_format.space_after = Pt(0)

# SALUD
p = doc.add_paragraph()
p.add_run("Enfermedad preexistente: {{ enfermedad_preexistente }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Vacuna Covid-19, cantidad de dosis: {{ vacuna_covid_dosis }} ")
p.add_run("Vacuna Gripe: {{ vacuna_gripe }} ")
p.add_run("Vacuna Neumonía: {{ vacuna_neumonia }}")
p.paragraph_format.space_after = Pt(0)

# ASIGNACIONES
p = doc.add_paragraph()
p.add_run("¿Va a percibir Asig. Fliares en el HCD?: {{ percibe_asig_fliares_hcd }} ")
p.add_run("Servicios Anteriores: {{ servicios_anteriores }} ")
p.add_run("Antigüedad: {{ antiguedad_servicios }} ")
p.add_run("Trabaja/o en otra repartición: {{ trabaja_otra_reparticion }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Donde: {{ otra_reparticion_donde }} ")
p.add_run("Percibe Asig. Fliares: {{ percibe_asig_fliares_otro }}")
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# ========================================
# ESTADO CIVIL Y CÓNYUGE
# ========================================
p = doc.add_paragraph()
p.add_run("PARA LOS AGENTES PARA EL COBRO DE ASIGNACIONES PERSONALES Y DEL GRUPO FAMILIAR").bold = True
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("(Remarcar lo que corresponde) {{ soltero }} Soltero - {{ casado }} Casado - {{ divorciado }} Divorciado - {{ separado }} Separado - {{ viudo }} Viudo - {{ conviviente }} Conviviente.-")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Apellido y Nombre del Cónyuge: {{ conyuge_apellido_nombre }}")
p.add_run("C.I./L.E./L.C./DNI Nº: {{ conyuge_dni }} ")
p.add_run("Enlace Contraído el: {{ conyuge_fecha_enlace }} ")
p.add_run("En: {{ conyuge_lugar_enlace }} ")
p.add_run("Trabaja: {{ conyuge_trabaja }} ")
p.add_run("Razón Social: {{ conyuge_razon_social }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("¿Percibe el cónyuge Asignaciones Familiares? {{ conyuge_percibe_asignaciones }} ")
p.add_run("Observaciones: {{ conyuge_observaciones }}")
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# ========================================
# FAMILIARES A CARGO
# ========================================
p = doc.add_paragraph()
p.add_run("Declaro Bajo Juramento que los Mencionados a Continuación están a Mi Cargo").bold = True
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("*Colocar en número, el grado o año de primaria o secundaria que concurre.").italic = True
p.paragraph_format.space_after = Pt(0)

# Tabla de familiares
table = doc.add_table(rows=8, cols=10)
table.style = 'Table Grid'
table.autofit = False
table.allow_autofit = False

# Anchos de columna
widths = [Cm(4.5), Cm(2.5), Cm(1.5), Cm(1.5), Cm(1.2), Cm(2.0), Cm(1.2), Cm(1.2), Cm(1.2), Cm(2.0)]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# Encabezados
headers = ["APELLIDO Y NOMBRE D/LOS HIJOS U OTRO FAMILIAR A CARGO", "Parentesco", "Convive", "F.NAC", "Edad", "DNI", "N.I.", "PRI", "SEC", "TER O UNIV"]
for i, header in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = header
    cell.paragraphs[0].runs[0].font.size = Pt(7)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

for i in range(1, 8):
    row_data = [
        f"{{{{ f{i}.apellido }}}}",
        f"{{{{ f{i}.par }}}}",
        f"{{{{ f{i}.convive }}}}",
        f"{{{{ f{i}.fe_nac }}}}",
        f"{{{{ f{i}.ed }}}}",
        f"{{{{ f{i}.dni }}}}",
        f"{{{{ f{i}.ni }}}}",
        f"{{{{ f{i}.pri }}}}",
        f"{{{{ f{i}.sec }}}}",
        f"{{{{ f{i}.ter }}}}"
    ]
    for j, text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = text
        cell.paragraphs[0].runs[0].font.size = Pt(7)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# ========================================
# NOTIFICACIÓN
# ========================================
p = doc.add_paragraph()
p.add_run("ME NOTIFICO QUE EN EL CASO DE CONCURRIR MIS HIJOS A ESTABLECIMIENTOS EDUCATIVOS QUE ORIGINEN EL PAGO DE ASIGNACIONES POR ESCOLARIDAD, PRESENTARÉ AL INICIAR Y AL FINALIZAR EL AÑO LECTIVO LOS CERTIFICADOS CORRSPONDIENTES: ME COMPROMETO A COMUNICAR DE INMEDIATO A LA DIRECCIÓN DE PERSONAL DEL HCD LAS VARIANTES QUE SE PRODUZCAN Y MANTENDRÉ PERMANENTEMENTE ACTUALIZADA LA PRESENTE DECLARACIÓN JURADA.-").italic = True
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

doc.add_paragraph()

# LUGAR Y FECHA
p = doc.add_paragraph()
p.add_run("LUGAR Y FECHA:  {{ lugar_declaracion }};{{ fecha_declaracion }}")
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# PRIMERA FIRMA
p = doc.add_paragraph()
p.add_run("      {{ sello_firma_director }}                              …{{ sello_organismo }}…      .                          {{ firma_agente }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("SELLO Y FIRMA DEL ORGANISMO                      SELLO DEL ORGANISMO                                     FIRMA DEL AGENTE")
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# NOTA
p = doc.add_paragraph()
p.add_run("NOTA: la demostración de la existencia de los cargos denunciados debe efectuarse exhibiendo original de: TITULOS, PARTIDAS DE MATRIMONIO y NACIMIENTO DE LOS HIJOS y entregando fotocopias de las mismas que quedaran en poder de la DIRECCIÓN DE PERSONAL del H.C.D. Posadas. -").italic = True
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

# ========================================
# SALTO DE PÁGINA
# ========================================
doc.add_page_break()

# ========================================
# PÁGINA 2 - ENCABEZADO
# ========================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HONORABLE CONCEJO DELIBERANTE  DE LA CIUDAD DE POSADAS  DIRECCIÓN DE PERSONAL")
run.bold = True
run.font.size = Pt(10)
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# Tabla de padres y hermanos
table2 = doc.add_table(rows=7, cols=6)
table2.style = 'Table Grid'
table2.autofit = False
table2.allow_autofit = False

# Anchos de columna
widths2 = [Cm(5.0), Cm(2.5), Cm(1.5), Cm(2.0), Cm(2.0), Cm(4.0)]
for row in table2.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths2[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

headers2 = ["Nombre y Apellido de Padres y Hermanos.", "Parentesco", "Convive", "F. Nac.", "DNI", "Observación:"]
for i, header in enumerate(headers2):
    cell = table2.cell(0, i)
    cell.text = header
    cell.paragraphs[0].runs[0].font.size = Pt(7)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

for i in range(1, 7):
    row_data = [
        f"{{{{ ph_{i}_apellido }}}}",
        f"{{{{ ph_{i}_parentesco }}}}",
        f"{{{{ ph_{i}_convive }}}}",
        f"{{{{ ph_{i}_fecha_nac }}}}",
        f"{{{{ ph_{i}_dni }}}}",
        f"{{{{ ph_{i}_observacion }}}}"
    ]
    for j, text in enumerate(row_data):
        cell = table2.cell(i, j)
        cell.text = text
        cell.paragraphs[0].runs[0].font.size = Pt(7)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# ========================================
# ANEXO LEGAL
# ========================================
p = doc.add_paragraph()
p.add_run("ANEXO A DECLARACIÓN JURADA").bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Declaro bajo juramento de ley, que expresamente reconozco y no me comprende las limitaciones y prohibiciones contenidas en el Artículo 11 Inciso m) Apartado 1) del Estatuto del Empleado Municipal Ordenanza XI- N° 15 (antes Ordenanza Nº: 1059/03) y el Artículo 1º de la Ley Provincial Nº: 3.200 que textualmente dicen:")
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("ESTATUTO DEL EMPLEADO MUNICIPAL – CAPITULO III – INCISO m) – APARTADO 1):").bold = True
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Los Empleados Públicos están obligados a declarar bajo Juramento: “Su situación patrimonial y modificaciones posteriores en el tiempo y forma que prescriben la Carta Orgánica y su reglamentación, proporcionando los informes y documentos que al respecto lo requieran. -")
p.add_run("La nómina de familiares a cargo, y comunicar dentro del plazo de quince (15) días de producido el cambio de estado civil o variantes de carácter familiar, acompañando la documentación correspondiente. -")
p.add_run("Sus actividades fuera y dentro de la Administración Pública Provincial y Nacional, si posee otros ingresos a través de empleo de la Administración Pública, a fin de establecer si son compatibles con sus funciones en los términos del presente Estatuto”.-")
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("LEY PROVINCIAL Nº 3.200 Artículo 1º:").bold = True
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("“No podrá acumularse en una misma persona, dos o más empleos así sean de la Administración Nacional, Provincial o Municipal, estando comprendidos en la prohibición:")
p.add_run("a) Los empleados y funcionarios de todas las categorías y niveles;")
p.add_run("b) Los que se desempeñan en cualquiera de los Poderes del Estado, Organismos de la Constitución, Organismos Descentralizados o Autárquicos, Sociedades con Participación Accionaría Mayoritaria del Estado, Economía Mixta, o Sociedades con Participación Accionaria Mayoritaria del Estado Nacional y/o Provincial;")
p.add_run("c) El personal permanente o temporario, con relación de dependencia, contratados o nombrado, cualquiera que fuere la modalidad jurídica que adopte;")
p.add_run("d) Los agentes que acumulen un beneficio provisional o retiro de cualquier índole, Nacional, Provincial o Municipal con algún empleo provincial definido en los incisos anteriores, salvo que estuviesen expresamente autorizados por la ley;")
p.add_run("e) Los casos que no estuvieren establecidos expresamente por ley como comprendidos en la excepción del Artículo 78º de la Constitución Provincial, o cuando estuviere expresamente autorizados por ley y hubiere incompatibilidad horaria total o parcial, para el desempeño de ambos cargos”.-")
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

doc.add_paragraph()

# ========================================
# SEGUNDA FIRMA
# ========================================
p = doc.add_paragraph()
p.add_run("{{ sello_firma_director2 }}               …	{{ sello_organismo2 }}…                             {{ firma_agente2 }}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("SELLO Y FIRMA DEL ORGANISMO                                       SELLO DEL ORGANISMO                                               FIRMA DEL AGENTE")
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# ========================================
# QR Y CÓDIGO
# ========================================
p = doc.add_paragraph()
p.add_run("{{qr_validacion}}")
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
p.add_run("Código de verificación: {{hash_documento}}")
p.paragraph_format.space_after = Pt(0)

# Guardar
doc.save(plantilla_path)

print("=" * 60)
print(f"✅ Plantilla EXACTA creada exitosamente")
print(f"📁 Ruta: {plantilla_path}")
print(f"📊 Tamaño: {os.path.getsize(plantilla_path)} bytes")
print("=" * 60)
print("\n🎯 Características:")
print("   • 2 páginas exactas como el modelo")
print("   • Texto compacto sin espacios en blanco")
print("   • Tablas con anchos fijos")
print("   • Formato profesional")
print("=" * 60)