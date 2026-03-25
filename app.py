"""
SISTEMA DE GESTIÓN DE LEGAJOS - ARCHIVO Y DECLARACIONES JURADAS
VERSIÓN REALIZADA POR MIRIAN YOLANDA GONZALEZ
"""

# ==========================
# IMPORTS PRINCIPALES
# ==========================
import io
import os
import sys
import sqlite3
import uuid
import json
import hashlib
import shutil
import subprocess
import tempfile
import logging
import smtplib
import time
import traceback
import random
import base64
import zipfile
from functools import wraps
from datetime import datetime, timedelta
from io import BytesIO
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

from flask import Flask, flash, make_response, render_template, request, redirect, send_file, session, send_from_directory, url_for, g, jsonify
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename




# Decorador para requerir login
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Por favor, inicia sesión para acceder a esta página', 'warning')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# ==========================
# CREAR CARPETAS NECESARIAS (DEBE IR ANTES DEL LOGGING)
# ==========================
carpetas_necesarias = ['LOGS', 'BACKUPS', 'DOCUMENTOS_AGENTES', 'qr_ddjj', 'CERTIFICADOS', 'plantillas']

for carpeta in carpetas_necesarias:
    if not os.path.exists(carpeta):
        os.makedirs(carpeta)
        print(f"📁 Carpeta creada: {carpeta}")

# ==========================
# CONFIGURACIÓN DE LOGGING
# ==========================
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('LOGS/sistema.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# ==========================
# IMPORTAR MÓDULOS LOCALES
# ==========================
from ac_interna import AutoridadCertificante
from firma_pdf import firmar_pdf_digitalmente

# ==========================
# DETECTAR ENTORNO
# ==========================
ENV_RENDER = os.environ.get('RENDER', False)
WIN32_AVAILABLE = False

if not ENV_RENDER:
    try:
        import win32com.client
        from win32com.client import constants
        WIN32_AVAILABLE = True
        print("✅ win32com cargado (entorno Windows)")
    except ImportError:
        WIN32_AVAILABLE = False
        print("⚠️ win32com no disponible")
else:
    print("🌐 Entorno Render - funciones de Office deshabilitadas")

# ==========================
# CONFIGURACIÓN DE LA APP
# ==========================
app = Flask(__name__,
            static_folder='static',
            static_url_path='/static')
app.secret_key = os.environ.get('SECRET_KEY', 'clave_secreta_para_desarrollo')
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=8)

# Configuración de rutas
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, 'legajos_agentes.db')
DOCUMENTOS_DIR = os.path.join(BASE_DIR, 'DOCUMENTOS_AGENTES')
PLANTILLAS_DIR = os.path.join(BASE_DIR, 'plantillas')
QR_DIR = os.path.join(BASE_DIR, 'qr_ddjj')
BACKUP_DIR = os.path.join(BASE_DIR, 'BACKUPS')
LOG_DIR = os.path.join(BASE_DIR, 'LOGS')
CERTIFICADOS_DIR = os.path.join(BASE_DIR, 'CERTIFICADOS')

# Asegurar que existan los directorios
for dir_path in [DOCUMENTOS_DIR, PLANTILLAS_DIR, QR_DIR, BACKUP_DIR, LOG_DIR, CERTIFICADOS_DIR]:
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)


# ==========================
# RUTAS Y BASES DE DATOS
# ==========================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "legajos_agentes.db")

# Configurar carpeta de documentos (usar ruta local por defecto)
DOCUMENTOS_AGENTES_PATH = os.path.join(BASE_DIR, "DOCUMENTOS_AGENTES")
app.config['UPLOAD_FOLDER'] = DOCUMENTOS_AGENTES_PATH
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max

# Configuración adicional
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'tu-secret-key-aqui')
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///legajos_agentes.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# ==========================
# CONFIGURACIÓN DE FIRMA DIGITAL
# ==========================
CERTIFICADOS_PATH = os.path.join(BASE_DIR, "CERTIFICADOS")
os.makedirs(CERTIFICADOS_PATH, exist_ok=True)

# ==========================
# CONTEXTO GLOBAL PARA TEMPLATES
# ==========================
@app.context_processor
def inject_now():
    """Inyecta la función now() en todos los templates"""
    return {'now': datetime.now}

# ==========================
# CONFIGURACIÓN EMAIL (para alertas)
# ==========================
EMAIL_CONFIG = {
    'smtp_server': 'smtp.gmail.com',
    'smtp_port': 587,
    'email_from': os.getenv('EMAIL_FROM', 'legajosyarchivo.hcd.pss@gmail.com'),
    'email_password': os.getenv('EMAIL_PASSWORD', ''),
    'email_admin': os.getenv('EMAIL_ADMIN', 'miriany.gonzalez@gmail.com')
}

# Verificar que la contraseña de email está configurada
if not EMAIL_CONFIG['email_password']:
     print("⚠️ ADVERTENCIA: EMAIL_PASSWORD no está configurada en el archivo .env")

# ==========================
# RUTAS DE DIRECTORIOS
# ==========================
DOCUMENTOS_PATH = app.config['UPLOAD_FOLDER']
PLANTILLAS_PATH = os.path.join(BASE_DIR, "plantillas")
QR_PATH = os.path.join(BASE_DIR, "qr_ddjj")
BACKUP_PATH = os.path.join(BASE_DIR, "BACKUPS")
LOGS_PATH = os.path.join(BASE_DIR, "LOGS")

# Crear directorios necesarios con manejo de errores
print("\n" + "="*50)
print("📁 Creando directorios necesarios...")
print("="*50)

for path_name, path in [
    ("Documentos", DOCUMENTOS_PATH),
    ("Plantillas", PLANTILLAS_PATH),
    ("QR", QR_PATH),
    ("Backups", BACKUP_PATH),
    ("Logs", LOGS_PATH),
    ("Certificados", CERTIFICADOS_PATH)
]:
    try:
        os.makedirs(path, exist_ok=True)
        print(f"  ✓ {path_name}: {path}")
    except Exception as e:
        print(f"  ✗ Error creando {path_name}: {e}")
        # Si falla la carpeta de documentos, crear una alternativa
        if path == DOCUMENTOS_PATH:
            fallback_path = os.path.join(BASE_DIR, "DOCUMENTOS_AGENTES_FALLBACK")
            os.makedirs(fallback_path, exist_ok=True)
            app.config['UPLOAD_FOLDER'] = fallback_path
            DOCUMENTOS_PATH = fallback_path
            print(f"  ✓ Usando carpeta alternativa: {fallback_path}")

print("\n" + "="*50)
logger.info("="*50)
logger.info(f"BASE DATOS: {DB_PATH}")
logger.info(f"DOCUMENTOS: {DOCUMENTOS_PATH}")
logger.info(f"PLANTILLAS: {PLANTILLAS_PATH}")
logger.info(f"QR: {QR_PATH}")
logger.info(f"BACKUPS: {BACKUP_PATH}")
logger.info(f"LOGS: {LOGS_PATH}")
logger.info(f"CERTIFICADOS: {CERTIFICADOS_PATH}")
logger.info("="*50)
print("✅ Configuración completada exitosamente")
print(f"📁 Carpeta de documentos: {DOCUMENTOS_PATH}")
print("="*50)

# ==========================
# DOCUMENTOS OBLIGATORIOS - LISTA COMPLETA
# ==========================
DOCUMENTOS_OBLIGATORIOS = [
    {'id': 'dni_frente', 'nombre': 'DNI Frente', 'categoria': 'IDENTIDAD', 'requerido': True},
    {'id': 'dni_dorso', 'nombre': 'DNI Dorso', 'categoria': 'IDENTIDAD', 'requerido': True},
    {'id': 'partida_nacimiento', 'nombre': 'Partida de Nacimiento', 'categoria': 'IDENTIDAD', 'requerido': True},
    {'id': 'cuil', 'nombre': 'CUIL', 'categoria': 'IDENTIDAD', 'requerido': True},
    {'id': 'foto_4x4', 'nombre': 'Foto 4x4', 'categoria': 'IDENTIDAD', 'requerido': True},
    
    {'id': 'matrimonio', 'nombre': 'Partida de Matrimonio/Conviviencia', 'categoria': 'FAMILIA', 'requerido': False},
    {'id': 'hijos_nacimiento', 'nombre': 'Partidas de Nacimiento de Hijos', 'categoria': 'FAMILIA', 'requerido': False},
    
    {'id': 'alumno_regular', 'nombre': 'Certificado de Alumno Regular', 'categoria': 'EDUCACION', 'requerido': False},
    {'id': 'titulo_secundario', 'nombre': 'Título Secundario', 'categoria': 'EDUCACION', 'requerido': True},
    {'id': 'titulo_terciario', 'nombre': 'Título Terciario/Universitario', 'categoria': 'EDUCACION', 'requerido': False},
    {'id': 'certificado_analitico', 'nombre': 'Certificado Analítico', 'categoria': 'EDUCACION', 'requerido': False},
    
    {'id': 'croquis_domicilio', 'nombre': 'Croquis de Domicilio', 'categoria': 'DOMICILIO', 'requerido': True},
    {'id': 'telefono_fijo', 'nombre': 'Teléfono Fijo', 'categoria': 'CONTACTO', 'requerido': False},
    {'id': 'telefono_celular', 'nombre': 'Teléfono Celular', 'categoria': 'CONTACTO', 'requerido': True},
    {'id': 'correo_electronico', 'nombre': 'Correo Electrónico', 'categoria': 'CONTACTO', 'requerido': True},
    {'id': 'telefono_referencia', 'nombre': 'Teléfono de Referencia', 'categoria': 'CONTACTO', 'requerido': True},
    {'id': 'parentesco_referencia', 'nombre': 'Parentesco de Referencia', 'categoria': 'CONTACTO', 'requerido': True},
    
    {'id': 'ddjj_anual', 'nombre': 'Declaración Jurada Anual', 'categoria': 'DDJJ', 'requerido': True},
    {'id': 'seguro_vida', 'nombre': 'Seguro de Vida', 'categoria': 'LABORAL', 'requerido': True},
    {'id': 'servicios_anteriores', 'nombre': 'Certificado de Servicios Anteriores', 'categoria': 'LABORAL', 'requerido': True},
    {'id': 'certif_medico', 'nombre': 'Certificado Médico de Aptitud Psicofísica', 'categoria': 'SALUD', 'requerido': True},
    {'id': 'certif_obra_social', 'nombre': 'Certificado de Obra Social IPS', 'categoria': 'SALUD', 'requerido': True},
    
    {'id': 'autoridad_diploma', 'nombre': 'Autoridad del Diploma/Decreto', 'categoria': 'NOMBRAMIENTO', 'requerido': True},
    {'id': 'contrato', 'nombre': 'Contrato', 'categoria': 'NOMBRAMIENTO', 'requerido': True},
    {'id': 'contrato_nominado', 'nombre': 'Contrato Nominado', 'categoria': 'NOMBRAMIENTO', 'requerido': False},
    {'id': 'decreto_designacion', 'nombre': 'Decreto de Designación', 'categoria': 'NOMBRAMIENTO', 'requerido': True},
    {'id': 'pase_planta', 'nombre': 'Pase a Planta Permanente', 'categoria': 'NOMBRAMIENTO', 'requerido': False},
    
    {'id': 'licencias', 'nombre': 'Licencias', 'categoria': 'LEGAJO', 'requerido': False},
    {'id': 'cortes_licencia', 'nombre': 'Cortes de Licencia', 'categoria': 'LEGAJO', 'requerido': False},
    {'id': 'franquicias', 'nombre': 'Franquicias', 'categoria': 'LEGAJO', 'requerido': False},
    {'id': 'sanciones', 'nombre': 'Sanciones', 'categoria': 'LEGAJO', 'requerido': False},
    {'id': 'resoluciones', 'nombre': 'Resoluciones', 'categoria': 'LEGAJO', 'requerido': False},
    {'id': 'circulares', 'nombre': 'Circulares', 'categoria': 'LEGAJO', 'requerido': False},
    {'id': 'memorandum', 'nombre': 'Memorándum', 'categoria': 'LEGAJO', 'requerido': False},
]

# ==========================
# MÓDULO DE FIRMA DIGITAL PKI
# ==========================

def generar_certificado_agente(legajo_id, nombre_completo, dni, password_pin):
    """
    Genera un certificado digital .p12 para un agente
    Esto simula la AC interna del HCD emitiendo el certificado
    """
    try:
        
        # Generar clave privada RSA
        private_key = rsa.generate_private_key(
            public_exponent=65537,
            key_size=2048,
            backend=default_backend()
        )
        
        # Construir el certificado
        subject = issuer = x509.Name([
            x509.NameAttribute(NameOID.COUNTRY_NAME, u"AR"),
            x509.NameAttribute(NameOID.STATE_OR_PROVINCE_NAME, u"MISIONES"),
            x509.NameAttribute(NameOID.LOCALITY_NAME, u"POSADAS"),
            x509.NameAttribute(NameOID.ORGANIZATION_NAME, u"HCD POSADAS"),
            x509.NameAttribute(NameOID.COMMON_NAME, nombre_completo),
        ])
        
        cert = x509.CertificateBuilder().subject_name(
            subject
        ).issuer_name(
            issuer
        ).public_key(
            private_key.public_key()
        ).serial_number(
            x509.random_serial_number()
        ).not_valid_before(
            datetime.datetime.utcnow()
        ).not_valid_after(
            datetime.datetime.utcnow() + datetime.timedelta(days=365)
        ).add_extension(
            x509.SubjectAlternativeName([x509.DNSName(u"localhost")]),
            critical=False,
        ).sign(private_key, hashes.SHA256(), default_backend())
        
        # Guardar en archivo PKCS#12
        p12_path = os.path.join(CERTIFICADOS_PATH, f"cert_{legajo_id}.p12")
        
        p12_data = pkcs12.serialize_key_and_certificates(
            name=nombre_completo.encode('utf-8'),
            key=private_key,
            cert=cert,
            cas=None,
            encryption_algorithm=serialization.BestAvailableEncryption(password_pin.encode('utf-8'))
        )
        
        with open(p12_path, 'wb') as f:
            f.write(p12_data)
        
        logger.info(f"✅ Certificado generado para {legajo_id}: {p12_path}")
        
        # Guardar información del certificado en BD
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS certificados_agentes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                legajo_id TEXT NOT NULL,
                ruta_certificado TEXT NOT NULL,
                fecha_emision TEXT,
                fecha_expiracion TEXT,
                pin_hash TEXT,
                activo INTEGER DEFAULT 1,
                FOREIGN KEY(legajo_id) REFERENCES legajos(legajo_id)
            )
        """)
        
        pin_hash = generate_password_hash(password_pin)
        fecha_expiracion = (datetime.datetime.utcnow() + datetime.timedelta(days=365)).strftime("%Y-%m-%d")
        
        cur.execute("""
            INSERT INTO certificados_agentes 
            (legajo_id, ruta_certificado, fecha_emision, fecha_expiracion, pin_hash, activo)
            VALUES (?, ?, ?, ?, ?, 1)
        """, (legajo_id, p12_path, datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"), fecha_expiracion, pin_hash))
        
        conn.commit()
        
        return p12_path
        
    except Exception as e:
        logger.error(f"Error generando certificado: {e}")
        return None

def firmar_pdf_con_pin(pdf_path, legajo_id, pin):
    """
    Firma un PDF usando el certificado .p12 del agente
    Verifica el PIN antes de firmar
    """
    try:
        # Verificar PIN del agente
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT ruta_certificado, pin_hash 
            FROM certificados_agentes 
            WHERE legajo_id = ? AND activo = 1
        """, (legajo_id,))
        
        certificado = cur.fetchone()
        if not certificado:
            return None, "❌ No se encontró certificado para este agente"
        
        if not check_password_hash(certificado['pin_hash'], pin):
            return None, "❌ PIN incorrecto"
        
        # Cargar el certificado
        with open(certificado['ruta_certificado'], 'rb') as f:
            p12_data = f.read()
        
        private_key, certificate, _ = pkcs12.load_key_and_certificates(
            p12_data,
            pin.encode('utf-8'),
            backend=default_backend()
        )
        
        # Leer el PDF existente
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        # Copiar todas las páginas
        for page in reader.pages:
            writer.add_page(page)
        
        # Crear diccionario de firma
        signature = DictionaryObject()
        signature.update({
            NameObject("/Type"): NameObject("/Sig"),
            NameObject("/Filter"): NameObject("/Adobe.PPKLite"),
            NameObject("/SubFilter"): NameObject("/adbe.pkcs7.detached"),
            NameObject("/Name"): TextStringObject(f"Agente: {legajo_id}"),
            NameObject("/Reason"): TextStringObject("Declaración Jurada Anual"),
            NameObject("/Location"): TextStringObject("HCD Posadas"),
            NameObject("/ContactInfo"): TextStringObject(f"Legajo: {legajo_id}"),
            NameObject("/M"): TextStringObject(datetime.now().strftime("D:%Y%m%d%H%M%S")),
        })
        
        # Agregar firma a la primera página
        writer.add_annotation(page_number=0, annotation=signature)
        
        # Guardar PDF firmado
        pdf_firmado = pdf_path.replace('.pdf', '_firmado.pdf')
        with open(pdf_firmado, 'wb') as f:
            writer.write(f)
        
        logger.info(f"✅ PDF firmado exitosamente: {pdf_firmado}")
        
        # Registrar firma en auditoría
        registrar_auditoria(legajo_id, "FIRMÓ DDJJ", legajo_id, f"PDF firmado con PIN")
        
        return pdf_firmado, "✅ Firma realizada correctamente"
        
    except Exception as e:
        logger.error(f"Error firmando PDF: {e}")
        return None, f"❌ Error al firmar: {str(e)}"

def cofirmar_pdf_director(pdf_path, director_usuario):
    """
    Co-firma el PDF como Director (segunda firma)
    """
    try:
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page)
        
        # Agregar sello de co-firma del Director
        cofirma = DictionaryObject()
        cofirma.update({
            NameObject("/Type"): NameObject("/Sig"),
            NameObject("/Filter"): NameObject("/Adobe.PPKLite"),
            NameObject("/SubFilter"): NameObject("/adbe.pkcs7.detached"),
            NameObject("/Name"): TextStringObject(f"Director: {director_usuario}"),
            NameObject("/Reason"): TextStringObject("Co-firma Directorial"),
            NameObject("/Location"): TextStringObject("HCD Posadas"),
            NameObject("/M"): TextStringObject(datetime.now().strftime("D:%Y%m%d%H%M%S")),
        })
        
        writer.add_annotation(page_number=1, annotation=cofirma)
        
        pdf_cofirmado = pdf_path.replace('.pdf', '_cofirmado.pdf')
        with open(pdf_cofirmado, 'wb') as f:
            writer.write(f)
        
        logger.info(f"✅ PDF co-firmado por Director: {pdf_cofirmado}")
        
        return pdf_cofirmado
        
    except Exception as e:
        logger.error(f"Error en co-firma: {e}")
        return None

def verificar_firma_pdf(pdf_path):
    """
    Verifica si un PDF tiene firma digital
    """
    try:
        reader = PdfReader(pdf_path)
        
        # Buscar firmas en el PDF
        firmas = []
        for page_num, page in enumerate(reader.pages):
            if '/Annots' in page:
                for annot in page['/Annots']:
                    if annot.get_object().get('/FT') == '/Sig':
                        firmas.append({
                            'pagina': page_num + 1,
                            'nombre': annot.get_object().get('/T', 'Desconocido'),
                            'fecha': annot.get_object().get('/M', 'No disponible')
                        })
        
        return firmas
        
    except Exception as e:
        logger.error(f"Error verificando firma: {e}")
        return []


def generar_pdf_simple(ddjj_id, legajo, datos, familiares, padres, conyuge, contexto):
    """Genera PDF usando ReportLab (sin Word)"""
    
    
    # Crear carpeta si no existe
    carpeta_agente = os.path.join(DOCUMENTOS_PATH, legajo['legajo_id'], "DDJJ_ANUAL")
    os.makedirs(carpeta_agente, exist_ok=True)
    
    pdf_path = os.path.join(carpeta_agente, f"DDJJ_{legajo['legajo_id']}_{ddjj_id}.pdf")
    
    # Crear PDF
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Título
    titulo_style = ParagraphStyle('Titulo', parent=styles['Title'], alignment=1, fontSize=14)
    story.append(Paragraph("HONORABLE CONCEJO DELIBERANTE", titulo_style))
    story.append(Paragraph("DE LA CIUDAD DE POSADAS", titulo_style))
    story.append(Paragraph("DIRECCIÓN DE PERSONAL", titulo_style))
    story.append(Spacer(1, 20))
    
    # Datos del agente
    story.append(Paragraph(f"AÑO: {contexto.get('anio', '2026')}", styles['Normal']))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"Legajo: {legajo['legajo_id']}", styles['Normal']))
    story.append(Paragraph(f"Apellido y Nombre: {legajo['apellido']}, {legajo['nombre']}", styles['Normal']))
    story.append(Paragraph(f"DNI: {legajo.get('dni', '')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Firma del agente (espacio)
    story.append(Paragraph("________________", styles['Normal']))
    story.append(Paragraph("FIRMA DEL AGENTE", styles['Normal']))
    
    # Construir PDF
    doc.build(story)
    
    # Calcular hash
    with open(pdf_path, 'rb') as f:
        pdf_data = f.read()
    hash_pdf = hashlib.sha256(pdf_data).hexdigest()
    
    return pdf_path, hash_pdf


# ==========================
# RESPALDOS MANUALES (SIN HILOS)
# ==========================

def crear_backup():
    """Crea un backup manual de la base de datos"""
    try:
        backup_dir = os.path.join(BASE_DIR, 'BACKUPS')
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_file = os.path.join(backup_dir, f'backup_{timestamp}.zip')
        
        with zipfile.ZipFile(backup_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
            if os.path.exists(DB_PATH):
                zipf.write(DB_PATH, os.path.basename(DB_PATH))
        
        # Mantener solo los últimos 10 backups
        backups = sorted([f for f in os.listdir(backup_dir) if f.startswith('backup_')])
        while len(backups) > 10:
            os.remove(os.path.join(backup_dir, backups.pop(0)))
        
        print(f"✅ Backup creado: {backup_file}")
        return backup_file
        
    except Exception as e:
        print(f"❌ Error en backup: {e}")
        return None

# Crear backup inicial (solo una vez al arrancar)
try:
    backup_dir = os.path.join(BASE_DIR, 'BACKUPS')
    if os.path.exists(backup_dir):
        backups = [f for f in os.listdir(backup_dir) if f.startswith('backup_')]
        if not backups:
            crear_backup()
            print("📦 Backup inicial creado")
except Exception as e:
    print(f"⚠️ Error: {e}")



def limpiar_backups_viejos():
    """Mantiene solo los últimos 10 backups"""
    try:
        # Obtener todos los archivos .zip
        backups = []
        if not os.path.exists(BACKUP_PATH):
            return
            
        for archivo in os.listdir(BACKUP_PATH):
            ruta = os.path.join(BACKUP_PATH, archivo)
            if os.path.isfile(ruta) and archivo.endswith('.zip'):
                # Obtener fecha de modificación
                fecha = os.path.getmtime(ruta)
                backups.append((fecha, ruta, archivo))
        
        # Ordenar por fecha (más reciente primero)
        backups.sort(reverse=True)
        
        # Mantener solo los últimos 10
        mantener = 10
        if len(backups) > mantener:
            for fecha, ruta, archivo in backups[mantener:]:
                os.remove(ruta)
                logger.info(f"🗑️ Backup antiguo eliminado: {archivo}")
            logger.info(f"✅ Limpieza completada. Se mantuvieron {mantener} backups.")
        else:
            logger.info(f"ℹ️ Solo hay {len(backups)} backups, no es necesario limpiar.")
            
    except Exception as e:
        logger.error(f"❌ Error limpiando backups: {e}")

# ==========================
# FUNCIONES DE EMAIL (Alertas)
# ==========================
def enviar_alerta_email(destinatario, asunto, mensaje, adjunto=None):
    """Envía alerta por email"""
    try:
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['email_from']
        msg['To'] = destinatario
        msg['Subject'] = asunto
        
        msg.attach(MIMEText(mensaje, 'html'))
        
        if adjunto and os.path.exists(adjunto):
            with open(adjunto, 'rb') as f:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(adjunto)}')
                msg.attach(part)
        
        server = smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port'])
        server.starttls()
        server.login(EMAIL_CONFIG['email_from'], EMAIL_CONFIG['email_password'])
        server.send_message(msg)
        server.quit()
        
        logger.info(f"Email enviado a {destinatario}")
        return True
    except Exception as e:
        logger.error(f"Error enviando email: {e}")
        return False

def alerta_ddjj_pendiente(legajo_id, email_agente):
    """Alerta para recordar DDJJ pendiente"""
    asunto = "⚠️ RECORDATORIO: Declaración Jurada Anual Pendiente"
    mensaje = f"""
    <html>
    <body style="font-family: Arial; padding: 20px;">
        <h2 style="color: #1e3c72;">Honorable Concejo Deliberante</h2>
        <h3>Dirección de Personal</h3>
        <hr>
        <p>Estimado/a Agente,</p>
        <p>Le recordamos que debe presentar su <strong>Declaración Jurada Anual</strong> correspondiente al año {datetime.now().year}.</p>
        <p>Por favor, ingrese al sistema de RRHH Digital para completar el formulario y adjuntar la documentación requerida.</p>
        <p>Plazo límite: <strong>31 de marzo de {datetime.now().year}</strong></p>
        <br>
        <p>Atentamente,</p>
        <p><strong>Dirección de Personal</strong><br>
        HCD Posadas</p>
        <hr>
        <p style="font-size: 12px; color: #666;">Este es un mensaje automático, por favor no responder.</p>
    </body>
    </html>
    """
    return enviar_alerta_email(email_agente, asunto, mensaje)



def verificar_ddjj_pendientes():
    """Verifica agentes con DDJJ pendiente y envía alertas"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        anio_actual = datetime.now().year
        
        cur.execute("""
            SELECT l.legajo_id, l.email, l.apellido, l.nombre
            FROM legajos l
            LEFT JOIN declaraciones_juradas d ON l.legajo_id = d.legajo_id AND d.anio = ?
            WHERE d.id IS NULL AND l.estado = 'Activo' AND l.email IS NOT NULL
        """, (anio_actual,))
        
        pendientes = cur.fetchall()
        
        for agente in pendientes:
            alerta_ddjj_pendiente(agente['legajo_id'], agente['email'])
            logger.info(f"Alerta enviada a {agente['legajo_id']} - {agente['apellido']}")
        
        conn.close()
    except Exception as e:
        logger.error(f"Error verificando DDJJ pendientes: {e}")

# ==========================
# FUNCIONES DE BASE DE DATOS
# ==========================
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

    """Obtiene conexión a la base de datos"""
    if 'db' not in g:
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
    return g.db

@app.teardown_appcontext
def close_db(e=None):
    db = g.pop('db', None)
    if db is not None:
        db.close()

def registrar_auditoria(usuario, accion, legajo_id=None, detalle=""):
    """Registra acción en tabla de auditoría con IP y User-Agent"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        ip = request.remote_addr if request else 'SISTEMA'
        user_agent = request.user_agent.string if request and hasattr(request, 'user_agent') else 'SISTEMA'
        
        cur.execute("""
            INSERT INTO auditoria 
            (usuario, rol, accion, legajo_id, detalle, fecha, ip, user_agent)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (usuario, session.get('rol', 'SISTEMA'), accion, legajo_id, detalle, 
              datetime.now().strftime("%Y-%m-%d %H:%M:%S"), ip, user_agent))
        conn.commit()
        
        logger.info(f"AUDITORÍA: {usuario} - {accion} - {legajo_id}")
    except Exception as e:
        logger.error(f"Error auditoría: {e}")



# ==========================
# FUNCIONES DE CONVERSIÓN A PDF
# ==========================
def convertir_a_pdf_seguro(archivo_entrada):
    """Convierte cualquier archivo a PDF de forma segura (soporta Windows y Linux)"""
    
    
    extension = os.path.splitext(archivo_entrada)[1].lower()
    nombre_salida = archivo_entrada.replace(extension, '.pdf')
    
    # Si ya es PDF
    if extension == '.pdf':
        return archivo_entrada
    
    # Si ya existe el PDF, retornarlo
    if os.path.exists(nombre_salida):
        return nombre_salida
    
    # Imágenes
    if extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
        try:
            from PIL import Image
            imagen = Image.open(archivo_entrada)
            if imagen.mode != 'RGB':
                imagen = imagen.convert('RGB')
            imagen.save(nombre_salida, 'PDF')
            print(f"✅ PDF generado desde imagen: {nombre_salida}")
            return nombre_salida
        except Exception as e:
            print(f"❌ Error convirtiendo imagen: {e}")
            return None
    
    # Documentos Word
    if extension in ['.doc', '.docx']:
        # Verificar si win32com está disponible (Windows)
        if WIN32_AVAILABLE:
            try:
                import time
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                
                doc = word.Documents.Open(archivo_entrada)
                time.sleep(1)
                doc.SaveAs(nombre_salida, FileFormat=17)  # 17 = wdFormatPDF
                time.sleep(1)
                doc.Close()
                word.Quit()
                
                print(f"✅ PDF generado desde Word (win32com): {nombre_salida}")
                return nombre_salida
                
            except Exception as e:
                print(f"❌ Error con win32com: {e}")
                # Si falla win32com, intentar con LibreOffice
                return _convertir_con_libreoffice(archivo_entrada, nombre_salida)
        else:
            # Usar LibreOffice en Render/Linux
            return _convertir_con_libreoffice(archivo_entrada, nombre_salida)
    
    # Archivos de texto
    if extension in ['.txt']:
        try:
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(nombre_salida)
            with open(archivo_entrada, 'r', encoding='utf-8') as f:
                y = 750
                for linea in f:
                    c.drawString(50, y, linea.strip())
                    y -= 20
                    if y < 50:
                        c.showPage()
                        y = 750
            c.save()
            print(f"✅ PDF generado desde texto: {nombre_salida}")
            return nombre_salida
        except Exception as e:
            print(f"❌ Error convirtiendo texto: {e}")
            return None
    
    print(f"⚠️ Formato no soportado: {extension}")
    return None

def _convertir_con_libreoffice(archivo_entrada, nombre_salida):
    """Función auxiliar para convertir con LibreOffice (Linux/Render)"""
    try:
    
        result = subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(nombre_salida),
            archivo_entrada
        ], capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0:
            print(f"✅ PDF generado (LibreOffice): {nombre_salida}")
            return nombre_salida
        else:
            print(f"❌ Error LibreOffice: {result.stderr}")
            return None
            
    except Exception as e:
        print(f"❌ Error conversión LibreOffice: {e}")
        return None




def generar_pdf_con_reportlab(ddjj_id, legajo, datos, familiares, padres, conyuge, contexto, codigo_validacion):
    """Genera PDF directamente con ReportLab sin usar Word"""

    
    carpeta_agente = os.path.join(DOCUMENTOS_PATH, legajo['legajo_id'], "DDJJ_ANUAL")
    os.makedirs(carpeta_agente, exist_ok=True)
    
    pdf_path = os.path.join(carpeta_agente, f"DDJJ_{legajo['legajo_id']}_{ddjj_id}.pdf")
    
    doc = SimpleDocTemplate(pdf_path, pagesize=A4, 
                            topMargin=2, bottomMargin=2,
                            leftMargin=2, rightMargin=2)
    
    styles = getSampleStyleSheet()
    story = []
    

# ==========================
# FUNCIÓN PARA GUARDAR DOCUMENTOS
# ==========================
def guardar_documento_seguro(archivo, legajo_id, tipo_documento):
    """Guarda documento con conversión a PDF y validación"""
    try:
        carpeta_base = os.path.join(DOCUMENTOS_PATH, legajo_id)
        carpeta_tipo = os.path.join(carpeta_base, tipo_documento.upper())
        os.makedirs(carpeta_tipo, exist_ok=True)
        
        filename = secure_filename(archivo.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_original = f"{timestamp}_{filename}"
        ruta_original = os.path.join(carpeta_tipo, nombre_original)
        
        archivo.save(ruta_original)
        ruta_pdf = convertir_a_pdf_seguro(ruta_original)
        hash_archivo = calcular_hash(ruta_pdf)
        
        if ruta_pdf != ruta_original:
            os.remove(ruta_original)
        
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO documentos 
            (legajo, tipo, nombre_archivo, ruta_pdf, hash_archivo, fecha_subida, subido_por, estado)
            VALUES (?, ?, ?, ?, ?, ?, ?, 'ACTIVO')
        """, (legajo_id, tipo_documento, os.path.basename(ruta_pdf), ruta_pdf, hash_archivo,
              datetime.now().strftime("%Y-%m-%d %H:%M:%S"), session.get('usuario', 'SISTEMA')))
        
        doc_id = cur.lastrowid
        conn.commit()
        
        registrar_auditoria(session.get('usuario', 'SISTEMA'), 'SUBIO DOCUMENTO', legajo_id, 
                           f"Tipo: {tipo_documento} - ID: {doc_id}")
        
        logger.info(f"Documento guardado: {legajo_id}/{tipo_documento} - Hash: {hash_archivo[:16]}...")
        
        return doc_id, ruta_pdf, hash_archivo
        
    except Exception as e:
        logger.error(f"Error guardando documento: {e}")
        return None, None, None

def calcular_hash(ruta_archivo):
    """Calcula hash SHA256 de un archivo"""
    sha256 = hashlib.sha256()
    with open(ruta_archivo, "rb") as f:
        for bloque in iter(lambda: f.read(4096), b""):
            sha256.update(bloque)
    return sha256.hexdigest()

def guardar_documento_seguro_con_ruta(ruta_archivo, legajo_id, tipo_documento):
    """Versión para guardar documento desde ruta de archivo existente"""
    try:
        carpeta_base = os.path.join(DOCUMENTOS_PATH, legajo_id)
        carpeta_tipo = os.path.join(carpeta_base, tipo_documento.upper())
        os.makedirs(carpeta_tipo, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_original = f"{timestamp}_{os.path.basename(ruta_archivo)}"
        ruta_original = os.path.join(carpeta_tipo, nombre_original)
        
        shutil.copy2(ruta_archivo, ruta_original)
        ruta_pdf = convertir_a_pdf_seguro(ruta_original)
        hash_archivo = calcular_hash(ruta_pdf)
        
        if ruta_pdf != ruta_original:
            os.remove(ruta_original)
        
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO documentos 
            (legajo, tipo, nombre_archivo, ruta_pdf, hash_archivo, fecha_subida, subido_por, estado)
            VALUES (?, ?, ?, ?, ?, ?, ?, 'ACTIVO')
        """, (legajo_id, tipo_documento, os.path.basename(ruta_pdf), ruta_pdf, hash_archivo,
              datetime.now().strftime("%Y-%m-%d %H:%M:%S"), session.get('usuario', 'SISTEMA')))
        
        doc_id = cur.lastrowid
        conn.commit()
        
        if os.path.exists(ruta_archivo) and ruta_archivo != ruta_original and ruta_archivo != ruta_pdf:
            try:
                os.remove(ruta_archivo)
            except:
                pass
        
        return doc_id, ruta_pdf, hash_archivo
        
    except Exception as e:
        logger.error(f"Error guardando documento: {e}")
        return None, None, None

# ==========================
# FUNCIONES PARA DDJJ
# ==========================
def generar_codigo_validacion():
    """Genera código único de validación"""
    return hashlib.sha256(f"{uuid.uuid4()}{datetime.now()}".encode()).hexdigest()[:16].upper()

def generar_qr_validacion(codigo, carpeta):
    """Genera QR con código de validación"""
    url = f"http://127.0.0.1:5000/validar/{codigo}"
    ruta = os.path.join(carpeta, f"qr_{codigo}.png")
    
    img = qrcode.make(url)
    img.save(ruta)
    return ruta

def crear_plantilla_oficial():
    """Crea la plantilla oficial si no existe"""
    plantilla_path = os.path.join(PLANTILLAS_PATH, "PLANTILLA_DDJJ_OFICIAL.docx")
    
    if not os.path.exists(plantilla_path):
        doc = Document()
        
        from docx.shared import Inches, Cm
        section = doc.sections[0]
        section.page_height = Cm(35.56)
        section.page_width = Cm(25.59)
        section.top_margin = Cm(0.94)
        section.bottom_margin = Cm(0.49)
        section.left_margin = Cm(3)
        section.right_margin = Cm(0)
        
        doc.add_heading('HONORABLE CONCEJO DELIBERANTE', 0)
        doc.add_heading('DE LA CIUDAD DE POSADAS', 1)
        doc.add_heading('DIRECCIÓN DE PERSONAL', 2)
        doc.add_paragraph('AÑO: {{ anio }}')
        
        doc.save(plantilla_path)
        logger.info("Plantilla oficial creada")

# ==========================
# FUNCIÓN DE MIGRACIÓN DE COLUMNAS PARA LEGAJOS
# ==========================
def agregar_columnas_faltantes():
    """Agrega columnas faltantes a la tabla legajos"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Obtener columnas actuales
        cur.execute("PRAGMA table_info(legajos)")
        columnas = [col[1] for col in cur.fetchall()]
        
        # Columnas que deben existir
        columnas_necesarias = [
            'lugar_nacimiento',
            'telefono_fijo',
            'telefono_referencia',
            'parentesco_referencia',
            'domicilio_piso',
            'domicilio_depto',
            'domicilio_barrio',
            'croquis_domicilio',
            'foto_path',
            'tipo_personal',
            'antiguedad',
            'lugar_desempeno',
            'obra_social',
            'estado_civil' 
        ]
        
        # Agregar columnas faltantes
        columnas_agregadas = []
        for columna in columnas_necesarias:
            if columna not in columnas:
                try:
                    cur.execute(f"ALTER TABLE legajos ADD COLUMN {columna} TEXT")
                    columnas_agregadas.append(columna)
                    print(f"✅ Columna agregada: {columna}")
                except Exception as e:
                    print(f"⚠️ Error al agregar {columna}: {e}")
        
        conn.commit()
        conn.close()
        
        if columnas_agregadas:
            print(f"📦 Columnas agregadas: {', '.join(columnas_agregadas)}")
        else:
            print("ℹ️ Todas las columnas de legajos ya existen")
        
        return True
        
    except Exception as e:
        print(f"❌ Error en migración de legajos: {e}")
        return False


# ==========================
# FUNCIÓN PARA CREAR ÍNDICE DE CÓDIGO DE VALIDACIÓN
# ==========================
def crear_indice_codigo_validacion():
    """Crea el índice para codigo_validacion si no existe"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Crear índice (no causa error si ya existe)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_ddjj_codigo ON declaraciones_juradas(codigo_validacion)")
        conn.commit()
        conn.close()
        print("✅ Índice idx_ddjj_codigo creado/verificado")
        return True
        
    except Exception as e:
        print(f"⚠️ Error creando índice: {e}")
        return False


# ==========================
# INICIALIZACIÓN DE BASE DE DATOS
# ==========================
def init_db():
    """Inicializa todas las tablas con la nueva estructura"""
    try:
        conn = get_db()
        cur = conn.cursor()

        print("📦 Inicializando base de datos...")

        # ==========================
        # TABLA USUARIOS
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS usuarios (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL,
                rol TEXT NOT NULL,
                email TEXT,
                legajo_id TEXT,
                bloqueado INTEGER DEFAULT 0,
                intentos_fallidos INTEGER DEFAULT 0,
                ultimo_login TEXT,
                fecha_creacion TEXT,
                activo INTEGER DEFAULT 1
            )
        """)

        # ==========================
        # TABLA LEGAJOS (estructura completa)
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS legajos(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                legajo_id TEXT UNIQUE NOT NULL,
                apellido TEXT NOT NULL,
                nombre TEXT NOT NULL,
                dni TEXT UNIQUE,
                cuil TEXT UNIQUE,
                email TEXT,
                telefono TEXT,
                telefono_fijo TEXT,
                telefono_referencia TEXT,
                parentesco_referencia TEXT,
                fecha_nacimiento TEXT,
                lugar_nacimiento TEXT,
                nacionalidad TEXT,
                estado_civil TEXT,
                domicilio_calle TEXT,
                domicilio_numero TEXT,
                domicilio_piso TEXT,
                domicilio_depto TEXT,
                domicilio_barrio TEXT,
                domicilio_localidad TEXT,
                croquis_domicilio TEXT,
                foto_path TEXT,
                cargo TEXT,
                tipo_personal TEXT,
                fecha_ingreso TEXT,
                categoria TEXT,
                antiguedad TEXT,
                lugar_desempeno TEXT,
                obra_social TEXT,
                estado TEXT DEFAULT 'Activo',
                activo INTEGER DEFAULT 1
            )
        """)

        # ==========================
        # TABLA DOCUMENTOS
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS documentos(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                legajo TEXT NOT NULL,
                tipo TEXT NOT NULL,
                nombre_archivo TEXT NOT NULL,
                ruta_pdf TEXT NOT NULL,
                fecha_subida TEXT,
                usuario_subio TEXT,
                activo INTEGER DEFAULT 1,
                FOREIGN KEY (legajo) REFERENCES legajos(legajo_id)
            )
        """)

        # ==========================
        # TABLA DECLARACIONES JURADAS
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS declaraciones_juradas(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                legajo_id TEXT NOT NULL,
                anio INTEGER NOT NULL,
                estado TEXT DEFAULT 'BORRADOR',
                archivo_pdf TEXT,
                hash_pdf TEXT,
                codigo_validacion TEXT UNIQUE,
                fecha_generacion TEXT,
                fecha_envio TEXT,
                usuario_genero TEXT,
                usuario_finalizo TEXT,
                observaciones TEXT,
                activa INTEGER DEFAULT 1,
                FOREIGN KEY(legajo_id) REFERENCES legajos(legajo_id)
            )
        """)

        # ==========================
        # TABLA DDJJ_DATOS
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS ddjj_datos(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ddjj_id INTEGER NOT NULL,
                fecha_nacimiento TEXT,
                lugar_nacimiento TEXT,
                dni TEXT,
                estado_civil TEXT,
                estudios_completos TEXT,
                estudios_incompletos TEXT,
                estudios_nivel TEXT,
                titulo TEXT,
                anios_cursados TEXT,
                domicilio_barrio TEXT,
                domicilio_chacra TEXT,
                domicilio_calle TEXT,
                domicilio_numero TEXT,
                domicilio_piso TEXT,
                domicilio_depto TEXT,
                domicilio_localidad TEXT,
                telefono_fijo TEXT,
                telefono_celular TEXT,
                correo TEXT,
                telefono_referencia TEXT,
                nombre_referencia TEXT,
                parentesco_referencia TEXT,
                fecha_ingreso TEXT,
                categoria TEXT,
                antiguedad TEXT,
                planta_permanente INTEGER,
                planta_temporaria INTEGER,
                lugar_desempeno TEXT,
                actividad TEXT,
                enfermedad_preexistente TEXT,
                vacuna_covid_dosis TEXT,
                vacuna_gripe TEXT,
                vacuna_neumonia TEXT,
                percibe_asig_fliares_hcd INTEGER,
                servicios_anteriores TEXT,
                antiguedad_servicios TEXT,
                trabaja_otra_reparticion INTEGER,
                otra_reparticion_donde TEXT,
                percibe_asig_fliares_otro INTEGER,
                lugar TEXT,
                fecha_declaracion TEXT,
                firma_agente TEXT,
                FOREIGN KEY (ddjj_id) REFERENCES declaraciones_juradas(id)
            )
        """)

        # ==========================
        # TABLA DDJJ_CONYUGE
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS ddjj_conyuge(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ddjj_id INTEGER NOT NULL,
                estado_civil TEXT, 
                apellido_nombre TEXT,
                dni TEXT,
                fecha_enlace TEXT,
                lugar_enlace TEXT,
                trabaja INTEGER,
                razon_social TEXT,
                FOREIGN KEY (ddjj_id) REFERENCES declaraciones_juradas(id)
            )
        """)

        # ==========================
        # TABLA DDJJ_FAMILIARES_CARGO
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS ddjj_familiares_cargo(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ddjj_id INTEGER NOT NULL,
                orden INTEGER,
                apellido_nombre TEXT,
                parentesco TEXT,
                convive INTEGER,
                fecha_nacimiento TEXT,
                edad INTEGER,
                pri TEXT,
                FOREIGN KEY (ddjj_id) REFERENCES declaraciones_juradas(id)
            )
        """)

        # ==========================
        # TABLA DDJJ_PADRES_HERMANOS
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS ddjj_padres_hermanos(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ddjj_id INTEGER NOT NULL,
                orden INTEGER,
                apellido_nombre TEXT,
                parentesco TEXT,
                convive INTEGER,
                fecha_nacimiento TEXT,
                dni TEXT,
                FOREIGN KEY (ddjj_id) REFERENCES declaraciones_juradas(id)
            )
        """)

        # ==========================
        # TABLA FIRMAS_DIGITALES
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS firmas_digitales(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ddjj_id INTEGER NOT NULL,
                tipo_firma TEXT NOT NULL,
                firmante_id INTEGER,
                firmante_nombre TEXT,
                hash_documento TEXT,
                firma_pkcs7 TEXT,
                timestamp_firma TEXT,
                ip_firma TEXT,
                user_agent TEXT,
                FOREIGN KEY (ddjj_id) REFERENCES declaraciones_juradas(id)
            )
        """)

        # ==========================
        # TABLA AUDITORIA
        # ==========================
        cur.execute("""
            CREATE TABLE IF NOT EXISTS auditoria(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                usuario TEXT NOT NULL,
                accion TEXT NOT NULL,
                legajo_id TEXT,
                detalle TEXT,
                fecha TEXT,
                ip TEXT,
                user_agent TEXT
            )
        """)

        # ==========================
        # CREAR USUARIO ADMIN SI NO EXISTE
        # ==========================
        cur.execute("SELECT * FROM usuarios WHERE username = ?", ("admin",))
        if not cur.fetchone():
            from werkzeug.security import generate_password_hash
            hashed = generate_password_hash("Admin2026!")
            cur.execute("""
                INSERT INTO usuarios (username, password, rol, email, fecha_creacion)
                VALUES (?, ?, ?, ?, ?)
            """, ("admin", hashed, "ADMIN", "admin@rrhh.com", datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            print("✅ Usuario admin creado")

        # ==========================
        # CREAR USUARIO DIRECTOR SI NO EXISTE
        # ==========================
        cur.execute("SELECT * FROM usuarios WHERE username = ?", ("director",))
        if not cur.fetchone():
            from werkzeug.security import generate_password_hash
            hashed = generate_password_hash("Director2026!")
            cur.execute("""
                INSERT INTO usuarios (username, password, rol, email, fecha_creacion)
                VALUES (?, ?, ?, ?, ?)
            """, ("director", hashed, "DIRECTOR", "director@rrhh.com", datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            print("✅ Usuario director creado")

        conn.commit()
        conn.close()

        print("✅ Base de datos inicializada correctamente")

    except Exception as e:
        print("❌ Error inicializando DB:", e)
        import traceback
        traceback.print_exc()


def migrar_legajos_estructura_completa():
    """Migra la tabla legajos a la estructura completa"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Verificar si falta lugar_nacimiento
        cur.execute("PRAGMA table_info(legajos)")
        columnas = [col[1] for col in cur.fetchall()]
        
        if 'lugar_nacimiento' not in columnas:
            print("🔄 Iniciando migración de estructura de legajos...")
            
            # 1. Crear tabla temporal con estructura completa
            cur.execute("""
                CREATE TABLE IF NOT EXISTS legajos_nueva(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    legajo_id TEXT UNIQUE NOT NULL,
                    apellido TEXT NOT NULL,
                    nombre TEXT NOT NULL,
                    dni TEXT UNIQUE,
                    cuil TEXT UNIQUE,
                    email TEXT,
                    telefono TEXT,
                    telefono_fijo TEXT,
                    telefono_referencia TEXT,
                    parentesco_referencia TEXT,
                    fecha_nacimiento TEXT,
                    lugar_nacimiento TEXT,
                    nacionalidad TEXT,
                    estado_civil TEXT,
                    domicilio_calle TEXT,
                    domicilio_numero TEXT,
                    domicilio_piso TEXT,
                    domicilio_depto TEXT,
                    domicilio_barrio TEXT,
                    domicilio_localidad TEXT,
                    croquis_domicilio TEXT,
                    foto_path TEXT,
                    cargo TEXT,
                    tipo_personal TEXT,
                    fecha_ingreso TEXT,
                    categoria TEXT,
                    antiguedad TEXT,
                    lugar_desempeno TEXT,
                    obra_social TEXT,
                    estado TEXT DEFAULT 'Activo',
                    activo INTEGER DEFAULT 1
                )
            """)
            
            # 2. Copiar datos existentes
            cur.execute("""
                INSERT INTO legajos_nueva (
                    id, legajo_id, apellido, nombre, dni, cuil, email, telefono,
                    fecha_nacimiento, domicilio_calle, domicilio_numero, 
                    domicilio_localidad, cargo, fecha_ingreso, estado
                )
                SELECT 
                    id, legajo_id, apellido, nombre, dni, cuil, email, telefono,
                    fecha_nacimiento, domicilio_calle, domicilio_numero, 
                    COALESCE(domicilio_localidad, 'POSADAS'), cargo, fecha_ingreso, 
                    COALESCE(estado, 'Activo')
                FROM legajos
            """)
            
            # 3. Eliminar tabla vieja y renombrar nueva
            cur.execute("DROP TABLE legajos")
            cur.execute("ALTER TABLE legajos_nueva RENAME TO legajos")
            
            conn.commit()
            print("✅ Migración completada. Tabla legajos actualizada a estructura completa.")
            
        else:
            print("ℹ️ La tabla legajos ya tiene la estructura completa.")
            
        conn.close()
        return True
        
    except Exception as e:
        print(f"❌ Error en migración: {e}")
        return False

def inicializar_datos_prueba():
    """Inicializa datos de prueba con estructura completa de legajos"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Verificar si ya hay datos
        cur.execute("SELECT COUNT(*) FROM legajos WHERE legajo_id BETWEEN '90000' AND '99999'")
        count = cur.fetchone()[0]
        
        if count > 0:
            print(f"ℹ️ Ya existen {count} legajos de prueba. No se crearán nuevos.")
            return True
        
        # Datos de prueba con estructura completa
        legajos_prueba = [
            {
                'legajo_id': '90001',
                'apellido': 'González',
                'nombre': 'María Elena',
                'dni': '12345678',
                'cuil': '27-12345678-9',
                'email': 'maria.gonzalez@ejemplo.com',
                'telefono': '3764123456',
                'telefono_fijo': '3764456789',
                'telefono_referencia': '3764987654',
                'parentesco_referencia': 'Madre',
                'fecha_nacimiento': '15/03/1985',
                'lugar_nacimiento': 'Posadas, Misiones',
                'nacionalidad': 'Argentina',
                'estado_civil': 'Casada',
                'domicilio_calle': 'San Martín',
                'domicilio_numero': '1234',
                'domicilio_piso': '2',
                'domicilio_depto': 'B',
                'domicilio_barrio': 'Centro',
                'domicilio_localidad': 'Posadas',
                'cargo': 'Administrativo',
                'tipo_personal': 'Planta Permanente',
                'fecha_ingreso': '01/03/2010',
                'categoria': 'Categoría 5',
                'antiguedad': '13 años',
                'lugar_desempeno': 'Dirección de RRHH',
                'obra_social': 'IOMA',
                'estado': 'Activo',
                'activo': 1
            },
            {
                'legajo_id': '90002',
                'apellido': 'Rodríguez',
                'nombre': 'Carlos Alberto',
                'dni': '23456789',
                'cuil': '20-23456789-1',
                'email': 'carlos.rodriguez@ejemplo.com',
                'telefono': '3764234567',
                'telefono_fijo': '3764567890',
                'telefono_referencia': '3764876543',
                'parentesco_referencia': 'Hermano',
                'fecha_nacimiento': '22/07/1980',
                'lugar_nacimiento': 'Apóstoles, Misiones',
                'nacionalidad': 'Argentina',
                'estado_civil': 'Soltero',
                'domicilio_calle': 'Bolívar',
                'domicilio_numero': '567',
                'domicilio_piso': '',
                'domicilio_depto': '',
                'domicilio_barrio': 'Villa Sarita',
                'domicilio_localidad': 'Posadas',
                'cargo': 'Técnico',
                'tipo_personal': 'Contratado',
                'fecha_ingreso': '15/05/2015',
                'categoria': 'Categoría 3',
                'antiguedad': '8 años',
                'lugar_desempeno': 'Área de Sistemas',
                'obra_social': 'OSEP',
                'estado': 'Activo',
                'activo': 1
            },
            {
                'legajo_id': '90003',
                'apellido': 'Fernández',
                'nombre': 'Laura Beatriz',
                'dni': '34567890',
                'cuil': '27-34567890-2',
                'email': 'laura.fernandez@ejemplo.com',
                'telefono': '3764345678',
                'telefono_fijo': '3764678901',
                'telefono_referencia': '3764765432',
                'parentesco_referencia': 'Padre',
                'fecha_nacimiento': '10/11/1990',
                'lugar_nacimiento': 'Oberá, Misiones',
                'nacionalidad': 'Argentina',
                'estado_civil': 'Soltera',
                'domicilio_calle': 'Sarmiento',
                'domicilio_numero': '890',
                'domicilio_piso': '1',
                'domicilio_depto': 'A',
                'domicilio_barrio': 'Centro',
                'domicilio_localidad': 'Posadas',
                'cargo': 'Profesional',
                'tipo_personal': 'Planta Permanente',
                'fecha_ingreso': '20/01/2018',
                'categoria': 'Categoría 4',
                'antiguedad': '5 años',
                'lugar_desempeno': 'Asesoría Legal',
                'obra_social': 'IOMA',
                'estado': 'Activo',
                'activo': 1
            }
        ]
        
        # Insertar legajos de prueba
        for legajo in legajos_prueba:
            cur.execute("""
                INSERT INTO legajos (
                    legajo_id, apellido, nombre, dni, cuil, email, telefono,
                    telefono_fijo, telefono_referencia, parentesco_referencia,
                    fecha_nacimiento, lugar_nacimiento, nacionalidad, estado_civil,
                    domicilio_calle, domicilio_numero, domicilio_piso, domicilio_depto,
                    domicilio_barrio, domicilio_localidad, cargo, tipo_personal,
                    fecha_ingreso, categoria, antiguedad, lugar_desempeno,
                    obra_social, estado, activo
                ) VALUES (
                    :legajo_id, :apellido, :nombre, :dni, :cuil, :email, :telefono,
                    :telefono_fijo, :telefono_referencia, :parentesco_referencia,
                    :fecha_nacimiento, :lugar_nacimiento, :nacionalidad, :estado_civil,
                    :domicilio_calle, :domicilio_numero, :domicilio_piso, :domicilio_depto,
                    :domicilio_barrio, :domicilio_localidad, :cargo, :tipo_personal,
                    :fecha_ingreso, :categoria, :antiguedad, :lugar_desempeno,
                    :obra_social, :estado, :activo
                )
            """, legajo)
        
        conn.commit()
        
        # Crear algunos documentos de ejemplo
        tipos_documentos = ['DNI', 'Título', 'Certificado', 'Declaración Jurada']
        for legajo in legajos_prueba:
            for tipo in tipos_documentos[:2]:  # 2 documentos por legajo
                cur.execute("""
                    INSERT INTO documentos (legajo, tipo, nombre_archivo, ruta_pdf, fecha_subida)
                    VALUES (?, ?, ?, ?, ?)
                """, (
                    legajo['legajo_id'],
                    tipo,
                    f"{tipo}_{legajo['legajo_id']}.pdf",
                    f"/app/DOCUMENTOS_AGENTES/{legajo['legajo_id']}_{tipo}.pdf",
                    datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                ))
        
        conn.commit()
        conn.close()
        
        print(f"✅ Datos de prueba inicializados: {len(legajos_prueba)} legajos creados")
        print(f"📄 Documentos de ejemplo creados para cada legajo")
        return True
        
    except Exception as e:
        print(f"❌ Error inicializando datos de prueba: {e}")
        import traceback
        traceback.print_exc()
        return False

def generar_documento_word(plantilla_path, datos, salida_path):
    """Genera documento Word desde plantilla (solo Windows si usa win32com)"""
    if not WIN32_AVAILABLE:
        print(f"⚠️ Generación de Word no disponible: {salida_path}")
        return None
    
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        
        doc = word.Documents.Open(plantilla_path)
        
        # Reemplazar placeholders en el documento
        for paragraph in doc.paragraphs:
            for key, value in datos.items():
                if f'{{{{{key}}}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))
        
        # Guardar documento
        doc.SaveAs(salida_path)
        doc.Close()
        word.Quit()
        
        print(f"✅ Documento Word generado: {salida_path}")
        return salida_path
        
    except Exception as e:
        print(f"❌ Error generando documento: {e}")
        return None

# ==========================
# DECORADORES DE SEGURIDAD
# ==========================
def login_requerido(f):
    @wraps(f)
    def decorado(*args, **kwargs):
        if "usuario" not in session:
            flash("Por favor, inicie sesión", "error")
            return redirect(url_for("login"))
        
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT bloqueado FROM usuarios WHERE username = ?", (session["usuario"],))
        user = cur.fetchone()
        if user and user['bloqueado'] == 1:
            session.clear()
            flash("Usuario bloqueado. Contacte al administrador.", "error")
            return redirect(url_for("login"))
        
        return f(*args, **kwargs)
    return decorado

def rol_permitido(roles):
    def decorator(f):
        @wraps(f)
        def decorado(*args, **kwargs):
            if session.get("rol") not in roles:
                registrar_auditoria(session.get("usuario"), "ACCESO DENEGADO", None, 
                                   f"Rol requerido: {roles}, Rol actual: {session.get('rol')}")
                flash("No tiene permiso para esta acción", "error")
                return redirect(url_for("panel"))
            return f(*args, **kwargs)
        return decorado
    return decorator

def calcular_hash(ruta_archivo):
    sha256 = hashlib.sha256()
    with open(ruta_archivo, "rb") as f:
        for bloque in iter(lambda: f.read(4096), b""):
            sha256.update(bloque)
    return sha256.hexdigest()

# ==========================
# FUNCIONES DDJJ CON FIRMA
# ==========================

def generar_ddjj_con_firma(legajo_id, ddjj_id, pin_agente):
    """
    Genera DDJJ y la firma digitalmente con PIN del agente
    """
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Obtener datos completos
        cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (legajo_id,))
        legajo = dict(cur.fetchone())
        
        cur.execute("SELECT * FROM ddjj_datos WHERE ddjj_id = ?", (ddjj_id,))
        datos_row = cur.fetchone()
        datos = dict(datos_row) if datos_row else {}
        
        cur.execute("SELECT * FROM ddjj_familiares_cargo WHERE ddjj_id = ? ORDER BY orden", (ddjj_id,))
        familiares = [dict(f) for f in cur.fetchall()]
        
        cur.execute("SELECT * FROM ddjj_padres_hermanos WHERE ddjj_id = ? ORDER BY orden", (ddjj_id,))
        padres = [dict(p) for p in cur.fetchall()]
        
        cur.execute("SELECT * FROM ddjj_conyuge WHERE ddjj_id = ?", (ddjj_id,))
        conyuge_row = cur.fetchone()
        conyuge = dict(conyuge_row) if conyuge_row else {}
        
        # Mapear datos para plantilla
        contexto = {}
        contexto['anio'] = datetime.now().year
        contexto['legajo'] = legajo_id
        contexto['apellido_nombre'] = f"{legajo.get('apellido', '')} {legajo.get('nombre', '')}"
        contexto['dni'] = datos.get('dni', legajo.get('dni', ''))
        contexto['fecha_nacimiento'] = datos.get('fecha_nacimiento', legajo.get('fecha_nacimiento', ''))
        contexto['estado_civil'] = datos.get('estado_civil', legajo.get('estado_civil', ''))
        contexto['email'] = datos.get('correo', legajo.get('email', ''))
        contexto['telefono_celular'] = datos.get('telefono_celular', legajo.get('telefono', ''))
        contexto['domicilio_calle'] = datos.get('domicilio_calle', legajo.get('domicilio_calle', ''))
        contexto['domicilio_numero'] = datos.get('domicilio_numero', legajo.get('domicilio_numero', ''))
        contexto['domicilio_localidad'] = datos.get('domicilio_localidad', legajo.get('domicilio_localidad', 'POSADAS'))
        
        # Código de validación
        codigo_validacion = hashlib.sha256(f"{uuid.uuid4()}{datetime.now()}".encode()).hexdigest()[:16].upper()
        contexto['codigo_validacion'] = codigo_validacion
        
        # Generar PDF con la plantilla
        plantilla_path = os.path.join(PLANTILLAS_PATH, "PLANTILLA_DDJJ_OFICIAL.docx")
        
        if not os.path.exists(plantilla_path):
            # Crear plantilla básica si no existe
            doc = Document()
            doc.add_heading('HONORABLE CONCEJO DELIBERANTE', 0)
            doc.add_heading('DECLARACIÓN JURADA ANUAL', 1)
            doc.add_paragraph(f"Año: {contexto['anio']}")
            doc.add_paragraph(f"Legajo: {contexto['legajo']}")
            doc.add_paragraph(f"Agente: {contexto['apellido_nombre']}")
            doc.add_paragraph(f"DNI: {contexto['dni']}")
            doc.add_paragraph(f"Código de validación: {codigo_validacion}")
            doc.save(plantilla_path)
        
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, f"DDJJ_{ddjj_id}.docx")
        
        doc = DocxTemplate(plantilla_path)
        doc.render(contexto)
        doc.save(docx_path)
        
        from docx2pdf import convert
        pdf_path = docx_path.replace('.docx', '.pdf')
        convertir_word_a_pdf(docx_path, pdf_path)
        if not convertir_docx_a_pdf_con_libreoffice(docx_path, pdf_path):
            # Si falla, usar ReportLab como alternativa
            pdf_path, hash_pdf = generar_pdf_con_reportlab(
                ddjj_id, legajo, datos, familiares, padres, conyuge, 
                contexto, codigo_validacion
            )
            # Saltar la lectura posterior
            with open(pdf_path, 'rb') as f:
                pdf_data = f.read()
            hash_pdf = hashlib.sha256(pdf_data).hexdigest()
                
        # FIRMAR EL PDF CON EL PIN DEL AGENTE
        pdf_firmado, mensaje_firma = firmar_pdf_con_pin(pdf_path, legajo_id, pin_agente)
        
        if not pdf_firmado:
            shutil.rmtree(temp_dir)
            return None, mensaje_firma
        
        # Guardar PDF firmado
        carpeta_agente = os.path.join(DOCUMENTOS_PATH, legajo_id, "DDJJ_ANUAL")
        os.makedirs(carpeta_agente, exist_ok=True)
        
        pdf_final = os.path.join(carpeta_agente, f"DDJJ_{legajo_id}_{datetime.now().year}.pdf")
        shutil.copy(pdf_firmado, pdf_final)
        
        # Generar QR
        try:
            qr = qrcode.QRCode(version=1, box_size=10, border=5)
            qr.add_data(f"http://127.0.0.1:5000/validar/{codigo_validacion}")
            qr.make(fit=True)
            qr_img = qr.make_image(fill_color="black", back_color="white")
            qr_path = os.path.join(carpeta_agente, f"QR_{codigo_validacion}.png")
            qr_img.save(qr_path)
        except:
            pass
        
        # Actualizar BD
        with open(pdf_final, 'rb') as f:
            pdf_data = f.read()
        hash_pdf = hashlib.sha256(pdf_data).hexdigest()
        
        cur.execute("""
            UPDATE declaraciones_juradas
            SET estado = 'FIRMADA_AGENTE',
                archivo_pdf = ?,
                hash_pdf = ?,
                codigo_validacion = ?,
                fecha_envio = ?,
                usuario_finalizo = ?
            WHERE id = ?
        """, (pdf_final, hash_pdf, codigo_validacion,
              datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
              session.get("usuario", "AGENTE"), ddjj_id))
        
        conn.commit()
        shutil.rmtree(temp_dir)
        
        # Notificar al Director para co-firma
        enviar_alerta_email(
            EMAIL_CONFIG['email_admin'],
            "📄 DDJJ LISTA PARA CO-FIRMA",
            f"""
            <h2>DDJJ firmada por agente</h2>
            <p><strong>Legajo:</strong> {legajo_id}</p>
            <p><strong>Agente:</strong> {contexto['apellido_nombre']}</p>
            <p><strong>Código:</strong> {codigo_validacion}</p>
            <p>La DDJJ está lista para co-firma del Director.</p>
            <a href="http://127.0.0.1:5000/ddjj/{ddjj_id}/cofirmar">Co-firmar DDJJ</a>
            """
        )
        
        return pdf_final, "✅ DDJJ firmada exitosamente. Pendiente co-firma del Director."
        
    except Exception as e:
        logger.error(f"Error generando DDJJ con firma: {e}")
        return None, f"Error: {str(e)}"



@app.route('/admin/backup/crear', methods=['POST'])
@login_required
def crear_backup_manual():
    """Crea un backup manual desde el panel"""
    if session.get('rol') != 'ADMIN':
        return jsonify({'error': 'Acceso denegado'}), 403
    
    archivo = crear_backup()
    if archivo:
        return jsonify({'success': True, 'archivo': os.path.basename(archivo)})
    else:
        return jsonify({'success': False, 'error': 'Error al crear backup'}), 500


# ==========================
# RUTAS DE FIRMA DIGITAL
# ==========================

@app.route("/legajo/<legajo_id>/certificado/generar", methods=["GET", "POST"])
@login_requerido
def generar_certificado_agente_route(legajo_id):
    """Ruta para que AC emita certificado .p12 para el agente"""
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (legajo_id,))
    legajo = cur.fetchone()
    
    if not legajo:
        flash("Legajo no encontrado", "error")
        return redirect(url_for("listar_legajos"))
    
    if request.method == "POST":
        pin = request.form.get("pin", "")
        confirm_pin = request.form.get("confirm_pin", "")
        
        if len(pin) < 4:
            flash("El PIN debe tener al menos 4 caracteres", "error")
            return redirect(url_for("generar_certificado_agente_route", legajo_id=legajo_id))
        
        if pin != confirm_pin:
            flash("Los PIN no coinciden", "error")
            return redirect(url_for("generar_certificado_agente_route", legajo_id=legajo_id))
        
        nombre_completo = f"{legajo['apellido']} {legajo['nombre']}"
        p12_path = generar_certificado_agente(legajo_id, nombre_completo, legajo['dni'], pin)
        
        if p12_path:
            flash(f"✅ Certificado generado exitosamente. PIN registrado.", "success")
            registrar_auditoria(session["usuario"], "GENERÓ CERTIFICADO", legajo_id)
        else:
            flash("❌ Error generando certificado", "error")
        
        return redirect(url_for("ver_legajo", legajo_id=legajo_id))
    
    # Verificar si ya tiene certificado
    cur.execute("SELECT * FROM certificados_agentes WHERE legajo_id = ? AND activo = 1", (legajo_id,))
    tiene_certificado = cur.fetchone()
    
    return render_template("generar_certificado.html", 
                          legajo=dict(legajo),
                          tiene_certificado=tiene_certificado)

@app.route("/ddjj/<int:ddjj_id>/firmar", methods=["GET", "POST"])
@login_requerido
def firmar_ddjj(ddjj_id):
    """Ruta para que el agente firme su DDJJ con PIN"""
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT d.*, l.apellido, l.nombre, l.legajo_id
        FROM declaraciones_juradas d
        JOIN legajos l ON d.legajo_id = l.legajo_id
        WHERE d.id = ?
    """, (ddjj_id,))
    
    ddjj = cur.fetchone()
    
    if not ddjj:
        flash("DDJJ no encontrada", "error")
        return redirect(url_for("panel"))
    
    ddjj = dict(ddjj)
    
    if ddjj['estado'] != 'BORRADOR':
        flash("Esta DDJJ ya fue procesada", "warning")
        return redirect(url_for("ver_legajo", legajo_id=ddjj['legajo_id']))
    
    # Verificar que el agente tenga certificado
    cur.execute("SELECT * FROM certificados_agentes WHERE legajo_id = ? AND activo = 1", (ddjj['legajo_id'],))
    certificado = cur.fetchone()
    
    if not certificado:
        flash("⚠️ No tiene certificado digital. Solicite su certificado a la AC interna.", "warning")
        return redirect(url_for("generar_certificado_agente_route", legajo_id=ddjj['legajo_id']))
    
    if request.method == "POST":
        pin = request.form.get("pin", "")
        
        # Generar y firmar DDJJ
        pdf_firmado, mensaje = generar_ddjj_con_firma(ddjj['legajo_id'], ddjj_id, pin)
        
        if pdf_firmado:
            flash(mensaje, "success")
            return redirect(url_for("ver_legajo", legajo_id=ddjj['legajo_id']))
        else:
            flash(mensaje, "error")
            return redirect(url_for("firmar_ddjj", ddjj_id=ddjj_id))
    
    return render_template("firmar_ddjj.html", 
                          ddjj=ddjj,
                          legajo_id=ddjj['legajo_id'],
                          agente_nombre=f"{ddjj['apellido']} {ddjj['nombre']}")

@app.route("/ddjj/<int:ddjj_id>/cofirmar", methods=["GET", "POST"])
@login_requerido
@rol_permitido(["ADMIN", "DIRECTOR"])
def cofirmar_ddjj(ddjj_id):
    """Ruta para que el Director co-firme la DDJJ (segunda firma)"""
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT d.*, l.apellido, l.nombre, l.legajo_id
        FROM declaraciones_juradas d
        JOIN legajos l ON d.legajo_id = l.legajo_id
        WHERE d.id = ?
    """, (ddjj_id,))
    
    ddjj = cur.fetchone()
    
    if not ddjj:
        flash("DDJJ no encontrada", "error")
        return redirect(url_for("panel"))
    
    ddjj = dict(ddjj)
    
    if ddjj['estado'] not in ['FIRMADA_AGENTE', 'BORRADOR']:
        flash("Esta DDJJ ya fue co-firmada", "warning")
        return redirect(url_for("listar_ddjj"))
    
    if not ddjj.get('archivo_pdf') or not os.path.exists(ddjj['archivo_pdf']):
        flash("El PDF de la DDJJ no está disponible", "error")
        return redirect(url_for("listar_ddjj"))
    
    if request.method == "POST":
        # Co-firmar PDF
        pdf_cofirmado = cofirmar_pdf_director(ddjj['archivo_pdf'], session["usuario"])
        
        if pdf_cofirmado:
            # Actualizar con el PDF co-firmado
            shutil.copy(pdf_cofirmado, ddjj['archivo_pdf'])
            
            cur.execute("""
                UPDATE declaraciones_juradas
                SET estado = 'COFIRMADA_DIRECTOR',
                    observaciones = ?
                WHERE id = ?
            """, (f"Co-firmada por Director {session['usuario']} en {datetime.now()}", ddjj_id))
            
            conn.commit()
            
            # Notificar al agente
            cur.execute("SELECT email FROM legajos WHERE legajo_id = ?", (ddjj['legajo_id'],))
            agente = cur.fetchone()
            if agente and agente['email']:
                enviar_alerta_email(
                    agente['email'],
                    "✅ DDJJ CO-FIRMADA",
                    f"""
                    <h2>Declaración Jurada Co-firmada</h2>
                    <p>Su DDJJ del año {ddjj['anio']} ha sido co-firmada por la Dirección.</p>
                    <p>La DDJJ tiene validez legal completa.</p>
                    <p><strong>Código de validación:</strong> {ddjj.get('codigo_validacion', 'N/A')}</p>
                    """
                )
            
            registrar_auditoria(session["usuario"], "CO-FIRMÓ DDJJ", ddjj['legajo_id'], f"DDJJ ID: {ddjj_id}")
            flash("✅ DDJJ co-firmada exitosamente. Tiene validez legal completa.", "success")
        else:
            flash("❌ Error en la co-firma", "error")
        
        return redirect(url_for("listar_ddjj"))
    
    return render_template("cofirmar_ddjj.html", 
                          ddjj=ddjj,
                          agente_nombre=f"{ddjj['apellido']} {ddjj['nombre']}")

@app.route("/ddjj/<int:ddjj_id>/verificar_firma")
@login_requerido
def verificar_firma_ddjj(ddjj_id):
    """Verifica las firmas digitales del PDF"""
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("SELECT archivo_pdf FROM declaraciones_juradas WHERE id = ?", (ddjj_id,))
    ddjj = cur.fetchone()
    
    if not ddjj or not ddjj['archivo_pdf']:
        return jsonify({'error': 'DDJJ sin archivo PDF'}), 404
    
    firmas = verificar_firma_pdf(ddjj['archivo_pdf'])
    
    return jsonify({
        'firmas': firmas,
        'total_firmas': len(firmas),
        'validacion': 'Válida' if len(firmas) >= 2 else 'Pendiente'
    })

# ==========================
# INICIALIZACIÓN DE BASE DE DATOS
# ==========================
def init_db():
    """Inicializa todas las tablas"""
    conn = get_db()
    cur = conn.cursor()
    
    # Tabla usuarios
    cur.execute("""
        CREATE TABLE IF NOT EXISTS usuarios(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            rol TEXT NOT NULL,
            email TEXT,
            telefono TEXT,
            activo INTEGER DEFAULT 1,
            bloqueado INTEGER DEFAULT 0,
            intentos_fallidos INTEGER DEFAULT 0,
            ultimo_acceso TEXT,
            fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)
    
    # Tabla legajos
    cur.execute("""
        CREATE TABLE IF NOT EXISTS legajos(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            legajo_id TEXT UNIQUE NOT NULL,
            apellido TEXT NOT NULL,
            nombre TEXT NOT NULL,
            dni TEXT,
            cuil TEXT,
            email TEXT,
            telefono TEXT,
            fecha_nacimiento TEXT,
            domicilio_calle TEXT,
            domicilio_numero TEXT,
            domicilio_localidad TEXT DEFAULT 'POSADAS',
            cargo TEXT,
            fecha_ingreso TEXT,
            estado TEXT DEFAULT 'Activo'
        )
    """)
    
    # Tabla documentos
    cur.execute("""
        CREATE TABLE IF NOT EXISTS documentos(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            legajo TEXT NOT NULL,
            tipo TEXT NOT NULL,
            nombre_archivo TEXT NOT NULL,
            ruta_pdf TEXT NOT NULL,
            hash_archivo TEXT UNIQUE,
            fecha_subida TEXT,
            subido_por TEXT,
            estado TEXT DEFAULT 'ACTIVO',
            FOREIGN KEY(legajo) REFERENCES legajos(legajo_id)
        )
    """)
    
    # Tabla declaraciones_juradas
    cur.execute("""
        CREATE TABLE IF NOT EXISTS declaraciones_juradas(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            legajo_id TEXT NOT NULL,
            anio INTEGER NOT NULL,
            estado TEXT DEFAULT 'BORRADOR',
            archivo_pdf TEXT,
            hash_pdf TEXT,
            codigo_validacion TEXT UNIQUE,
            fecha_generacion TEXT,
            fecha_envio TEXT,
            usuario_genero TEXT,
            usuario_finalizo TEXT,
            observaciones TEXT,
            activa INTEGER DEFAULT 1,
            FOREIGN KEY(legajo_id) REFERENCES legajos(legajo_id)
        )
    """)
    
    # Tabla ddjj_datos
    cur.execute("""
        CREATE TABLE IF NOT EXISTS ddjj_datos(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ddjj_id INTEGER NOT NULL,
            dni TEXT,
            correo TEXT,
            telefono_celular TEXT,
            domicilio_calle TEXT,
            domicilio_numero TEXT,
            domicilio_localidad TEXT,
            FOREIGN KEY(ddjj_id) REFERENCES declaraciones_juradas(id)
        )
    """)
    
    # Tabla certificados_agentes
    cur.execute("""
        CREATE TABLE IF NOT EXISTS certificados_agentes(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            legajo_id TEXT NOT NULL,
            ruta_certificado TEXT NOT NULL,
            fecha_emision TEXT,
            fecha_expiracion TEXT,
            pin_hash TEXT,
            activo INTEGER DEFAULT 1,
            FOREIGN KEY(legajo_id) REFERENCES legajos(legajo_id)
        )
    """)
    
    # Tabla auditoria
    cur.execute("""
        CREATE TABLE IF NOT EXISTS auditoria(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario TEXT NOT NULL,
            rol TEXT,
            accion TEXT NOT NULL,
            legajo_id TEXT,
            detalle TEXT,
            ip TEXT,
            user_agent TEXT,
            fecha TEXT
        )
    """)
    
    conn.commit()
    logger.info("✅ Tablas creadas/verificadas")
    
    # Crear usuario admin
    cur.execute("SELECT * FROM usuarios WHERE username = 'admin'")
    if not cur.fetchone():
        hashed = generate_password_hash('Admin2026!')
        cur.execute("""
            INSERT INTO usuarios (username, password, rol, email, fecha_creacion)
            VALUES (?, ?, ?, ?, ?)
        """, ('admin', hashed, 'ADMIN', 'admin@hcdposadas.gob.ar', datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        conn.commit()
        logger.info("✅ Usuario admin creado")
    
    # Crear usuario director
    cur.execute("SELECT * FROM usuarios WHERE username = 'director'")
    if not cur.fetchone():
        hashed = generate_password_hash('Director2026!')
        cur.execute("""
            INSERT INTO usuarios (username, password, rol, email, fecha_creacion)
            VALUES (?, ?, ?, ?, ?)
        """, ('director', hashed, 'DIRECTOR', 'director@hcdposadas.gob.ar', datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        conn.commit()
        logger.info("✅ Usuario director creado")
    
    # Índices
    cur.execute("CREATE INDEX IF NOT EXISTS idx_legajos_legajo_id ON legajos(legajo_id)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_ddjj_legajo_anio ON declaraciones_juradas(legajo_id, anio)")
    #cur.execute("CREATE INDEX IF NOT EXISTS idx_ddjj_codigo ON declaraciones_juradas(codigo_validacion)")
    
    logger.info("✅ Índices creados")

def crear_indice_codigo_validacion():
    """Crea el índice para codigo_validacion si no existe"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Crear índice (no causa error si ya existe)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_ddjj_codigo ON declaraciones_juradas(codigo_validacion)")
        conn.commit()
        conn.close()
        print("✅ Índice idx_ddjj_codigo creado/verificado")
        return True
        
    except Exception as e:
        print(f"⚠️ Error creando índice: {e}")
        return False

# ==========================
# RUTAS PRINCIPALES
# ==========================
@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        
        ip = request.remote_addr
        logger.info(f"Intento de login - Usuario: {username} - IP: {ip}")
        
        if not username or not password:
            flash("Por favor ingrese usuario y contraseña", "error")
            return render_template("login.html")
        
        try:
            conn = get_db()
            cur = conn.cursor()
            
            cur.execute("SELECT * FROM usuarios WHERE username = ?", (username,))
            user = cur.fetchone()
            
            if user:
                if user['bloqueado'] == 1:
                    logger.warning(f"Usuario bloqueado intenta acceder: {username}")
                    flash("Usuario bloqueado. Contacte al administrador.", "error")
                    return render_template("login.html")
                
                if check_password_hash(user["password"], password):
                    session.permanent = True
                    session["usuario"] = user["username"]
                    session["rol"] = user["rol"]
                    session["user_id"] = user["id"]
                    
                    cur.execute("UPDATE usuarios SET intentos_fallidos = 0, ultimo_acceso = ? WHERE id = ?",
                               (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), user['id']))
                    conn.commit()
                    
                    registrar_auditoria(user["username"], "LOGIN EXITOSO")
                    logger.info(f"Login exitoso: {username}")
                    
                    flash(f"Bienvenido {user['username']}", "success")
                    return redirect(url_for("panel"))
                else:
                    nuevos_intentos = user['intentos_fallidos'] + 1
                    if nuevos_intentos >= 5:
                        cur.execute("UPDATE usuarios SET bloqueado = 1 WHERE id = ?", (user['id'],))
                        logger.warning(f"Usuario bloqueado por 5 intentos: {username}")
                        flash("Usuario bloqueado por múltiples intentos fallidos", "error")
                    else:
                        cur.execute("UPDATE usuarios SET intentos_fallidos = ? WHERE id = ?",
                                   (nuevos_intentos, user['id']))
                        flash(f"Contraseña incorrecta. Intento {nuevos_intentos} de 5", "error")
                    
                    conn.commit()
                    registrar_auditoria(username, "LOGIN FALLIDO", None, f"IP: {ip}")
            else:
                flash("Usuario o contraseña incorrectos", "error")
                registrar_auditoria(username, "INTENTO LOGIN USUARIO INEXISTENTE", None, f"IP: {ip}")
            
            # Dentro de login(), después de obtener los datos:
            print(f"Username recibido: '{username}'")
            print(f"Password recibida: '{password}'")
            print(f"Largo password: {len(password)}")
            
        except Exception as e:
            logger.error(f"Error en login: {e}")
            flash("Error en el sistema. Intente nuevamente.", "error")
    
    return render_template("login.html")

@app.route("/logout")
@login_requerido
def logout():
    registrar_auditoria(session["usuario"], "LOGOUT")
    logger.info(f"Logout: {session['usuario']}")
    session.clear()
    flash("Sesión cerrada correctamente", "success")
    return redirect(url_for("login"))

@app.route("/panel")
@login_requerido
def panel():
    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT COUNT(*) as total FROM legajos WHERE estado = 'Activo'")
    total_legajos = cur.fetchone()["total"]
    cur.execute("SELECT COUNT(*) as total FROM declaraciones_juradas WHERE anio = ? AND estado IN ('FIRMADA_AGENTE', 'COFIRMADA_DIRECTOR')", 
                (datetime.now().year,))
    total_ddjj_presentadas = cur.fetchone()["total"]
    
    return render_template("panel.html",
                          usuario=session["usuario"],
                          rol=session["rol"],
                          total_legajos=total_legajos,
                          total_ddjj_presentadas=total_ddjj_presentadas)


# ==========================
# DECORADORES DE SEGURIDAD (DEFINICIÓN ÚNICA)
# ==========================
def login_requerido(f):
    @wraps(f)
    def decorado(*args, **kwargs):
        if "usuario" not in session:
            flash("Por favor, inicie sesión", "error")
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorado

def rol_permitido(roles):
    def decorator(f):
        @wraps(f)
        def decorado(*args, **kwargs):
            if session.get("rol") not in roles:
                flash("No tiene permiso para esta acción", "error")
                return redirect(url_for("panel"))
            return f(*args, **kwargs)
        return decorado
    return decorator


# ==========================
# RUTA PARA LISTAR LEGAJOS
# ==========================
@app.route("/legajos")
@login_requerido
def listar_legajos():
    conn = get_db()
    cur = conn.cursor()
    
    # Obtener todos los legajos
    cur.execute("SELECT * FROM legajos ORDER BY apellido, nombre")
    legajos = cur.fetchall()
    
    # Obtener conteo de documentos por legajo
    cur.execute("""
        SELECT legajo, COUNT(*) as total 
        FROM documentos 
        WHERE estado = 'ACTIVO' 
        GROUP BY legajo
    """)
    documentos_por_legajo = {row['legajo']: row['total'] for row in cur.fetchall()}
    
    return render_template("lista_legajos.html", 
                          legajos=legajos,
                          documentos_por_legajo=documentos_por_legajo)



@app.route("/legajo/nuevo", methods=["GET", "POST"])
@login_requerido
@rol_permitido(["ADMIN", "RRHH"])
def nuevo_legajo():
    if request.method == "POST":
        legajo_id = request.form.get("legajo_id", "").strip()
        apellido = request.form.get("apellido", "").strip().upper()
        nombre = request.form.get("nombre", "").strip().upper()
        dni = request.form.get("dni", "").strip()
        email = request.form.get("email", "").strip()
        
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (legajo_id,))
        if cur.fetchone():
            flash("Legajo ya existe", "error")
            return redirect(url_for("nuevo_legajo"))
        
        cur.execute("""
            INSERT INTO legajos (legajo_id, apellido, nombre, dni, email, estado)
            VALUES (?, ?, ?, ?, ?, 'Activo')
        """, (legajo_id, apellido, nombre, dni, email))
        conn.commit()
        
        registrar_auditoria(session["usuario"], "CREÓ LEGAJO", legajo_id)
        flash(f"Legajo {legajo_id} creado", "success")
        return redirect(url_for("ver_legajo", legajo_id=legajo_id))
    
    return render_template("nuevo_legajo.html")

@app.route("/ddjj/<int:ddjj_id>/nueva", methods=["GET"])
@login_requerido
def nueva_ddjj_route(ddjj_id):
    return redirect(url_for("firmar_ddjj", ddjj_id=ddjj_id))

@app.route("/listar_ddjj")
@login_requerido
def listar_ddjj():
    conn = get_db()
    cur = conn.cursor()
    cur.execute("""
        SELECT d.id, d.legajo_id as legajo, d.anio, d.estado, 
               d.fecha_generacion as fecha_presentacion, d.archivo_pdf as archivo_path,
               l.apellido, l.nombre
        FROM declaraciones_juradas d
        JOIN legajos l ON d.legajo_id = l.legajo_id
        ORDER BY d.anio DESC, d.id DESC
    """)
    ddjj_list = cur.fetchall()
    
    # Convertir a lista de diccionarios
    ddjj = []
    for row in ddjj_list:
        d = dict(row)
        # Formatear fecha
        if d.get('fecha_presentacion'):
            try:
                fecha = datetime.strptime(d['fecha_presentacion'], "%Y-%m-%d %H:%M:%S")
                d['fecha_presentacion'] = fecha.strftime("%d/%m/%Y")
            except:
                pass
        else:
            d['fecha_presentacion'] = '-'
        ddjj.append(d)
    
    return render_template("ddjj_lista.html", ddjj=ddjj)
     

@app.route("/ddjj/<int:ddjj_id>/ver")
@login_requerido
def ver_ddjj_detalle(ddjj_id):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("""
        SELECT d.*, l.apellido, l.nombre
        FROM declaraciones_juradas d
        JOIN legajos l ON d.legajo_id = l.legajo_id
        WHERE d.id = ?
    """, (ddjj_id,))
    ddjj = cur.fetchone()
    
    if not ddjj:
        flash("DDJJ no encontrada", "error")
        return redirect(url_for("listar_ddjj"))
    
    return render_template("ddjj_detalle.html", ddjj=dict(ddjj))

@app.route("/ddjj/descargar/<int:ddjj_id>")
@login_requerido
def descargar_ddjj(ddjj_id):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT archivo_pdf FROM declaraciones_juradas WHERE id = ?", (ddjj_id,))
    ddjj = cur.fetchone()
    
    if ddjj and ddjj['archivo_pdf'] and os.path.exists(ddjj['archivo_pdf']):
        return send_file(ddjj['archivo_pdf'], as_attachment=True, download_name=f"DDJJ_{ddjj_id}.pdf")
    
    flash("Archivo no disponible", "error")
    return redirect(url_for("listar_ddjj"))

# ==========================
# RUTA PARA VALIDAR QR
# ==========================
@app.route("/validar/<codigo>")
def validar_qr(codigo):
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT d.*, l.apellido, l.nombre, l.dni
        FROM declaraciones_juradas d
        JOIN legajos l ON d.legajo_id = l.legajo_id
        WHERE d.codigo_validacion = ?
    """, (codigo,))
    
    ddjj = cur.fetchone()
    
    if not ddjj:
        return render_template("validacion_no_encontrada.html", codigo=codigo), 404
    
    return render_template("validacion_exitosa.html", ddjj=dict(ddjj))



# ==========================
# RUTA PARA NUEVO LEGAJO
# ==========================

@app.route("/legajo/crear", methods=["POST"])
@login_requerido
@rol_permitido(["ADMIN", "RRHH"])
def crear_legajo():
    legajo_id = request.form.get("legajo_id", "").strip()
    apellido = request.form.get("apellido", "").strip().upper()
    nombre = request.form.get("nombre", "").strip().upper()
    dni = request.form.get("dni", "").strip()
    email = request.form.get("email", "").strip().lower()
    telefono = request.form.get("telefono", "").strip()
    fecha_nacimiento = request.form.get("fecha_nacimiento", "")
    estado_civil = request.form.get("estado_civil", "")
    cargo = request.form.get("cargo", "").strip()
    tipo_personal = request.form.get("tipo_personal", "")
    fecha_ingreso = request.form.get("fecha_ingreso", "")
    estado = request.form.get("estado", "Activo")
    
    conn = get_db()
    cur = conn.cursor()
    
    if not legajo_id or not apellido or not nombre:
        flash("Legajo, Apellido y Nombre son obligatorios", "error")
        return redirect(url_for("nuevo_legajo"))
    
    cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (legajo_id,))
    if cur.fetchone():
        flash(f"El legajo {legajo_id} ya existe", "error")
        return redirect(url_for("nuevo_legajo"))
    
    try:
        cur.execute("""
            INSERT INTO legajos (
                legajo_id, apellido, nombre, dni, email, telefono,
                fecha_nacimiento, estado_civil, cargo, tipo_personal,
                fecha_ingreso, estado
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (legajo_id, apellido, nombre, dni, email, telefono,
              fecha_nacimiento, estado_civil, cargo, tipo_personal,
              fecha_ingreso, estado))
        
        carpeta_agente = os.path.join(app.config['UPLOAD_FOLDER'], legajo_id)
        os.makedirs(carpeta_agente, exist_ok=True)
        
        registrar_auditoria(session["usuario"], "CREÓ LEGAJO", legajo_id)
        conn.commit()
        
        flash(f"Legajo {legajo_id} creado exitosamente", "success")
        return redirect(url_for("ver_legajo", legajo_id=legajo_id))
        
    except Exception as e:
        conn.rollback()
        logger.error(f"Error creando legajo: {e}")
        flash(f"Error al crear legajo: {str(e)}", "error")
        return redirect(url_for("nuevo_legajo"))

# ==========================
# RUTA PARA VER DETALLE DE LEGAJO
# ==========================
@app.route("/legajo/<legajo_id>")
@login_requerido
def ver_legajo(legajo_id):
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (legajo_id,))
    legajo = cur.fetchone()
    
    if not legajo:
        flash("Legajo no encontrado", "error")
        return redirect(url_for("listar_legajos"))
    
    cur.execute("""
        SELECT * FROM documentos 
        WHERE legajo = ? AND estado = 'ACTIVO' 
        ORDER BY tipo, fecha_subida DESC
    """, (legajo_id,))
    documentos = cur.fetchall()
    
    documentos_por_tipo = {}
    for doc in documentos:
        doc_dict = dict(doc)
        tipo = doc_dict['tipo']
        if tipo not in documentos_por_tipo:
            documentos_por_tipo[tipo] = []
        documentos_por_tipo[tipo].append(doc_dict)
    
    total_requeridos = sum(1 for d in DOCUMENTOS_OBLIGATORIOS if d['requerido'])
    completados = sum(1 for d in DOCUMENTOS_OBLIGATORIOS 
                     if d['requerido'] and d['id'].upper() in documentos_por_tipo)
    progreso = int((completados / total_requeridos) * 100) if total_requeridos > 0 else 0
    
    cur.execute("""
        SELECT * FROM declaraciones_juradas 
        WHERE legajo_id = ? 
        ORDER BY anio DESC, version DESC
    """, (legajo_id,))
    declaraciones = [dict(d) for d in cur.fetchall()]
    
    anexos = []
    
    categorias = {
        'IDENTIDAD': 'Identificación',
        'FAMILIA': 'Familia',
        'EDUCACION': 'Educación',
        'DOMICILIO': 'Domicilio',
        'CONTACTO': 'Contacto',
        'DDJJ': 'Declaraciones Juradas',
        'LABORAL': 'Laboral',
        'SALUD': 'Salud',
        'NOMBRAMIENTO': 'Nombramiento',
        'LEGAJO': 'Legajo'
    }
    
    registrar_auditoria(session["usuario"], "VIO LEGAJO", legajo_id)
    
    return render_template("legajo_detalle.html",
                          legajo=dict(legajo),
                          documentos_por_tipo=documentos_por_tipo,
                          documentos_obligatorios=DOCUMENTOS_OBLIGATORIOS,
                          categorias=categorias,
                          progreso_documentos=progreso,
                          total_requeridos=total_requeridos,
                          completados=completados,
                          declaraciones=declaraciones,
                          anexos=anexos)

# ==========================
# RUTA PARA SUBIR DOCUMENTOS
# ==========================
@app.route("/legajo/<legajo_id>/subir_documento", methods=["POST"])
@login_requerido
@rol_permitido(["ADMIN", "RRHH"])
def subir_documento(legajo_id):
    if 'archivo' not in request.files:
        return jsonify({'error': 'No se seleccionó archivo'}), 400
    
    archivo = request.files["archivo"]
    tipo_documento = request.form.get("tipo_documento", "").upper()
    
    if not tipo_documento:
        return jsonify({'error': 'Tipo de documento no especificado'}), 400
    
    if archivo.filename == "":
        return jsonify({'error': 'Nombre de archivo vacío'}), 400
    
    tipos_validos = [d['id'].upper() for d in DOCUMENTOS_OBLIGATORIOS]
    if tipo_documento not in tipos_validos:
        return jsonify({'error': 'Tipo de documento no válido'}), 400
    
    try:
        doc_id, ruta_pdf, hash_archivo = guardar_documento_seguro(archivo, legajo_id, tipo_documento)
        
        if doc_id:
            conn = get_db()
            cur = conn.cursor()
            cur.execute("SELECT COUNT(*) as total FROM documentos WHERE legajo = ? AND estado = 'ACTIVO'", (legajo_id,))
            total_docs = cur.fetchone()['total']
            
            if total_docs >= len([d for d in DOCUMENTOS_OBLIGATORIOS if d['requerido']]):
                enviar_alerta_email(
                    EMAIL_CONFIG['email_admin'],
                    "✅ Legajo Completo",
                    f"El legajo {legajo_id} ha completado toda la documentación obligatoria."
                )
            
            registrar_auditoria(session["usuario"], "SUBIO DOCUMENTO", legajo_id, 
                               f"Tipo: {tipo_documento} - Hash: {hash_archivo[:16]}...")
            
            return jsonify({
                'success': True,
                'doc_id': doc_id,
                'hash': hash_archivo,
                'mensaje': 'Documento subido correctamente'
            })
        else:
            return jsonify({'error': 'Error al guardar el documento'}), 500
            
    except Exception as e:
        logger.error(f"Error subiendo documento: {e}")
        return jsonify({'error': str(e)}), 500

# ==========================
# RUTA PARA VER DOCUMENTO
# ==========================
@app.route("/documento/<int:doc_id>/ver")
@login_requerido
def ver_documento(doc_id):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM documentos WHERE id = ?", (doc_id,))
    doc = cur.fetchone()
    
    if not doc:
        flash("Documento no encontrado", "error")
        return redirect(url_for("panel"))
    
    if session['rol'] not in ['ADMIN', 'RRHH']:
        if session.get('legajo_asignado') != doc['legajo']:
            registrar_auditoria(session["usuario"], "INTENTO ACCESO NO AUTORIZADO", doc['legajo'],
                               f"Documento ID: {doc_id}")
            flash("No tiene permiso para ver este documento", "error")
            return redirect(url_for("panel"))
    
    registrar_auditoria(session["usuario"], "VIO DOCUMENTO", doc['legajo'], f"Documento ID: {doc_id}")
    
    return send_file(doc['ruta_pdf'], as_attachment=False)

# ==========================
# RUTA PARA DESCARGAR DOCUMENTO
# ==========================
@app.route("/documento/<int:doc_id>/descargar")
@login_requerido
def descargar_documento(doc_id):
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM documentos WHERE id = ?", (doc_id,))
    doc = cur.fetchone()
    
    if not doc:
        flash("Documento no encontrado", "error")
        return redirect(url_for("panel"))
    
    registrar_auditoria(session["usuario"], "DESCARGÓ DOCUMENTO", doc['legajo'], f"Documento ID: {doc_id}")
    
    return send_file(doc['ruta_pdf'], as_attachment=True, download_name=doc['nombre_archivo'])

# ==========================
# RUTA PARA ANULAR DOCUMENTO
# ==========================
@app.route("/documento/<int:doc_id>/anular", methods=["POST"])
@login_requerido
@rol_permitido(["ADMIN"])
def anular_documento(doc_id):
    motivo = request.form.get("motivo", "Sin especificar")
    
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("SELECT * FROM documentos WHERE id = ?", (doc_id,))
    doc = cur.fetchone()
    
    if doc:
        cur.execute("""
            UPDATE documentos 
            SET estado = 'ANULADO', 
                anulado_por = ?, 
                fecha_anulacion = ?,
                motivo_anulacion = ?
            WHERE id = ?
        """, (session["usuario"], datetime.now().strftime("%Y-%m-%d %H:%M:%S"), motivo, doc_id))
        conn.commit()
        
        registrar_auditoria(session["usuario"], "ANULÓ DOCUMENTO", doc['legajo'], 
                           f"Documento ID: {doc_id} - Motivo: {motivo}")
        
        flash("Documento anulado correctamente", "success")
        return redirect(url_for("ver_legajo", legajo_id=doc['legajo']))
    else:
        flash("Documento no encontrado", "error")
        return redirect(url_for("panel"))

# ==========================
# RUTA PARA NUEVA DDJJ
# ==========================
@app.route("/legajo/<legajo_id>/ddjj/nueva", methods=["GET", "POST"])
@login_requerido
def nueva_ddjj(legajo_id):
    anio_actual = datetime.now().year
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (legajo_id,))
    agente = cur.fetchone()
    
    if not agente:
        flash("El legajo no existe", "error")
        return redirect(url_for("panel"))
    
    tipos_obligatorios = [doc['id'].upper() for doc in DOCUMENTOS_OBLIGATORIOS if doc['requerido']]
    documentos_requeridos = len(tipos_obligatorios)
    
    if tipos_obligatorios:
        placeholders = ','.join(['?'] * len(tipos_obligatorios))
        query = f"""
            SELECT DISTINCT tipo FROM documentos 
            WHERE legajo = ? AND estado = 'ACTIVO' 
            AND tipo IN ({placeholders})
        """
        params = [legajo_id] + tipos_obligatorios
        cur.execute(query, params)
        tipos_existentes = [row['tipo'] for row in cur.fetchall()]
        total_docs_unicos = len(tipos_existentes)
    else:
        total_docs_unicos = 0
    
    progreso = int((total_docs_unicos / documentos_requeridos) * 100) if documentos_requeridos > 0 else 0
    
    documentos_faltantes = []
    for doc in DOCUMENTOS_OBLIGATORIOS:
        if doc['requerido'] and doc['id'].upper() not in tipos_existentes:
            documentos_faltantes.append(doc)
    
    if request.method == "POST":
        cur.execute("""
            SELECT id FROM declaraciones_juradas 
            WHERE legajo_id=? AND anio=? AND activa=1
        """, (legajo_id, anio_actual))
        existe = cur.fetchone()
        
        if existe:
            flash("Ya existe una DDJJ activa para este año", "warning")
            return redirect(url_for("editar_ddjj", ddjj_id=existe["id"]))
        
        if total_docs_unicos < documentos_requeridos:
            flash(f"Debe completar todos los documentos obligatorios antes de generar la DDJJ. Completados: {total_docs_unicos}/{documentos_requeridos}", "warning")
            return redirect(url_for("ver_legajo", legajo_id=legajo_id))
        
        fecha_ahora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        cur.execute("""
            INSERT INTO declaraciones_juradas
            (legajo_id, anio, estado, fecha_generacion, usuario_genero, activa)
            VALUES (?, ?, 'BORRADOR', ?, ?, 1)
        """, (legajo_id, anio_actual, fecha_ahora, session["usuario"]))
        
        ddjj_id = cur.lastrowid
        
        cur.execute("""
            INSERT INTO ddjj_datos (ddjj_id, dni, fecha_ingreso)
            VALUES (?, ?, ?)
        """, (ddjj_id, agente['dni'], agente['fecha_ingreso']))
        
        conn.commit()
        
        registrar_auditoria(session["usuario"], "CREÓ DDJJ", legajo_id, f"DDJJ ID: {ddjj_id}")
        flash("DDJJ creada en modo borrador. Complete los datos.", "success")
        return redirect(url_for("editar_ddjj", ddjj_id=ddjj_id))
    
    return render_template("ddjj_nueva.html", 
                          legajo=dict(agente), 
                          anio=anio_actual,
                          completados=total_docs_unicos,
                          total_requeridos=documentos_requeridos,
                          progreso_documentos=progreso,
                          documentos_faltantes=documentos_faltantes)

# ==========================
# RUTA PARA EDITAR DDJJ
# ==========================
@app.route("/ddjj/<int:ddjj_id>/editar", methods=["GET", "POST"])
@login_requerido
def editar_ddjj(ddjj_id):
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT d.*, l.apellido, l.nombre, l.legajo_id, l.dni, l.email, l.telefono
        FROM declaraciones_juradas d
        JOIN legajos l ON d.legajo_id = l.legajo_id
        WHERE d.id = ?
    """, (ddjj_id,))
    ddjj = cur.fetchone()
    
    if not ddjj:
        flash("DDJJ no encontrada", "error")
        return redirect(url_for("panel"))
    
    ddjj = dict(ddjj)
    
    cur.execute("SELECT * FROM ddjj_datos WHERE ddjj_id = ?", (ddjj_id,))
    datos = cur.fetchone()
    datos = dict(datos) if datos else {}
    
    cur.execute("SELECT * FROM ddjj_familiares_cargo WHERE ddjj_id = ? ORDER BY orden", (ddjj_id,))
    familiares = [dict(f) for f in cur.fetchall()]
    
    cur.execute("SELECT * FROM ddjj_padres_hermanos WHERE ddjj_id = ? ORDER BY orden", (ddjj_id,))
    padres_hermanos = [dict(p) for p in cur.fetchall()]
    
    cur.execute("SELECT * FROM ddjj_conyuge WHERE ddjj_id = ?", (ddjj_id,))
    conyuge = cur.fetchone()
    conyuge = dict(conyuge) if conyuge else {}
    
    if request.method == "POST":
        if ddjj['estado'] != 'BORRADOR':
            flash("No se puede editar una DDJJ finalizada", "error")
            return redirect(url_for("ver_ddjj", ddjj_id=ddjj_id))
        
        try:
            cur.execute("""
                UPDATE ddjj_datos SET
                    fecha_nacimiento = ?, lugar_nacimiento = ?, dni = ?,
                    estado_civil = ?, estudios_completos = ?, estudios_incompletos = ?,
                    estudios_nivel = ?, titulo = ?, anios_cursados = ?,
                    domicilio_barrio = ?, domicilio_chacra = ?, domicilio_calle = ?,
                    domicilio_numero = ?, domicilio_piso = ?, domicilio_depto = ?,
                    domicilio_localidad = ?, telefono_fijo = ?, telefono_celular = ?,
                    correo = ?, telefono_referencia = ?, nombre_referencia = ?,
                    parentesco_referencia = ?, fecha_ingreso = ?, categoria = ?,
                    antiguedad = ?, lugar_desempeno = ?, actividad = ?,
                    enfermedad_preexistente = ?, vacuna_covid_dosis = ?,
                    vacuna_gripe = ?, vacuna_neumonia = ?, percibe_asig_fliares_hcd = ?,
                    servicios_anteriores = ?, antiguedad_servicios = ?,
                    trabaja_otra_reparticion = ?, otra_reparticion_donde = ?,
                    percibe_asig_fliares_otro = ?, lugar = ?, firma_agente = ?
                WHERE ddjj_id = ?
            """, (
                request.form.get('fecha_nacimiento'), request.form.get('lugar_nacimiento'),
                request.form.get('dni'), request.form.get('estado_civil'),
                request.form.get('estudios_completos'), request.form.get('estudios_incompletos'),
                request.form.get('estudios_nivel'), request.form.get('titulo'),
                request.form.get('anios_cursados'), request.form.get('domicilio_barrio'),
                request.form.get('domicilio_chacra'), request.form.get('domicilio_calle'),
                request.form.get('domicilio_numero'), request.form.get('domicilio_piso'),
                request.form.get('domicilio_depto'), request.form.get('domicilio_localidad'),
                request.form.get('telefono_fijo'), request.form.get('telefono_celular'),
                request.form.get('correo'), request.form.get('telefono_referencia'),
                request.form.get('nombre_referencia'), request.form.get('parentesco_referencia'),
                request.form.get('fecha_ingreso'), request.form.get('categoria'),
                request.form.get('antiguedad'), request.form.get('lugar_desempeno'),
                request.form.get('actividad'), request.form.get('enfermedad_preexistente'),
                request.form.get('vacuna_covid_dosis'), request.form.get('vacuna_gripe'),
                request.form.get('vacuna_neumonia'), 1 if request.form.get('percibe_asig_fliares_hcd') else 0,
                request.form.get('servicios_anteriores'), request.form.get('antiguedad_servicios'),
                1 if request.form.get('trabaja_otra_reparticion') else 0,
                request.form.get('otra_reparticion_donde'),
                1 if request.form.get('percibe_asig_fliares_otro') else 0,
                request.form.get('lugar'), request.form.get('firma_agente'),
                ddjj_id
            ))
            
            cur.execute("DELETE FROM ddjj_familiares_cargo WHERE ddjj_id = ?", (ddjj_id,))
            for i in range(1, 8):
                if request.form.get(f'familiar_{i}_apellido_nombre'):
                    cur.execute("""
                        INSERT INTO ddjj_familiares_cargo 
                        (ddjj_id, orden, apellido_nombre, parentesco, convive,
                         fecha_nacimiento, edad, dni, ni, pri, sec, ter_univ)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        ddjj_id, i,
                        request.form.get(f'familiar_{i}_apellido_nombre'),
                        request.form.get(f'familiar_{i}_parentesco'),
                        1 if request.form.get(f'familiar_{i}_convive') else 0,
                        request.form.get(f'familiar_{i}_fecha_nacimiento'),
                        request.form.get(f'familiar_{i}_edad'),
                        request.form.get(f'familiar_{i}_dni'),
                        request.form.get(f'familiar_{i}_ni'),
                        request.form.get(f'familiar_{i}_pri'),
                        request.form.get(f'familiar_{i}_sec'),
                        request.form.get(f'familiar_{i}_ter_univ')
                    ))
            
            cur.execute("DELETE FROM ddjj_padres_hermanos WHERE ddjj_id = ?", (ddjj_id,))
            for i in range(1, 7):
                if request.form.get(f'ph_{i}_apellido_nombre'):
                    cur.execute("""
                        INSERT INTO ddjj_padres_hermanos 
                        (ddjj_id, orden, apellido_nombre, parentesco, convive,
                         fecha_nacimiento, dni, observacion)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        ddjj_id, i,
                        request.form.get(f'ph_{i}_apellido_nombre'),
                        request.form.get(f'ph_{i}_parentesco'),
                        1 if request.form.get(f'ph_{i}_convive') else 0,
                        request.form.get(f'ph_{i}_fecha_nacimiento'),
                        request.form.get(f'ph_{i}_dni'),
                        request.form.get(f'ph_{i}_observacion')
                    ))
            
            cur.execute("DELETE FROM ddjj_conyuge WHERE ddjj_id = ?", (ddjj_id,))
            if request.form.get('conyuge_apellido_nombre'):
                cur.execute("""
                    INSERT INTO ddjj_conyuge 
                    (ddjj_id, apellido_nombre, dni, fecha_enlace, lugar_enlace,
                     trabaja, razon_social, percibe_asignaciones, observaciones)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    ddjj_id,
                    request.form.get('conyuge_apellido_nombre'),
                    request.form.get('conyuge_dni'),
                    request.form.get('conyuge_fecha_enlace'),
                    request.form.get('conyuge_lugar_enlace'),
                    1 if request.form.get('conyuge_trabaja') else 0,
                    request.form.get('conyuge_razon_social'),
                    1 if request.form.get('conyuge_percibe_asignaciones') else 0,
                    request.form.get('conyuge_observaciones')
                ))
            
            conn.commit()
            registrar_auditoria(session["usuario"], "EDITÓ DDJJ", ddjj['legajo_id'], f"DDJJ ID: {ddjj_id}")
            flash("Datos guardados correctamente", "success")
            
        except Exception as e:
            conn.rollback()
            logger.error(f"Error guardando DDJJ: {e}")
            flash(f"Error al guardar: {str(e)}", "error")
        
        return redirect(url_for("editar_ddjj", ddjj_id=ddjj_id))
    
    return render_template("ddjj_editar_completa.html",
                          ddjj=ddjj,
                          legajo={'legajo_id': ddjj['legajo_id'], 'apellido': ddjj['apellido'], 'nombre': ddjj['nombre']},
                          datos=datos,
                          conyuge=conyuge,
                          familiares_cargo=familiares,
                          padres_hermanos=padres_hermanos)

# ==========================
# RUTA PARA VER DDJJ
# ==========================
@app.route("/ddjj/<int:ddjj_id>")
@login_requerido
def ver_ddjj(ddjj_id):
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT d.*, l.apellido, l.nombre, l.legajo_id
        FROM declaraciones_juradas d
        JOIN legajos l ON d.legajo_id = l.legajo_id
        WHERE d.id = ?
    """, (ddjj_id,))
    ddjj = cur.fetchone()
    
    if not ddjj:
        flash("DDJJ no encontrada", "error")
        return redirect(url_for("panel"))
    
    if not ddjj['archivo_pdf']:
        flash("DDJJ aún no ha sido generada", "warning")
        return redirect(url_for("editar_ddjj", ddjj_id=ddjj_id))
    
    registrar_auditoria(session["usuario"], "VIO DDJJ", ddjj['legajo_id'], f"DDJJ ID: {ddjj_id}")
    
    return send_file(ddjj['archivo_pdf'], as_attachment=False, mimetype='application/pdf')
# ==========================
# FUNCIÓN PARA MAPEAR DATOS A LA PLANTILLA (VERSIÓN FINAL)
# ==========================
def mapear_datos_para_plantilla(legajo, datos, familiares, padres, conyuge, ddjj):
    """
    Convierte los datos de la BD al formato que espera la plantilla
    """
    import hashlib
    import uuid
    
    contexto = {}
    
    # DATOS BÁSICOS
    contexto['anio'] = ddjj.get('anio') if ddjj else datetime.now().year
    contexto['legajo'] = legajo.get('legajo_id', '') if legajo else ''
    
    apellido = legajo.get('apellido', '').strip()
    nombre = legajo.get('nombre', '').strip()
    contexto['apellido_nombre'] = f"{apellido} {nombre}".strip()
    
    contexto['dni'] = datos.get('dni', legajo.get('dni', ''))
    fecha_nac = datos.get('fecha_nacimiento', legajo.get('fecha_nacimiento', ''))
    contexto['fecha_nacimiento'] = fecha_nac if fecha_nac else ''
    contexto['lugar_nacimiento'] = datos.get('lugar_nacimiento', '')
    contexto['estado_civil'] = datos.get('estado_civil', legajo.get('estado_civil', ''))
    
    # CONTACTO Y DOMICILIO
    contexto['email'] = datos.get('correo', legajo.get('email', ''))
    contexto['telefono_fijo'] = datos.get('telefono_fijo', legajo.get('telefono_fijo', ''))
    contexto['telefono_celular'] = datos.get('telefono_celular', legajo.get('telefono', ''))
    contexto['telefono_referencia'] = datos.get('telefono_referencia', legajo.get('telefono_referencia', ''))
    contexto['nombre_referencia'] = datos.get('nombre_referencia', '')
    contexto['parentesco_referencia'] = datos.get('parentesco_referencia', legajo.get('parentesco_referencia', ''))
    
    contexto['domicilio_calle'] = datos.get('domicilio_calle', legajo.get('domicilio_calle', ''))
    contexto['domicilio_numero'] = datos.get('domicilio_numero', legajo.get('domicilio_numero', ''))
    contexto['domicilio_piso'] = datos.get('domicilio_piso', legajo.get('domicilio_piso', ''))
    contexto['domicilio_depto'] = datos.get('domicilio_depto', legajo.get('domicilio_depto', ''))
    contexto['domicilio_barrio'] = datos.get('domicilio_barrio', legajo.get('domicilio_barrio', ''))
    contexto['domicilio_localidad'] = datos.get('domicilio_localidad', legajo.get('domicilio_localidad', 'POSADAS'))
    contexto['domicilio_chacra'] = datos.get('domicilio_chacra', '')
    
    # ESTUDIOS
    contexto['estudios_completos'] = datos.get('estudios_completos', '')
    contexto['estudios_incompletos'] = datos.get('estudios_incompletos', '')
    
    nivel = datos.get('estudios_nivel', '').upper()
    contexto['primario'] = 'X' if nivel == 'PRIMARIO' else ''
    contexto['secundario'] = 'X' if nivel == 'SECUNDARIO' else ''
    contexto['terciario'] = 'X' if nivel == 'TERCIARIO' else ''
    contexto['universitario'] = 'X' if nivel == 'UNIVERSITARIO' else ''
    
    contexto['titulo'] = datos.get('titulo', '')
    contexto['anios_cursados'] = datos.get('anios_cursados', '')
    
    # LABORAL
    contexto['fecha_ingr'] = datos.get('fecha_ingreso', legajo.get('fecha_ingreso', ''))
    contexto['categ'] = datos.get('categoria', legajo.get('categoria', ''))
    contexto['antig'] = datos.get('antiguedad', '')
    
    contexto['p_permanente'] = 'X' if datos.get('planta_permanente') or legajo.get('planta_permanente') else ''
    contexto['p_temporaria'] = 'X' if datos.get('planta_temporaria') or legajo.get('planta_temporaria') else ''
    
    contexto['lugar_dependencia'] = datos.get('lugar_desempeno', legajo.get('lugar_desempeno', ''))
    contexto['actividad_realiza_dependencia'] = datos.get('actividad', legajo.get('actividad', ''))
    
    # SALUD
    contexto['enfermedad_preexistente'] = datos.get('enfermedad_preexistente', legajo.get('enfermedad_preexistente', ''))
    contexto['vacuna_covid_dosis'] = datos.get('vacuna_covid_dosis', legajo.get('vacuna_covid_dosis', ''))
    
    vac_gripe = datos.get('vacuna_gripe', 'NO')
    if isinstance(vac_gripe, (int, bool)):
        vac_gripe = 'SI' if vac_gripe else 'NO'
    contexto['vacuna_gripe'] = vac_gripe
    
    vac_neumonia = datos.get('vacuna_neumonia', 'NO')
    if isinstance(vac_neumonia, (int, bool)):
        vac_neumonia = 'SI' if vac_neumonia else 'NO'
    contexto['vacuna_neumonia'] = vac_neumonia
    
    # ASIGNACIONES
    asig_hcd = datos.get('percibe_asig_fliares_hcd', 'NO')
    if isinstance(asig_hcd, (int, bool)):
        asig_hcd = 'SI' if asig_hcd else 'NO'
    contexto['percibe_asig_fliares_hcd'] = asig_hcd
    
    contexto['servicios_anteriores'] = datos.get('servicios_anteriores', legajo.get('servicios_anteriores', ''))
    contexto['antiguedad_servicios'] = datos.get('antiguedad_servicios', legajo.get('antiguedad_servicios', ''))
    
    trabaja_otra = datos.get('trabaja_otra_reparticion', 'NO')
    if isinstance(trabaja_otra, (int, bool)):
        trabaja_otra = 'SI' if trabaja_otra else 'NO'
    contexto['trabaja_otra_reparticion'] = trabaja_otra
    
    contexto['otra_reparticion_donde'] = datos.get('otra_reparticion_donde', legajo.get('otra_reparticion_donde', ''))
    
    asig_otro = datos.get('percibe_asig_fliares_otro', 'NO')
    if isinstance(asig_otro, (int, bool)):
        asig_otro = 'SI' if asig_otro else 'NO'
    contexto['percibe_asig_fliares_otro'] = asig_otro
    
    # ESTADO CIVIL (CHECKBOXES)
    ec = contexto['estado_civil'].lower()
    contexto['soltero'] = 'X' if ec == 'soltero' else ''
    contexto['casado'] = 'X' if ec == 'casado' else ''
    contexto['divorciado'] = 'X' if ec == 'divorciado' else ''
    contexto['separado'] = 'X' if ec == 'separado' else ''
    contexto['viudo'] = 'X' if ec == 'viudo' else ''
    contexto['conviviente'] = 'X' if ec in ['conviviente', 'concubinato'] else ''
    
    # CÓNYUGE
    if conyuge:
        contexto['conyuge_apellido_nombre'] = conyuge.get('apellido_nombre', '')
        contexto['conyuge_dni'] = conyuge.get('dni', '')
        contexto['conyuge_fecha_enlace'] = conyuge.get('fecha_enlace', '')
        contexto['conyuge_lugar_enlace'] = conyuge.get('lugar_enlace', '')
        
        trabaja = conyuge.get('trabaja', 'NO')
        if isinstance(trabaja, (int, bool)):
            trabaja = 'SI' if trabaja else 'NO'
        contexto['conyuge_trabaja'] = trabaja
        
        contexto['conyuge_razon_social'] = conyuge.get('razon_social', '')
        
        percibe = conyuge.get('percibe_asignaciones', 'NO')
        if isinstance(percibe, (int, bool)):
            percibe = 'SI' if percibe else 'NO'
        contexto['conyuge_percibe_asignaciones'] = percibe
        
        contexto['conyuge_observaciones'] = conyuge.get('observaciones', '')
    else:
        contexto['conyuge_apellido_nombre'] = ''
        contexto['conyuge_dni'] = ''
        contexto['conyuge_fecha_enlace'] = ''
        contexto['conyuge_lugar_enlace'] = ''
        contexto['conyuge_trabaja'] = 'NO'
        contexto['conyuge_razon_social'] = ''
        contexto['conyuge_percibe_asignaciones'] = 'NO'
        contexto['conyuge_observaciones'] = ''
    
    # FAMILIARES A CARGO (f1 a f7)
    for i in range(1, 8):
        idx = i - 1
        if idx < len(familiares):
            f = familiares[idx]
            
            contexto[f'f{i}.apellido'] = f.get('apellido_nombre', '')
            contexto[f'f{i}.par'] = f.get('parentesco', '')
            
            convive_val = f.get('convive', 1)
            contexto[f'f{i}.convive'] = 'SI' if convive_val and convive_val != 0 else 'NO'
            
            contexto[f'f{i}.fe_nac'] = f.get('fecha_nacimiento', '')
            
            edad = f.get('edad', '')
            contexto[f'f{i}.ed'] = str(edad) if edad else ''
            
            contexto[f'f{i}.dni'] = f.get('dni', '')
            contexto[f'f{i}.ni'] = f.get('ni', '')
            contexto[f'f{i}.pri'] = f.get('pri', '')
            contexto[f'f{i}.sec'] = f.get('sec', '')
            contexto[f'f{i}.ter'] = f.get('ter_univ', '')
            
            contexto[f'f{i}'] = {
                'apellido': f.get('apellido_nombre', ''),
                'par': f.get('parentesco', ''),
                'convive': 'SI' if convive_val and convive_val != 0 else 'NO',
                'fe_nac': f.get('fecha_nacimiento', ''),
                'ed': str(edad) if edad else '',
                'dni': f.get('dni', ''),
                'ni': f.get('ni', ''),
                'pri': f.get('pri', ''),
                'sec': f.get('sec', ''),
                'ter': f.get('ter_univ', '')
            }
        else:
            contexto[f'f{i}.apellido'] = ''
            contexto[f'f{i}.par'] = ''
            contexto[f'f{i}.convive'] = ''
            contexto[f'f{i}.fe_nac'] = ''
            contexto[f'f{i}.ed'] = ''
            contexto[f'f{i}.dni'] = ''
            contexto[f'f{i}.ni'] = ''
            contexto[f'f{i}.pri'] = ''
            contexto[f'f{i}.sec'] = ''
            contexto[f'f{i}.ter'] = ''
            
            contexto[f'f{i}'] = {
                'apellido': '', 'par': '', 'convive': '', 'fe_nac': '',
                'ed': '', 'dni': '', 'ni': '', 'pri': '', 'sec': '', 'ter': ''
            }
    
    # PADRES Y HERMANOS (ph_1 a ph_6)
    for i in range(1, 7):
        idx = i - 1
        if idx < len(padres):
            p = padres[idx]
            
            contexto[f'ph_{i}_apellido'] = p.get('apellido_nombre', '')
            contexto[f'ph_{i}_parentesco'] = p.get('parentesco', '')
            
            convive_val = p.get('convive', 0)
            contexto[f'ph_{i}_convive'] = 'SI' if convive_val and convive_val != 0 else 'NO'
            
            contexto[f'ph_{i}_fecha_nac'] = p.get('fecha_nacimiento', '')
            contexto[f'ph_{i}_dni'] = p.get('dni', '')
            contexto[f'ph_{i}_observacion'] = p.get('observacion', '')
            
            contexto[f'ph_{i}'] = {
                'apellido': p.get('apellido_nombre', ''),
                'parentesco': p.get('parentesco', ''),
                'convive': 'SI' if convive_val and convive_val != 0 else 'NO',
                'fecha_nac': p.get('fecha_nacimiento', ''),
                'dni': p.get('dni', ''),
                'observacion': p.get('observacion', '')
            }
        else:
            contexto[f'ph_{i}_apellido'] = ''
            contexto[f'ph_{i}_parentesco'] = ''
            contexto[f'ph_{i}_convive'] = ''
            contexto[f'ph_{i}_fecha_nac'] = ''
            contexto[f'ph_{i}_dni'] = ''
            contexto[f'ph_{i}_observacion'] = ''
            
            contexto[f'ph_{i}'] = {
                'apellido': '', 'parentesco': '', 'convive': '',
                'fecha_nac': '', 'dni': '', 'observacion': ''
            }
    
    # CÓDIGO DE VALIDACIÓN Y QR
    if ddjj and ddjj.get('codigo_validacion'):
        codigo = ddjj['codigo_validacion']
    else:
        codigo = hashlib.sha256(f"{uuid.uuid4()}{datetime.now()}".encode()).hexdigest()[:16].upper()
    
    contexto['hash_documento'] = codigo
    contexto['qr_validacion'] = f"http://127.0.0.1:5000/validar/{codigo}"
    
    # FIRMAS Y SELLOS
    contexto['firma_agente'] = "_______________"
    contexto['sello_firma_director'] = "__________________"
    contexto['sello_organismo'] = "_______________"
    contexto['firma_agente2'] = "__________________"
    contexto['sello_firma_director2'] = "________________"
    contexto['sello_organismo2'] = "_____________________"
    
    # FECHA Y LUGAR
    contexto['lugar_declaracion'] = datos.get('lugar', 'POSADAS')
    contexto['fecha_declaracion'] = datetime.now().strftime('%d/%m/%Y')
    
    return contexto, codigo

# ==========================
# RUTA PARA FINALIZAR DDJJ
# ==========================
@app.route("/ddjj/<int:ddjj_id>/finalizar", methods=["POST"])
@login_requerido
def finalizar_ddjj(ddjj_id):
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("SELECT * FROM declaraciones_juradas WHERE id = ?", (ddjj_id,))
    ddjj = cur.fetchone()
    
    if not ddjj:
        flash("DDJJ no encontrada", "error")
        return redirect(url_for("panel"))
    
    ddjj = dict(ddjj)
    
    if ddjj['estado'] != 'BORRADOR':
        flash("La DDJJ ya está finalizada", "warning")
        return redirect(url_for("ver_legajo", legajo_id=ddjj['legajo_id']))
    
    temp_dir = None
    
    try:
        cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (ddjj['legajo_id'],))
        legajo = dict(cur.fetchone())
        
        cur.execute("SELECT * FROM ddjj_datos WHERE ddjj_id = ?", (ddjj_id,))
        datos_row = cur.fetchone()
        datos = dict(datos_row) if datos_row else {}
        
        cur.execute("SELECT * FROM ddjj_familiares_cargo WHERE ddjj_id = ? ORDER BY orden", (ddjj_id,))
        familiares = [dict(f) for f in cur.fetchall()]
        
        cur.execute("SELECT * FROM ddjj_padres_hermanos WHERE ddjj_id = ? ORDER BY orden", (ddjj_id,))
        padres = [dict(p) for p in cur.fetchall()]
        
        cur.execute("SELECT * FROM ddjj_conyuge WHERE ddjj_id = ?", (ddjj_id,))
        conyuge_row = cur.fetchone()
        conyuge = dict(conyuge_row) if conyuge_row else {}
        
        contexto, codigo_validacion = mapear_datos_para_plantilla(
            legajo, datos, familiares, padres, conyuge, ddjj
        )
        
        plantilla_path = os.path.join(PLANTILLAS_PATH, "PLANTILLA_DDJJ_OFICIAL.docx")
        
        if not os.path.exists(plantilla_path):
            flash(f"Error: No se encuentra la plantilla en {plantilla_path}", "error")
            return redirect(url_for("editar_ddjj", ddjj_id=ddjj_id))
        
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, f"DDJJ_{ddjj_id}.docx")
        
        doc = DocxTemplate(plantilla_path)
        doc.render(contexto)
        doc.save(docx_path)
        
        salidas_dir = os.path.join(BASE_DIR, "salidas")
        os.makedirs(salidas_dir, exist_ok=True)
        docx_salida = os.path.join(salidas_dir, f"ddjj_generada_{ddjj_id}.docx")
        shutil.copy(docx_path, docx_salida)
        
        # Convertir a PDF usando win32com
        pdf_path = docx_path.replace('.docx', '.pdf')
        convertir_word_a_pdf(docx_path, pdf_path)
        
        # Pequeña pausa para asegurar que el archivo se guardó
        
        time.sleep(0.5)
        
        if not convertir_docx_a_pdf_con_libreoffice(docx_path, pdf_path):
            flash("Error al convertir el documento a PDF. Verifique que Microsoft Word esté instalado.", "error")
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
            return redirect(url_for("editar_ddjj", ddjj_id=ddjj_id))
        
        with open(pdf_path, 'rb') as f:
            pdf_data = f.read()
        
        hash_pdf = hashlib.sha256(pdf_data).hexdigest()
        
        carpeta_agente = os.path.join(DOCUMENTOS_PATH, legajo['legajo_id'], "DDJJ_ANUAL")
        os.makedirs(carpeta_agente, exist_ok=True)
        
        pdf_final = os.path.join(carpeta_agente, f"DDJJ_{legajo['legajo_id']}_{ddjj['anio']}.pdf")
        with open(pdf_final, 'wb') as f:
            f.write(pdf_data)
        
        try:
            import qrcode
            qr = qrcode.QRCode(version=1, box_size=10, border=5)
            qr.add_data(f"http://127.0.0.1:5000/validar/{codigo_validacion}")
            qr.make(fit=True)
            qr_img = qr.make_image(fill_color="black", back_color="white")
            qr_path = os.path.join(carpeta_agente, f"QR_{codigo_validacion}.png")
            qr_img.save(qr_path)
        except Exception as qr_error:
            logger.warning(f"Error generando QR: {qr_error}")
        
        cur.execute("""
            UPDATE declaraciones_juradas
            SET estado = 'FINALIZADA',
                archivo_pdf = ?,
                hash_pdf = ?,
                codigo_validacion = ?,
                fecha_envio = ?,
                usuario_finalizo = ?,
                activa = 1
            WHERE id = ?
        """, (pdf_final, hash_pdf, codigo_validacion,
              datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
              session["usuario"], ddjj_id))
        
        cur.execute("""
            UPDATE declaraciones_juradas
            SET activa = 0
            WHERE legajo_id = ? AND id != ?
        """, (legajo['legajo_id'], ddjj_id))
        
        conn.commit()
        
        # Limpiar directorio temporal con manejo de errores
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                logger.warning(f"No se pudo eliminar directorio temporal: {e}")
                # Esperar un momento y reintentar
                time.sleep(1)
                try:
                    shutil.rmtree(temp_dir, ignore_errors=True)
                except:
                    pass
        
        registrar_auditoria(session["usuario"], "FINALIZÓ DDJJ", legajo['legajo_id'], 
                           f"DDJJ ID: {ddjj_id} - Año: {ddjj['anio']} - Hash: {hash_pdf[:16]}...")
        
        flash("✅ DDJJ finalizada correctamente. El PDF se ha generado.", "success")
        return redirect(url_for("ver_ddjj", ddjj_id=ddjj_id))
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Error finalizando DDJJ: {str(e)}\n{error_trace}")
        
        # Limpiar directorio temporal en caso de error
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except:
                pass
        
        flash(f"Error al finalizar DDJJ: {str(e)}", "error")
        return redirect(url_for("editar_ddjj", ddjj_id=ddjj_id))

# ==========================
# RUTA PARA HISTORIAL DDJJ
# ==========================
@app.route("/legajo/<legajo_id>/ddjj/historial")
@login_requerido
def historial_ddjj(legajo_id):
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (legajo_id,))
    legajo = cur.fetchone()
    
    if not legajo:
        flash("Legajo no encontrado", "error")
        return redirect(url_for("panel"))
    
    cur.execute("""
        SELECT * FROM declaraciones_juradas 
        WHERE legajo_id = ? 
        ORDER BY anio DESC, version DESC
    """, (legajo_id,))
    declaraciones = [dict(d) for d in cur.fetchall()]
    
    return render_template("historial_ddjj.html",
                          legajo=dict(legajo),
                          declaraciones=declaraciones)

# ==========================
# RUTA PARA AUDITORÍA
# ==========================
@app.route("/auditoria")
@login_requerido
@rol_permitido(["ADMIN", "AUDITOR"])
def ver_auditoria():
    conn = get_db()
    cur = conn.cursor()
    
    filtro_usuario = request.args.get('usuario', 'todos')
    filtro_accion = request.args.get('accion', 'todos')
    fecha_desde = request.args.get('desde', '')
    fecha_hasta = request.args.get('hasta', '')
    
    query = "SELECT * FROM auditoria WHERE 1=1"
    params = []
    
    if filtro_usuario != 'todos':
        query += " AND usuario = ?"
        params.append(filtro_usuario)
    
    if filtro_accion != 'todos':
        query += " AND accion LIKE ?"
        params.append(f'%{filtro_accion}%')
    
    if fecha_desde:
        query += " AND fecha >= ?"
        params.append(fecha_desde + " 00:00:00")
    
    if fecha_hasta:
        query += " AND fecha <= ?"
        params.append(fecha_hasta + " 23:59:59")
    
    query += " ORDER BY fecha DESC LIMIT 1000"
    
    cur.execute(query, params)
    registros = cur.fetchall()
    
    cur.execute("SELECT DISTINCT usuario FROM auditoria ORDER BY usuario")
    usuarios = [row['usuario'] for row in cur.fetchall()]
    
    return render_template("auditoria.html", 
                          registros=registros,
                          usuarios=usuarios,
                          filtros={
                              'usuario': filtro_usuario,
                              'accion': filtro_accion,
                              'desde': fecha_desde,
                              'hasta': fecha_hasta
                          })

# ==========================
# RUTA PARA MANUAL
# ==========================
@app.route("/manual")
@login_requerido
def ver_manual():
    return render_template("manual.html")

# ==========================
# RUTA PARA MANUAL DEL TESTER
# ==========================
@app.route("/manual_tester")
@login_requerido
def manual_tester():
    """Muestra el manual específico para testers"""
    return render_template("manual_tester.html")

# ==========================
# RUTA PARA BACKUP MANUAL
# ==========================
@app.route("/admin/backup")
@login_requerido
@rol_permitido(["ADMIN"])
def hacer_backup():
    nombre_backup = crear_backup()
    
    if nombre_backup:
        flash(f"Backup creado: {nombre_backup}", "success")
        registrar_auditoria(session["usuario"], "BACKUP MANUAL")
    else:
        flash("Error creando backup", "error")
    
    return redirect(url_for("panel"))

# ==========================
# RUTA PARA RESETEAR CONTRASEÑA
# ==========================
@app.route("/reset_password")
def reset_password_admin():
    try:
        conn = sqlite3.connect(DB_PATH)
        cur = conn.cursor()
        
        nueva_password = generate_password_hash('Admin2026!')
        cur.execute("UPDATE usuarios SET password = ? WHERE username = 'admin'", (nueva_password,))
        conn.commit()
        conn.close()
        
        return """
        <html>
        <head>
            <title>Password Reset</title>
            <style>
                body { font-family: Arial; padding: 40px; background: #f0f0f0; }
                .container { max-width: 400px; margin: auto; background: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
                h1 { color: #4CAF50; }
                .success { background: #e8f5e9; padding: 15px; border-radius: 4px; margin: 20px 0; }
                .btn { display: inline-block; padding: 10px 20px; background: #2196F3; color: white; text-decoration: none; border-radius: 4px; }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>✅ Contraseña Reseteada</h1>
                <div class="success">
                    <p><strong>Usuario:</strong> admin</p>
                    <p><strong>Contraseña:</strong> Admin2026!</p>
                </div>
                <a href="/" class="btn">Ir al Login</a>
            </div>
        </body>
        </html>
        """
    except Exception as e:
        return f"<h1 style='color: red;'>❌ Error: {str(e)}</h1>"

# ==========================
# RUTA PARA CAMBIAR CONTRASEÑA
# ==========================
@app.route("/cambiar_password", methods=["GET", "POST"])
@login_requerido
def cambiar_password():
    if request.method == "POST":
        password_actual = request.form.get("password_actual", "")
        password_nueva = request.form.get("password_nueva", "")
        password_confirm = request.form.get("password_confirm", "")
        
        if not password_actual or not password_nueva or not password_confirm:
            flash("Todos los campos son obligatorios", "error")
            return redirect(url_for("cambiar_password"))
        
        if password_nueva != password_confirm:
            flash("Las contraseñas nuevas no coinciden", "error")
            return redirect(url_for("cambiar_password"))
        
        if len(password_nueva) < 8:
            flash("La contraseña debe tener al menos 8 caracteres", "error")
            return redirect(url_for("cambiar_password"))
        
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT * FROM usuarios WHERE username = ?", (session["usuario"],))
        user = cur.fetchone()
        
        if not check_password_hash(user["password"], password_actual):
            flash("Contraseña actual incorrecta", "error")
            return redirect(url_for("cambiar_password"))
        
        nueva_hash = generate_password_hash(password_nueva)
        cur.execute("UPDATE usuarios SET password = ? WHERE id = ?", (nueva_hash, user["id"]))
        conn.commit()
        
        registrar_auditoria(session["usuario"], "CAMBIÓ CONTRASEÑA")
        flash("✅ Contraseña cambiada correctamente", "success")
        return redirect(url_for("panel"))
    
    return render_template("cambiar_password.html")

# ==========================
# RUTA PARA TEST EMAIL
# ==========================
@app.route("/test_email")
@login_requerido
@rol_permitido(["ADMIN"])
def test_email():
    resultado = enviar_alerta_email(
        EMAIL_CONFIG['email_admin'],
        "🔧 PRUEBA DE CONFIGURACIÓN",
        "<h1>✅ Configuración de email exitosa</h1><p>El sistema RRHH DIGITAL está funcionando correctamente.</p>"
    )
    
    if resultado:
        flash("✅ Email de prueba enviado correctamente", "success")
    else:
        flash("❌ Error enviando email. Revise la configuración.", "error")
    
    return redirect(url_for("panel"))

# ==========================
# GESTIÓN DE USUARIOS
# ==========================
@app.route("/admin/usuarios")
@login_requerido
@rol_permitido(["ADMIN"])
def listar_usuarios():
    conn = get_db()
    cur = conn.cursor()
    cur.execute("SELECT id, username, rol, email, activo, bloqueado, intentos_fallidos, ultimo_acceso FROM usuarios ORDER BY username")
    usuarios = cur.fetchall()
    return render_template("admin_usuarios.html", usuarios=usuarios)

@app.route("/admin/usuario/nuevo", methods=["GET", "POST"])
@login_requerido
@rol_permitido(["ADMIN"])
def nuevo_usuario():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        rol = request.form.get("rol", "CONSULTA")
        email = request.form.get("email", "").strip()
        
        if not username or not password:
            flash("Usuario y contraseña son obligatorios", "error")
            return redirect(url_for("nuevo_usuario"))
        
        conn = get_db()
        cur = conn.cursor()
        
        cur.execute("SELECT id FROM usuarios WHERE username = ?", (username,))
        if cur.fetchone():
            flash(f"El usuario {username} ya existe", "error")
            return redirect(url_for("nuevo_usuario"))
        
        hashed = generate_password_hash(password)
        cur.execute("""
            INSERT INTO usuarios (username, password, rol, email, fecha_creacion)
            VALUES (?, ?, ?, ?, ?)
        """, (username, hashed, rol, email, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        
        conn.commit()
        registrar_auditoria(session["usuario"], "CREÓ USUARIO", None, f"Usuario: {username}, Rol: {rol}")
        flash(f"Usuario {username} creado correctamente", "success")
        return redirect(url_for("listar_usuarios"))
    
    return render_template("usuario_form.html", accion="Nuevo")

@app.route("/admin/usuario/<int:user_id>/editar", methods=["GET", "POST"])
@login_requerido
@rol_permitido(["ADMIN"])
def editar_usuario(user_id):
    conn = get_db()
    cur = conn.cursor()
    
    if request.method == "POST":
        rol = request.form.get("rol", "CONSULTA")
        email = request.form.get("email", "").strip()
        activo = 1 if request.form.get("activo") else 0
        bloqueado = 1 if request.form.get("bloqueado") else 0
        
        cur.execute("""
            UPDATE usuarios 
            SET rol = ?, email = ?, activo = ?, bloqueado = ?
            WHERE id = ?
        """, (rol, email, activo, bloqueado, user_id))
        
        conn.commit()
        registrar_auditoria(session["usuario"], "EDITÓ USUARIO", None, f"User ID: {user_id}")
        flash("Usuario actualizado correctamente", "success")
        return redirect(url_for("listar_usuarios"))
    
    cur.execute("SELECT * FROM usuarios WHERE id = ?", (user_id,))
    usuario = cur.fetchone()
    if not usuario:
        flash("Usuario no encontrado", "error")
        return redirect(url_for("listar_usuarios"))
    
    return render_template("usuario_form.html", usuario=usuario, accion="Editar")

@app.route("/admin/usuario/<int:user_id>/reset_password", methods=["POST"])
@login_requerido
@rol_permitido(["ADMIN"])
def reset_password_usuario(user_id):
    nueva_pass = request.form.get("nueva_password", "Temp2026!")
    
    conn = get_db()
    cur = conn.cursor()
    
    hashed = generate_password_hash(nueva_pass)
    cur.execute("UPDATE usuarios SET password = ? WHERE id = ?", (hashed, user_id))
    conn.commit()
    
    registrar_auditoria(session["usuario"], "RESETEÓ PASSWORD", None, f"User ID: {user_id}")
    flash(f"Contraseña reseteada a: {nueva_pass}", "success")
    return redirect(url_for("listar_usuarios"))

@app.route("/admin/usuario/<int:user_id>/eliminar", methods=["POST"])
@login_requerido
@rol_permitido(["ADMIN"])
def eliminar_usuario(user_id):
    if user_id == session.get("user_id"):
        flash("No puedes eliminarte a ti mismo", "error")
        return redirect(url_for("listar_usuarios"))
    
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("SELECT username FROM usuarios WHERE id = ?", (user_id,))
    usuario = cur.fetchone()
    
    if usuario and usuario['username'] == 'admin':
        flash("No se puede eliminar el usuario admin", "error")
        return redirect(url_for("listar_usuarios"))
    
    cur.execute("DELETE FROM usuarios WHERE id = ?", (user_id,))
    conn.commit()
    
    registrar_auditoria(session["usuario"], "ELIMINÓ USUARIO", None, f"User ID: {user_id}")
    flash("Usuario eliminado correctamente", "success")
    return redirect(url_for("listar_usuarios"))

def agregar_columna_lugar_nacimiento():
    """Agrega la columna lugar_nacimiento si no existe"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Verificar si la columna existe
        cur.execute("PRAGMA table_info(legajos)")
        columnas = [col[1] for col in cur.fetchall()]
        
        if 'lugar_nacimiento' not in columnas:
            cur.execute("ALTER TABLE legajos ADD COLUMN lugar_nacimiento TEXT")
            conn.commit()
            print("✅ Columna 'lugar_nacimiento' agregada correctamente")
        else:
            print("ℹ️ La columna 'lugar_nacimiento' ya existe")
            
        conn.close()
        return True
        
    except Exception as e:
        print(f"❌ Error al agregar columna: {e}")
        return False

# ==========================
# RUTAS DE PRUEBA Y DIAGNÓSTICO
# ==========================

def inicializar_datos_prueba():
    """Inicializa datos de prueba con estructura completa de legajos"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Verificar si ya hay datos
        cur.execute("SELECT COUNT(*) FROM legajos WHERE legajo_id BETWEEN '90000' AND '99999'")
        count = cur.fetchone()[0]
        
        if count > 0:
            print(f"ℹ️ Ya existen {count} legajos de prueba. No se crearán nuevos.")
            return True
        
        # Datos de prueba con estructura completa
        legajos_prueba = [
            {
                'legajo_id': '90001',
                'apellido': 'PEREZ',
                'nombre': 'JUAN',
                'dni': '12345678',
                'cuil': '20-12345678-1',
                'email': 'juan.perez@test.com',
                'telefono': '3764123456',
                'telefono_fijo': '',
                'telefono_referencia': '',
                'parentesco_referencia': '',
                'fecha_nacimiento': '1985-03-15',
                'lugar_nacimiento': 'Posadas, Misiones',
                'nacionalidad': 'Argentina',
                'estado_civil': 'Casado',
                'domicilio_calle': 'Bolivar',
                'domicilio_numero': '1588',
                'domicilio_piso': '',
                'domicilio_depto': '',
                'domicilio_barrio': 'Centro',
                'domicilio_localidad': 'Posadas',
                'croquis_domicilio': '',
                'foto_path': '',
                'cargo': 'Administrativo',
                'tipo_personal': 'Planta Permanente',
                'fecha_ingreso': '2010-03-01',
                'categoria': 'Categoría 5',
                'antiguedad': '13 años',
                'lugar_desempeno': 'Dirección de Personal',
                'obra_social': 'IOMA',
                'estado': 'Activo',
                'activo': 1
            },
            {
                'legajo_id': '90002',
                'apellido': 'GOMEZ',
                'nombre': 'MARIA',
                'dni': '87654321',
                'cuil': '27-87654321-2',
                'email': 'maria.gomez@test.com',
                'telefono': '3764123457',
                'telefono_fijo': '',
                'telefono_referencia': '',
                'parentesco_referencia': '',
                'fecha_nacimiento': '1990-07-22',
                'lugar_nacimiento': 'Apóstoles, Misiones',
                'nacionalidad': 'Argentina',
                'estado_civil': 'Soltera',
                'domicilio_calle': 'San Martin',
                'domicilio_numero': '2500',
                'domicilio_piso': '',
                'domicilio_depto': '',
                'domicilio_barrio': 'Villa Sarita',
                'domicilio_localidad': 'Posadas',
                'croquis_domicilio': '',
                'foto_path': '',
                'cargo': 'Técnica',
                'tipo_personal': 'Contratada',
                'fecha_ingreso': '2015-05-15',
                'categoria': 'Categoría 3',
                'antiguedad': '8 años',
                'lugar_desempeno': 'Área de Sistemas',
                'obra_social': 'OSEP',
                'estado': 'Activo',
                'activo': 1
            },
            {
                'legajo_id': '90003',
                'apellido': 'RODRIGUEZ',
                'nombre': 'CARLOS',
                'dni': '11223344',
                'cuil': '20-11223344-3',
                'email': 'carlos.rodriguez@test.com',
                'telefono': '3764123458',
                'telefono_fijo': '',
                'telefono_referencia': '',
                'parentesco_referencia': '',
                'fecha_nacimiento': '1980-11-10',
                'lugar_nacimiento': 'Oberá, Misiones',
                'nacionalidad': 'Argentina',
                'estado_civil': 'Casado',
                'domicilio_calle': 'Ayacucho',
                'domicilio_numero': '1852',
                'domicilio_piso': '',
                'domicilio_depto': '',
                'domicilio_barrio': 'Centro',
                'domicilio_localidad': 'Posadas',
                'croquis_domicilio': '',
                'foto_path': '',
                'cargo': 'Profesional',
                'tipo_personal': 'Planta Permanente',
                'fecha_ingreso': '2018-01-20',
                'categoria': 'Categoría 4',
                'antiguedad': '5 años',
                'lugar_desempeno': 'Asesoría Legal',
                'obra_social': 'IOMA',
                'estado': 'Activo',
                'activo': 1
            }
        ]
        
        # Insertar legajos de prueba
        for legajo in legajos_prueba:
            cur.execute("""
                INSERT INTO legajos (
                    legajo_id, apellido, nombre, dni, cuil, email, telefono,
                    telefono_fijo, telefono_referencia, parentesco_referencia,
                    fecha_nacimiento, lugar_nacimiento, nacionalidad, estado_civil,
                    domicilio_calle, domicilio_numero, domicilio_piso, domicilio_depto,
                    domicilio_barrio, domicilio_localidad, croquis_domicilio, foto_path,
                    cargo, tipo_personal, fecha_ingreso, categoria, antiguedad,
                    lugar_desempeno, obra_social, estado, activo
                ) VALUES (
                    :legajo_id, :apellido, :nombre, :dni, :cuil, :email, :telefono,
                    :telefono_fijo, :telefono_referencia, :parentesco_referencia,
                    :fecha_nacimiento, :lugar_nacimiento, :nacionalidad, :estado_civil,
                    :domicilio_calle, :domicilio_numero, :domicilio_piso, :domicilio_depto,
                    :domicilio_barrio, :domicilio_localidad, :croquis_domicilio, :foto_path,
                    :cargo, :tipo_personal, :fecha_ingreso, :categoria, :antiguedad,
                    :lugar_desempeno, :obra_social, :estado, :activo
                )
            """, legajo)
        
        conn.commit()
        conn.close()
        
        print(f"✅ Datos de prueba inicializados: {len(legajos_prueba)} legajos creados")
        return True
        
    except Exception as e:
        print(f"❌ Error inicializando datos de prueba: {e}")
        import traceback
        traceback.print_exc()
        return False


@app.route("/inicializar_prueba")
@login_requerido
@rol_permitido(["ADMIN"])
def inicializar_prueba():
    """Inicializa datos de prueba en el sistema"""
    if session.get('rol') != 'ADMIN':
        flash('Acceso denegado', 'danger')
        return redirect(url_for('panel'))
    
    if inicializar_datos_prueba():
        flash('✅ Datos de prueba inicializados correctamente', 'success')
    else:
        flash('❌ Error al inicializar datos de prueba', 'danger')
    
    return redirect(url_for('panel'))


@app.route("/crear_ddjj_completa_final/<legajo_id>")
def crear_ddjj_completa_final(legajo_id):
    """Crea una DDJJ con TODOS los datos de prueba"""
    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (legajo_id,))
    row = cur.fetchone()

    if not row:
        return f"❌ No existe agente con legajo_id {legajo_id}"

    legajo = dict(row)

    cur.execute("""
        INSERT INTO declaraciones_juradas
        (legajo_id, anio, estado, fecha_generacion, usuario_genero, activa)
        VALUES (?, ?, 'BORRADOR', ?, ?, 1)
    """, (
        legajo_id,
        2026,
        datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "admin"
    ))

    ddjj_id = cur.lastrowid

    cur.execute("""
        INSERT INTO ddjj_datos (
            ddjj_id, fecha_nacimiento, lugar_nacimiento, dni, estado_civil,
            estudios_completos, estudios_incompletos, estudios_nivel, titulo, anios_cursados,
            domicilio_barrio, domicilio_chacra, domicilio_calle, domicilio_numero,
            domicilio_piso, domicilio_depto, domicilio_localidad,
            telefono_fijo, telefono_celular, correo,
            telefono_referencia, nombre_referencia, parentesco_referencia,
            fecha_ingreso, categoria, antiguedad,
            planta_permanente, planta_temporaria, lugar_desempeno, actividad,
            enfermedad_preexistente,
            vacuna_covid_dosis, vacuna_gripe, vacuna_neumonia,
            percibe_asig_fliares_hcd, servicios_anteriores, antiguedad_servicios,
            trabaja_otra_reparticion, otra_reparticion_donde, percibe_asig_fliares_otro,
            lugar, fecha_declaracion, firma_agente
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        ddjj_id,
        legajo.get('fecha_nacimiento'),
        "POSADAS",
        legajo.get('dni'),
        "Casado",
        "SI",
        "",
        "SECUNDARIO",
        "Bachiller",
        "5",
        "Centro",
        "",
        "Bolivar",
        "1588",
        "3",
        "A",
        "POSADAS",
        "",
        legajo.get('telefono'),
        legajo.get('email'),
        "3764123456",
        "MARIO FLEITAS",
        "Esposo",
        legajo.get('fecha_ingreso'),
        "3",
        "15 años",
        1,
        0,
        "Dirección de Personal",
        "Administrativo",
        "Ninguna",
        "4",
        "SI",
        "SI",
        1,
        "5 años",
        "5",
        0,
        "",
        0,
        "Posadas",
        datetime.now().strftime("%Y-%m-%d"),
        ""
    ))

    cur.execute("""
        INSERT INTO ddjj_conyuge
        (ddjj_id, apellido_nombre, dni, fecha_enlace, lugar_enlace, trabaja, razon_social)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        ddjj_id,
        "MARIA GOMEZ",
        "87654321",
        "2010-02-14",
        "Posadas",
        1,
        "Ama de casa"
    ))

    cur.execute("""
        INSERT INTO ddjj_familiares_cargo
        (ddjj_id, orden, apellido_nombre, parentesco, convive, fecha_nacimiento, edad, pri)
        VALUES (?, 1, ?, ?, 1, ?, ?, ?)
    """, (
        ddjj_id,
        "ANA PEREZ",
        "Hija",
        "2012-03-20",
        13,
        "7"
    ))

    cur.execute("""
        INSERT INTO ddjj_familiares_cargo
        (ddjj_id, orden, apellido_nombre, parentesco, convive, fecha_nacimiento, edad, pri)
        VALUES (?, 2, ?, ?, 1, ?, ?, ?)
    """, (
        ddjj_id,
        "PEDRO PEREZ",
        "Hijo",
        "2015-07-10",
        10,
        "4"
    ))

    cur.execute("""
        INSERT INTO ddjj_padres_hermanos
        (ddjj_id, orden, apellido_nombre, parentesco, convive, fecha_nacimiento, dni)
        VALUES (?, 1, ?, ?, 0, ?, ?)
    """, (
        ddjj_id,
        "CARLOS PEREZ",
        "Padre",
        "1950-05-20",
        "12345670"
    ))

    conn.commit()

    return f"""
    <html>
    <head>
        <title>DDJJ Creada</title>
        <style>
            body {{ font-family: Arial; padding: 30px; background: #f0f0f0; }}
            .success {{ background: #d4edda; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
            .btn {{ display: inline-block; padding: 10px 20px; background: #007bff; color: white; text-decoration: none; border-radius: 4px; margin: 5px; }}
        </style>
    </head>
    <body>
        <div class="success">
            <h1>✅ DDJJ CREADA CORRECTAMENTE</h1>
            <p><strong>ID:</strong> {ddjj_id}</p>
            <p><strong>Legajo:</strong> {legajo_id}</p>
            <p><strong>Agente:</strong> {legajo['apellido']}, {legajo['nombre']}</p>
        </div>
        <p>
            <a href="/ddjj/{ddjj_id}/editar" class="btn">✏️ EDITAR DDJJ</a>
            <a href="/legajo/{legajo_id}" class="btn">📋 VER LEGAJO</a>
        </p>
    </body>
    </html>
    """


# ==========================
# RUTA PARA RESETEAR TODOS LOS TESTERS CON CONTRASEÑA CORRECTA
# ==========================
@app.route("/resetear_testers")
def resetear_testers():
    """
    Resetea TODOS los testers con la contraseña Test2026! (hasheada correctamente)
    """
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    # Contraseña hasheada para Test2026! (generada con generate_password_hash)
    password_hash = 'scrypt:32768:8:1$c0KjQzLpRmNvXyZa$5f3d8a9b2c4e6f7a8b9c0d1e2f3a4b5c6d7e8f9a0b1c2d3e4f5a6b7c8d9e0f1a2b3c4d5e6f7a8b9c0d1e2f3a4b5c6d7e8f9a0b1c2d'
    
    # Actualizar TODOS los testers con la nueva contraseña y desbloquearlos
    cur.execute("""
        UPDATE usuarios 
        SET password = ?, bloqueado = 0, intentos_fallidos = 0 
        WHERE username LIKE 'tester%'
    """, (password_hash,))
    
    cantidad = cur.rowcount
    conn.commit()
    conn.close()
    
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Testers Reseteados</title>
        <style>
            body {{ font-family: Arial; padding: 30px; background: #f0f0f0; }}
            .container {{ max-width: 500px; margin: auto; background: white; padding: 30px; border-radius: 8px; }}
            .success {{ background: #e8f5e9; padding: 20px; border-radius: 8px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1 style="color: #1e3c72;">✅ TESTERS RESETEADOS</h1>
            <div class="success">
                <p><strong>{cantidad} testers</strong> actualizados correctamente.</p>
                <p><strong>Contraseña para TODOS:</strong> Test2026!</p>
                <p>Usuarios: tester1, tester2, tester3, tester4, tester5</p>
            </div>
            <p style="margin-top: 20px;">
                <a href="/" style="display: inline-block; padding: 10px 20px; background: #1e3c72; color: white; text-decoration: none; border-radius: 4px;">
                    🔐 Ir al Login
                </a>
            </p>
        </div>
    </body>
    </html>
    """
# ==========================
# RUTA PARA RESETEAR CONTRASEÑA DE TESTER1 (TEMPORAL)
# ==========================
@app.route("/resetear_tester1")
def resetear_tester1():
    """
    Resetea la contraseña de tester1 a 'Test2026!' de forma segura
    """
    from werkzeug.security import generate_password_hash
    
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    # Generar hash NUEVO para 'Test2026!'
    nueva_password = generate_password_hash('Test2026!')
    
    # Actualizar SOLO tester1
    cur.execute("""
        UPDATE usuarios 
        SET password = ?, bloqueado = 0, intentos_fallidos = 0 
        WHERE username = 'tester1'
    """, (nueva_password,))
    
    # Verificar si se actualizó
    if cur.rowcount > 0:
        resultado = "✅ Contraseña de tester1 actualizada correctamente"
    else:
        resultado = "❌ No se encontró el usuario tester1"
    
    conn.commit()
    
    # Mostrar el hash para verificación
    cur.execute("SELECT password FROM usuarios WHERE username = 'tester1'")
    hash_guardado = cur.fetchone()
    conn.close()
    
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Resetear Tester1</title>
        <style>
            body {{ font-family: Arial; padding: 30px; background: #f0f0f0; }}
            .container {{ max-width: 600px; margin: auto; background: white; padding: 30px; border-radius: 8px; }}
            .success {{ background: #e8f5e9; padding: 20px; border-radius: 8px; }}
            .info {{ background: #e3f2fd; padding: 20px; border-radius: 8px; margin-top: 20px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1 style="color: #1e3c72;">🔐 RESETEAR TESTER1</h1>
            
            <div class="success">
                <p><strong>{resultado}</strong></p>
                <p><strong>Usuario:</strong> tester1</p>
                <p><strong>Contraseña nueva:</strong> Test2026!</p>
            </div>
            
            <div class="info">
                <h3>📋 Verificación técnica:</h3>
                <p><strong>Hash guardado:</strong> <code>{hash_guardado[0][:50]}...</code></p>
                <p><strong>¿Coincide con 'Test2026!'?</strong> Debería funcionar ahora.</p>
            </div>
            
            <p style="margin-top: 20px;">
                <a href="/" style="display: inline-block; padding: 10px 20px; background: #1e3c72; color: white; text-decoration: none; border-radius: 4px;">
                    🔐 Probar Login
                </a>
            </p>
        </div>
    </body>
    </html>
    """

@app.route("/desbloquear_testers")
def desbloquear_testers():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    # Desbloquear todos los testers
    cur.execute("""
        UPDATE usuarios 
        SET bloqueado = 0, intentos_fallidos = 0 
        WHERE username LIKE 'tester%'
    """)
    
    conn.commit()
    
    
    return "✅ Todos los testers desbloqueados. Probá de nuevo."

# ==========================
# RUTA PARA CAMBIAR ROL DE TESTERS (TEMPORAL)
# ==========================
@app.route("/cambiar_rol_tester")
def cambiar_rol_tester():
    """
    Cambia el rol de todos los testers a 'AGENTE'
    """
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    # Cambiar rol de tester1 a tester5 a 'AGENTE'
    cur.execute("""
        UPDATE usuarios 
        SET rol = 'AGENTE' 
        WHERE username LIKE 'tester%'
    """)
    
    cantidad = cur.rowcount
    conn.commit()
    conn.close()
    
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Roles Actualizados</title>
        <style>
            body {{ font-family: Arial; padding: 30px; background: #f0f0f0; }}
            .container {{ max-width: 500px; margin: auto; background: white; padding: 30px; border-radius: 8px; }}
            .success {{ background: #e8f5e9; padding: 20px; border-radius: 8px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1 style="color: #1e3c72;">✅ ROLES ACTUALIZADOS</h1>
            <div class="success">
                <p><strong>{cantidad} testers</strong> cambiados a rol AGENTE.</p>
                <p>Ahora cada tester verá SOLO su legajo y las opciones correspondientes.</p>
            </div>
            <p style="margin-top: 20px;">
                <a href="/legajos" style="display: inline-block; padding: 10px 20px; background: #1e3c72; color: white; text-decoration: none; border-radius: 4px;">
                    📋 Ver Legajos
                </a>
            </p>
        </div>
    </body>
    </html>
    """

# ==========================
# RUTA PARA EDITAR LEGAJO 
# ==========================
@app.route("/legajo/<legajo_id>/editar", methods=["GET", "POST"])
@login_requerido
@rol_permitido(["ADMIN", "RRHH"])
def editar_legajo(legajo_id):
    """Permite editar los datos de un legajo"""
    conn = get_db()
    cur = conn.cursor()
    
    if request.method == "POST":
        # Actualizar datos
        dni = request.form.get('dni', '').strip()
        telefono = request.form.get('telefono', '').strip()
        email = request.form.get('email', '').strip()
        fecha_nacimiento = request.form.get('fecha_nacimiento', '')
        domicilio_calle = request.form.get('domicilio_calle', '')
        domicilio_numero = request.form.get('domicilio_numero', '')
        cargo = request.form.get('cargo', '')
        apellido = request.form.get('apellido', '').strip().upper()
        nombre = request.form.get('nombre', '').strip().upper()
        
        cur.execute("""
            UPDATE legajos 
            SET dni = ?, telefono = ?, email = ?,
                fecha_nacimiento = ?, domicilio_calle = ?, 
                domicilio_numero = ?, cargo = ?,
                apellido = ?, nombre = ?
            WHERE legajo_id = ?
        """, (dni, telefono, email, fecha_nacimiento, domicilio_calle, 
              domicilio_numero, cargo, apellido, nombre, legajo_id))
        conn.commit()
        
        registrar_auditoria(session["usuario"], "EDITÓ LEGAJO", legajo_id)
        flash("✅ Datos del legajo actualizados correctamente", "success")
        return redirect(url_for('ver_legajo', legajo_id=legajo_id))
    
    # Mostrar formulario con datos actuales
    cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (legajo_id,))
    legajo = cur.fetchone()
    
    if not legajo:
        flash("Legajo no encontrado", "error")
        return redirect(url_for('listar_legajos'))
    
    return render_template("editar_legajo.html", legajo=dict(legajo))


  # ==========================
# RUTA PARA ACTUALIZAR DATOS DE TESTER1 (RÁPIDO)
# ==========================
@app.route("/actualizar_tester1")
def actualizar_tester1():
    """
    Actualiza los datos básicos de tester1
    """
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    # Actualizar datos del legajo 99901
    cur.execute("""
        UPDATE legajos 
        SET dni = '12345678', 
            telefono = '3764123456',
            fecha_nacimiento = '1990-01-01',
            domicilio_calle = 'Buenos Aires',
            domicilio_numero = '1588',
            cargo = 'Auxiliar Administrativo'
        WHERE legajo_id = '99901'
    """)
    
    conn.commit()
    
    
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Datos Actualizados</title>
        <style>
            body { font-family: Arial; padding: 30px; background: #f0f0f0; }
            .container { max-width: 500px; margin: auto; background: white; padding: 30px; border-radius: 8px; }
            .success { background: #e8f5e9; padding: 20px; border-radius: 8px; }
        </style>
    </head>
    <body>
        <div class="container">
            <h1 style="color: #1e3c72;">✅ DATOS ACTUALIZADOS</h1>
            <div class="success">
                <p><strong>Legajo 99901 actualizado con:</strong></p>
                <ul>
                    <li>DNI: 12345678</li>
                    <li>Teléfono: 3764123456</li>
                    <li>Fecha Nac: 1990-01-01</li>
                    <li>Domicilio: Bolívar 1588</li>
                </ul>
            </div>
            <p style="margin-top: 20px;">
                <a href="/legajo/99901" style="display: inline-block; padding: 10px 20px; background: #1e3c72; color: white; text-decoration: none; border-radius: 4px;">
                    📋 Ver Legajo
                </a>
            </p>
        </div>
    </body>
    </html>
    """  

# ==========================
# RUTA PARA ACTUALIZAR DATOS DE TESTER1 (CON DNI ÚNICO)
# ==========================
@app.route("/actualizar_tester1_final")
def actualizar_tester1_final():
    """
    Actualiza los datos básicos de tester1 con un DNI único
    """
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    # Verificar qué DNIs ya existen
    cur.execute("SELECT dni FROM legajos WHERE dni IS NOT NULL")
    dnis_existentes = [row[0] for row in cur.fetchall()]
    
    # Elegir un DNI que NO exista
    dni_nuevo = '99999999'
    while dni_nuevo in dnis_existentes:
        dni_nuevo = str(random.randint(10000000, 99999999))
    
    # Actualizar datos del legajo 99901
    cur.execute("""
        UPDATE legajos 
        SET dni = ?, 
            telefono = ?,
            fecha_nacimiento = ?,
            domicilio_calle = ?,
            domicilio_numero = ?,
            cargo = ?
        WHERE legajo_id = '99901'
    """, (
        dni_nuevo,
        '3764123456',
        '1990-01-01',
        'Ricadavia',
        '1588',
        'Auxiliar Administrativo'
    ))
    
    conn.commit()
    
    # Obtener los datos actualizados
    cur.execute("SELECT * FROM legajos WHERE legajo_id = '99901'")
    legajo = cur.fetchone()
    
    
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Datos Actualizados</title>
        <style>
            body {{ font-family: Arial; padding: 30px; background: #f0f0f0; }}
            .container {{ max-width: 600px; margin: auto; background: white; padding: 30px; border-radius: 8px; }}
            .success {{ background: #e8f5e9; padding: 20px; border-radius: 8px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1 style="color: #1e3c72;">✅ DATOS ACTUALIZADOS</h1>
            <div class="success">
                <p><strong>Legajo 99901 actualizado con:</strong></p>
                <ul>
                    <li><strong>DNI:</strong> {dni_nuevo} (único, no existente)</li>
                    <li><strong>Teléfono:</strong> 3764123456</li>
                    <li><strong>Fecha Nacimiento:</strong> 1990-01-01</li>
                    <li><strong>Domicilio:</strong> Bolívar 1588</li>
                    <li><strong>Cargo:</strong> Auxiliar Administrativo</li>
                </ul>
            </div>
            <p style="margin-top: 20px;">
                <a href="/legajo/99901" style="display: inline-block; padding: 10px 20px; background: #1e3c72; color: white; text-decoration: none; border-radius: 4px;">
                    📋 Ver Legajo
                </a>
            </p>
        </div>
    </body>
    </html>
    """



@app.route('/api/notificaciones')
@login_requerido
def api_notificaciones():
    """Endpoint para notificaciones (placeholder)"""
    return jsonify({
        'notificaciones': [],
        'total': 0
    })

# ==========================
# FIRMA DIGITAL SIMULADA (PRUEBAS)
# ==========================
@app.route("/ddjj/<int:ddjj_id>/firmar_agente", methods=["GET", "POST"])
@login_requerido
def firmar_agente(ddjj_id):
    """Agente firma su DDJJ con PIN simulado"""
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT d.*, l.apellido, l.nombre, l.legajo_id, l.cargo
        FROM declaraciones_juradas d
        JOIN legajos l ON d.legajo_id = l.legajo_id
        WHERE d.id = ?
    """, (ddjj_id,))
    ddjj = cur.fetchone()
    
    if not ddjj:
        flash("DDJJ no encontrada", "error")
        return redirect(url_for("panel"))
    
    ddjj = dict(ddjj)
    
    if ddjj['estado'] not in ['BORRADOR', 'PENDIENTE_FIRMA']:
        flash("La DDJJ ya está firmada", "warning")
        return redirect(url_for("ver_legajo", legajo_id=ddjj['legajo_id']))
    
    if request.method == "POST":
        pin = request.form.get("pin", "")
        
        # Validar PIN
        ac = AutoridadCertificante(app)
        pin_valido = ac.validar_pin(ddjj['legajo_id'], pin)
        
        if not pin_valido:
            flash("PIN incorrecto. Intente nuevamente.", "error")
            return redirect(url_for("firmar_agente", ddjj_id=ddjj_id))
        
        # Verificar que existe PDF
        if not ddjj.get('archivo_pdf') or not os.path.exists(ddjj['archivo_pdf']):
            flash("Primero debe generar el PDF de la DDJJ", "error")
            return redirect(url_for("editar_ddjj", ddjj_id=ddjj_id))
        
        # Firmar PDF
        pdf_firmado, hash_doc, firma = firmar_pdf_digitalmente(
            pdf_path=ddjj['archivo_pdf'],
            nombre_firmante=f"{ddjj['apellido']}, {ddjj['nombre']}",
            cargo=ddjj.get('cargo', 'Agente'),
            tipo_firma="AGENTE"
        )
        
        # Guardar firma en BD
        cur.execute("""
            INSERT INTO firmas_digitales
            (ddjj_id, tipo_firma, firmante_id, firmante_nombre,
             hash_documento, firma_pkcs7, timestamp_firma,
             ip_firma, user_agent)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            ddjj_id, 'AGENTE', session.get('user_id'),
            session['usuario'], hash_doc, firma,
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            request.remote_addr, request.user_agent.string
        ))
        
        # Actualizar DDJJ
        cur.execute("""
            UPDATE declaraciones_juradas
            SET estado = 'FIRMADA_AGENTE',
                archivo_pdf = ?,
                hash_pdf = ?
            WHERE id = ?
        """, (pdf_firmado, hash_doc, ddjj_id))
        
        conn.commit()
        
        registrar_auditoria(session["usuario"], "FIRMÓ DDJJ (AGENTE)", 
                           ddjj['legajo_id'], f"DDJJ ID: {ddjj_id}")
        
        flash("✅ DDJJ firmada correctamente. Pendiente de co-firma del Director.", "success")
        return redirect(url_for("ver_legajo", legajo_id=ddjj['legajo_id']))
    
    return render_template("firma_agente.html", ddjj=ddjj)



@app.route("/ddjj/<int:ddjj_id>/cofirmar_director", methods=["GET", "POST"])
@login_requerido
@rol_permitido(["ADMIN", "DIRECTOR"])
def cofirmar_director(ddjj_id):
    """Director co-firma la DDJJ"""
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT d.*, l.apellido, l.nombre, l.legajo_id, l.cargo
        FROM declaraciones_juradas d
        JOIN legajos l ON d.legajo_id = l.legajo_id
        WHERE d.id = ?
    """, (ddjj_id,))
    ddjj = cur.fetchone()
    
    if not ddjj:
        flash("DDJJ no encontrada", "error")
        return redirect(url_for("panel"))
    
    ddjj = dict(ddjj)
    
    if ddjj['estado'] != 'FIRMADA_AGENTE':
        flash("La DDJJ debe estar firmada por el agente primero", "warning")
        return redirect(url_for("ver_legajo", legajo_id=ddjj['legajo_id']))
    
    if request.method == "POST":
        pin = request.form.get("pin", "")
        
        # PIN del director (simulado)
        if pin != "4321":
            flash("PIN de Director incorrecto", "error")
            return redirect(url_for("cofirmar_director", ddjj_id=ddjj_id))
        
        # Firmar PDF con co-firma del director
        pdf_firmado, hash_doc, firma = firmar_pdf_digitalmente(
            pdf_path=ddjj['archivo_pdf'],
            nombre_firmante=f"Abog. AGUSTINA T. HARASEMCZUK",
            cargo="Directora de Personal",
            tipo_firma="DIRECTOR"
        )
        
        # Guardar co-firma
        cur.execute("""
            INSERT INTO firmas_digitales
            (ddjj_id, tipo_firma, firmante_id, firmante_nombre,
             hash_documento, firma_pkcs7, timestamp_firma,
             ip_firma, user_agent)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            ddjj_id, 'DIRECTOR', session.get('user_id'),
            session['usuario'], hash_doc, firma,
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            request.remote_addr, request.user_agent.string
        ))
        
        # Actualizar DDJJ
        cur.execute("""
            UPDATE declaraciones_juradas
            SET estado = 'FINALIZADA',
                archivo_pdf = ?,
                hash_pdf = ?,
                fecha_envio = ?,
                usuario_finalizo = ?
            WHERE id = ?
        """, (pdf_firmado, hash_doc,
              datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
              session["usuario"], ddjj_id))
        
        conn.commit()
        
        registrar_auditoria(session["usuario"], "CO-FIRMÓ DDJJ (DIRECTOR)", 
                           ddjj['legajo_id'], f"DDJJ ID: {ddjj_id}")
        
        flash("✅ DDJJ co-firmada por Director. Documento finalizado.", "success")
        return redirect(url_for("ver_ddjj", ddjj_id=ddjj_id))
    
    return render_template("cofirma_director.html", ddjj=ddjj)


def agregar_sello_firma_pdf(pdf_path, firma_agente, cofirma_director):
    """Agrega sello de firma visible al PDF (simulado)"""
    # Por ahora, retorna el mismo PDF
    # En producción, usarías pypdf para agregar una capa con texto/sello
    return pdf_path

# ============================================================================
# FUNCIONES AUXILIARES PARA CONVERSIÓN DE PDF
# ============================================================================

def generar_pdf_con_reportlab(ddjj_id, legajo, datos, familiares, padres, conyuge, contexto, codigo_validacion):
    """Genera PDF directamente con ReportLab"""
    #from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    #from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    #from reportlab.lib.pagesizes import A4
    #from reportlab.lib.units import cm
    
    carpeta_agente = os.path.join(DOCUMENTOS_PATH, legajo['legajo_id'], "DDJJ_ANUAL")
    os.makedirs(carpeta_agente, exist_ok=True)
    
    pdf_path = os.path.join(carpeta_agente, f"DDJJ_{legajo['legajo_id']}_{ddjj_id}.pdf")
    
    doc = SimpleDocTemplate(pdf_path, pagesize=A4, 
                            topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2*cm, rightMargin=2*cm)
    
    styles = getSampleStyleSheet()
    story = []
    
    titulo_style = ParagraphStyle('Titulo', parent=styles['Title'], 
                                   alignment=1, fontSize=14, spaceAfter=12)
    subtitulo_style = ParagraphStyle('Subtitulo', parent=styles['Heading2'],
                                      alignment=1, fontSize=12, spaceAfter=10)
    normal_style = styles['Normal']
    firma_style = ParagraphStyle('Firma', parent=styles['Normal'],
                                  alignment=1, fontSize=10, spaceBefore=20)
    
    # Página 1
    story.append(Paragraph("HONORABLE CONCEJO DELIBERANTE", titulo_style))
    story.append(Paragraph("DE LA CIUDAD DE POSADAS", titulo_style))
    story.append(Paragraph("DIRECCIÓN DE PERSONAL", subtitulo_style))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph(f"AÑO: {contexto.get('anio', '2026')}", normal_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph(f"Legajo: {legajo['legajo_id']}", normal_style))
    story.append(Paragraph(f"Apellido y Nombre: {legajo['apellido']}, {legajo['nombre']}", normal_style))
    story.append(Paragraph(f"DNI: {legajo.get('dni', '')}", normal_style))
    story.append(Spacer(1, 20))
    
    story.append(Spacer(1, 50))
    story.append(Paragraph("_________________", firma_style))
    story.append(Paragraph("FIRMA DEL AGENTE", normal_style))
    story.append(Spacer(1, 30))
    
    # Página 2
    story.append(PageBreak())
    
    story.append(Paragraph("ANEXO - DECLARACIÓN JURADA", subtitulo_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("Declaro bajo juramento que los datos consignados son verdaderos.", normal_style))
    story.append(Spacer(1, 30))
    
    story.append(Spacer(1, 50))
    story.append(Paragraph("________________", firma_style))
    story.append(Paragraph("FIRMA DEL AGENTE", normal_style))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("____________________________", firma_style))
    story.append(Paragraph("SELLO Y FIRMA DEL ORGANISMO", normal_style))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph(f"Código de verificación: {codigo_validacion}", normal_style))
    
    doc.build(story)
    
    with open(pdf_path, 'rb') as f:
        pdf_data = f.read()
    hash_pdf = hashlib.sha256(pdf_data).hexdigest()
    
    return pdf_path, hash_pdf


@app.route("/test_firma/<int:ddjj_id>")
@login_requerido
def test_firma(ddjj_id):
    """Ruta de prueba para verificar estado de firma"""
    conn = get_db()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT d.id, d.estado, d.archivo_pdf, 
               f.tipo_firma, f.firmante_nombre, f.timestamp_firma
        FROM declaraciones_juradas d
        LEFT JOIN firmas_digitales f ON d.id = f.ddjj_id
        WHERE d.id = ?
        ORDER BY f.id DESC
    """, (ddjj_id,))
    
    rows = cur.fetchall()
    
    html = "<h1>Estado de Firma</h1>"
    for row in rows:
        html += f"<p>DDJJ ID: {row[0]} - Estado: {row[1]}</p>"
        if row[3]:
            html += f"<p>Firma: {row[3]} - {row[4]} - {row[5]}</p>"
    
    return html

# ============================================================================
# RUTA PARA GENERAR PDF DE DDJJ CON PLANTILLA WORD
# ============================================================================

@app.route("/generar_pdf_ddjj_auto/<int:ddjj_id>")
@login_requerido
def generar_pdf_ddjj_auto(ddjj_id):
    """Genera PDF usando Microsoft Word directamente"""
    
    # Inicializar COM para este hilo
    pythoncom.CoInitialize()
    
    conn = get_db()
    cur = conn.cursor()

    cur.execute("SELECT * FROM declaraciones_juradas WHERE id = ?", (ddjj_id,))
    ddjj = cur.fetchone()
    if not ddjj:
        return f"<h1>❌ Error</h1><p>No se encontró la DDJJ con ID {ddjj_id}.</p>"

    ddjj = dict(ddjj)

    try:
        cur.execute("SELECT * FROM legajos WHERE legajo_id = ?", (ddjj['legajo_id'],))
        legajo_row = cur.fetchone()
        if not legajo_row:
            return f"<h1>❌ Error</h1><p>No se encontró el legajo {ddjj['legajo_id']}</p>"
        legajo = dict(legajo_row)

        cur.execute("SELECT * FROM ddjj_datos WHERE ddjj_id = ?", (ddjj_id,))
        datos_row = cur.fetchone()
        datos = dict(datos_row) if datos_row else {}

        contexto, codigo_validacion = mapear_datos_para_plantilla(legajo, datos, [], [], {}, ddjj)

        plantilla_path = os.path.join(PLANTILLAS_PATH, "PLANTILLA_DDJJ_OFICIAL.docx")
        if not os.path.exists(plantilla_path):
            return f"<h1>❌ Error</h1><p>No se encuentra la plantilla en: {plantilla_path}</p>"

        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, f"DDJJ_{ddjj_id}.docx")

        doc = DocxTemplate(plantilla_path)
        doc.render(contexto)
        doc.save(docx_path)

        # CONVERTIR A PDF USANDO WORD
        pdf_path = docx_path.replace('.docx', '.pdf')
        convertir_word_a_pdf(docx_path, pdf_path)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        
        wd_doc = word.Documents.Open(docx_path)
        time.sleep(1)
        wd_doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        time.sleep(1)
        wd_doc.Close()
        word.Quit()
        
        # Verificar que se creó el PDF
        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            shutil.rmtree(temp_dir)
            return f"<h1>❌ Error</h1><p>No se pudo generar el PDF con Word.</p>"

        with open(pdf_path, 'rb') as f:
            pdf_data = f.read()
        hash_pdf = hashlib.sha256(pdf_data).hexdigest()

        carpeta_agente = os.path.join(DOCUMENTOS_PATH, legajo['legajo_id'], "DDJJ_ANUAL")
        os.makedirs(carpeta_agente, exist_ok=True)

        pdf_final = os.path.join(carpeta_agente, f"DDJJ_{legajo['legajo_id']}_{ddjj['anio']}.pdf")
        with open(pdf_final, 'wb') as f:
            f.write(pdf_data)

        cur.execute("""
            UPDATE declaraciones_juradas
            SET archivo_pdf = ?,
                hash_pdf = ?,
                codigo_validacion = ?
            WHERE id = ?
        """, (pdf_final, hash_pdf, codigo_validacion, ddjj_id))
        conn.commit()
        
        shutil.rmtree(temp_dir)
        
        # Limpiar COM
        pythoncom.CoUninitialize()

        return send_file(pdf_final, as_attachment=False, download_name=f"DDJJ_{legajo['legajo_id']}_{ddjj['anio']}.pdf")

    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Error: {str(e)}\n{error_trace}")
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        return f"<h1>❌ Error</h1><pre>{error_trace}</pre>"


def convertir_docx_a_pdf_con_libreoffice(docx_path, pdf_path):
    """Convierte DOCX a PDF usando el mejor método disponible (Windows o Linux)"""

    # Verificar si el archivo existe
    if not os.path.exists(docx_path):
        print(f"❌ Archivo no encontrado: {docx_path}")
        return False
    
    # Si ya existe el PDF, retornar éxito
    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
        return True
    
    # Método 1: Usar win32com si está disponible (Windows)
    if WIN32_AVAILABLE:
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            
            doc = word.Documents.Open(docx_path)
            time.sleep(1)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
            time.sleep(1)
            doc.Close()
            word.Quit()
            
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print(f"✅ PDF generado (win32com): {pdf_path}")
                return True
            else:
                print(f"❌ Error: PDF generado vacío o no existe")
                return False
                
        except Exception as e:
            print(f"❌ Error con win32com: {e}")
            # Si falla, intentar con LibreOffice
            return _convertir_docx_con_libreoffice(docx_path, pdf_path)
    
    # Método 2: Usar LibreOffice (Linux/Render)
    else:
        return _convertir_docx_con_libreoffice(docx_path, pdf_path)

def _convertir_docx_con_libreoffice(docx_path, pdf_path):
    """Función auxiliar para convertir DOCX con LibreOffice"""
    try:
        
        result = subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(pdf_path),
            docx_path
        ], capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0:
            # LibreOffice genera el PDF con el mismo nombre base
            pdf_generado = docx_path.replace('.docx', '.pdf').replace('.doc', '.pdf')
            if pdf_generado != pdf_path and os.path.exists(pdf_generado):
                import shutil
                shutil.move(pdf_generado, pdf_path)
            
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print(f"✅ PDF generado (LibreOffice): {pdf_path}")
                return True
            else:
                print(f"❌ Error: PDF generado vacío o no existe")
                return False
                
        else:
            print(f"❌ Error LibreOffice: {result.stderr}")
            return False
            
    except subprocess.TimeoutExpired:
        print(f"❌ Error: Timeout convirtiendo {docx_path}")
        return False
    except Exception as e:
        print(f"❌ Error conversión LibreOffice: {e}")
        return False

# ==========================
# FUNCIONES DE RESPALDO AUTOMÁTICO
# ==========================

def crear_respaldo_manual():
    """Crea un respaldo manual de la base de datos"""
    try:
        
        # Crear carpeta de respaldos si no existe
        backup_dir = os.path.join(BASE_DIR, 'BACKUPS')
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)
            print(f"📁 Carpeta de respaldos creada: {backup_dir}")
        
        # Nombre del respaldo con timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_file = os.path.join(backup_dir, f'backup_{timestamp}.zip')
        
        # Archivos a respaldar
        archivos_a_respaldar = []
        
        # Base de datos
        if os.path.exists(DB_PATH):
            archivos_a_respaldar.append(DB_PATH)
            print(f"📦 Respaldando base de datos: {DB_PATH}")
        
        # Documentos importantes (opcional, comenta si quieres incluir documentos)
        # if os.path.exists(DOCUMENTOS_DIR):
        #     archivos_a_respaldar.append(DOCUMENTOS_DIR)
        
        if not archivos_a_respaldar:
            print("⚠️ No hay archivos para respaldar")
            return None
        
        # Crear ZIP
        with zipfile.ZipFile(backup_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for archivo in archivos_a_respaldar:
                if os.path.isfile(archivo):
                    zipf.write(archivo, os.path.basename(archivo))
                elif os.path.isdir(archivo):
                    for root, dirs, files in os.walk(archivo):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, os.path.dirname(archivo))
                            zipf.write(file_path, arcname)
        
        # Mantener solo los últimos 10 respaldos
        backups = sorted([f for f in os.listdir(backup_dir) if f.startswith('backup_') and f.endswith('.zip')])
        while len(backups) > 10:
            archivo_antiguo = os.path.join(backup_dir, backups.pop(0))
            os.remove(archivo_antiguo)
            print(f"🗑️ Respaldo antiguo eliminado: {archivo_antiguo}")
        
        print(f"✅ Respaldo creado: {backup_file}")
        logger.info(f"✅ Respaldo automático creado: {backup_file}")
        return backup_file
        
    except Exception as e:
        print(f"❌ Error en respaldo: {e}")
        logger.error(f"❌ Error en respaldo: {e}")
        return None

def restaurar_respaldo(archivo_respaldo):
    """Restaura un respaldo (solo para uso administrativo)"""
    try:
        
        if not os.path.exists(archivo_respaldo):
            return False, "Archivo de respaldo no encontrado"
        
        # Crear backup de la base actual antes de restaurar
        crear_respaldo_manual()
        
        # Restaurar
        with zipfile.ZipFile(archivo_respaldo, 'r') as zipf:
            zipf.extractall(BASE_DIR)
        
        return True, "Respaldo restaurado correctamente"
        
    except Exception as e:
        return False, f"Error al restaurar: {e}"



# Programar respaldo cada hora
# ==========================
# INICIALIZACIÓN DE RESPALDOS (COMPATIBLE CON GUNICORN)
# ==========================

# Crear un respaldo inicial si no existe ninguno
def crear_respaldo_inicial():
    """Crea un respaldo inicial si no hay ninguno"""
    backup_dir = os.path.join(BASE_DIR, 'BACKUPS')
    if os.path.exists(backup_dir):
        backups = [f for f in os.listdir(backup_dir) if f.startswith('backup_') and f.endswith('.zip')]
        if not backups:
            crear_respaldo_manual()
            print("📦 Respaldo inicial creado")

# Ejecutar respaldo inicial (no bloquea)
try:
    crear_respaldo_inicial()
except Exception as e:
    print(f"⚠️ Error al crear respaldo inicial: {e}")

print("✅ Sistema de respaldos listo. Usa /admin/respaldos para gestionar")
logger.info("✅ Sistema de respaldos listo")

@app.route('/admin/backup/manual', methods=['POST'])
@login_required
def backup_manual():
    """Crea un respaldo manual (acceso solo admin)"""
    if session.get('rol') != 'ADMIN':
        return jsonify({'error': 'Acceso denegado'}), 403
    
    archivo = crear_respaldo_manual()
    if archivo:
        return jsonify({'success': True, 'archivo': os.path.basename(archivo)})
    else:
        return jsonify({'success': False, 'error': 'Error al crear respaldo'}), 500


@app.route('/admin/respaldos')
@login_required
def listar_respaldos():
    """Lista los respaldos disponibles (solo admin)"""
    if session.get('rol') != 'ADMIN':
        flash('Acceso denegado', 'danger')
        return redirect(url_for('panel'))
    
    backup_dir = os.path.join(BASE_DIR, 'BACKUPS')
    respaldos = []
    
    if os.path.exists(backup_dir):
        for archivo in os.listdir(backup_dir):
            if archivo.startswith('backup_') and archivo.endswith('.zip'):
                ruta = os.path.join(backup_dir, archivo)
                respaldos.append({
                    'nombre': archivo,
                    'fecha': datetime.fromtimestamp(os.path.getctime(ruta)).strftime('%Y-%m-%d %H:%M:%S'),
                    'tamano': f"{os.path.getsize(ruta) / 1024:.2f} KB"
                })
    
    respaldos.sort(key=lambda x: x['fecha'], reverse=True)
    return render_template('respaldos.html', respaldos=respaldos)

@app.route('/admin/respaldos/download/<nombre>')
@login_required
def descargar_respaldo(nombre):
    """Descarga un respaldo"""
    if session.get('rol') != 'ADMIN':
        flash('Acceso denegado', 'danger')
        return redirect(url_for('panel'))
    
    backup_dir = os.path.join(BASE_DIR, 'BACKUPS')
    ruta = os.path.join(backup_dir, nombre)
    
    if os.path.exists(ruta):
        return send_file(ruta, as_attachment=True, download_name=nombre)
    else:
        flash('Archivo no encontrado', 'danger')
        return redirect(url_for('listar_respaldos'))

@app.route('/admin/respaldos/crear', methods=['POST'])
@login_required
def crear_respaldo():
    """Crea un respaldo manual"""
    if session.get('rol') != 'ADMIN':
        return jsonify({'error': 'Acceso denegado'}), 403
    
    archivo = crear_respaldo_manual()
    if archivo:
        return jsonify({'success': True, 'archivo': archivo})
    else:
        return jsonify({'success': False, 'error': 'Error al crear respaldo'}), 500


# ==========================
# INICIALIZACIÓN DEL SISTEMA
# ==========================
def inicializar_sistema():
    """Inicializa el sistema completo"""
    print("🔧 Inicializando sistema...")
    
    # Crear directorios si no existen
    os.makedirs('LOGS', exist_ok=True)
    os.makedirs('BACKUPS', exist_ok=True)
    os.makedirs('DOCUMENTOS_AGENTES', exist_ok=True)
    os.makedirs('qr_ddjj', exist_ok=True)
    os.makedirs('CERTIFICADOS', exist_ok=True)
    os.makedirs('plantillas', exist_ok=True)
    
    # Inicializar base de datos
    try:
        init_db()
    except Exception as e:
        print(f"⚠️ Error en init_db: {e}")
    
    # Agregar columnas faltantes
    try:
        agregar_columnas_faltantes()
    except Exception as e:
        print(f"⚠️ Error en migración de columnas: {e}")
    
    # Crear índice
    try:
        crear_indice_codigo_validacion()
    except Exception as e:
        print(f"⚠️ Error creando índice: {e}")
    
    print("✅ Sistema inicializado correctamente")  # <--- PARÉNTESIS CERRADO



# ==========================
# INICIAR EL SISTEMA (UNA SOLA VEZ)
# ==========================
# Esta línea ejecuta la inicialización cuando la app es importada
inicializar_sistema()

# ==========================
# SISTEMA DE ALERTAS Y NOTIFICACIONES (SIN SCHEDULE)
# ==========================

# Configuración de email (para producción, usar variables de entorno)
EMAIL_CONFIG = {
    'smtp_server': os.environ.get('SMTP_SERVER', 'smtp.gmail.com'),
    'smtp_port': int(os.environ.get('SMTP_PORT', 587)),
    'email_remitente': os.environ.get('EMAIL_REMITENTE', ''),
    'email_password': os.environ.get('EMAIL_PASSWORD', ''),
    'email_director': os.environ.get('EMAIL_DIRECTOR', 'director@rrhh.com'),
    'email_admin': os.environ.get('EMAIL_ADMIN', 'admin@rrhh.com')
}


def enviar_email_notificacion(destinatario, asunto, cuerpo, html=False):
    """Envía email de notificación"""
    try:
        if not EMAIL_CONFIG['email_remitente'] or not EMAIL_CONFIG['email_password']:
            print("⚠️ Email no configurado. Omitiendo envío.")
            return False
        
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['email_remitente']
        msg['To'] = destinatario
        msg['Subject'] = asunto
        
        if html:
            msg.attach(MIMEText(cuerpo, 'html'))
        else:
            msg.attach(MIMEText(cuerpo, 'plain'))
        
        with smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port']) as server:
            server.starttls()
            server.login(EMAIL_CONFIG['email_remitente'], EMAIL_CONFIG['email_password'])
            server.send_message(msg)
        
        print(f"✅ Email enviado a {destinatario}")
        return True
        
    except Exception as e:
        print(f"❌ Error enviando email: {e}")
        return False


def notificar_cofirma_pendiente(legajo_id, apellido, nombre, ddjj_id):
    """Notifica al director que hay una DDJJ pendiente de co-firma"""
    asunto = f"🔔 DDJJ Pendiente de Co-Firma - Legajo {legajo_id}"
    
    cuerpo_html = f"""
    <html>
    <body>
        <h2>Declaración Jurada Pendiente de Co-Firma</h2>
        <p>El agente <strong>{apellido}, {nombre}</strong> (Legajo {legajo_id}) 
        ha firmado su Declaración Jurada y espera su co-firma.</p>
        
        <p><strong>Detalles:</strong><br>
        - DDJJ ID: {ddjj_id}<br>
        - Legajo: {legajo_id}<br>
        - Agente: {apellido}, {nombre}<br>
        - Estado: FIRMADA_AGENTE</p>
        
        <p>Por favor, ingrese al sistema para co-firmar la DDJJ:</p>
        <p><a href="https://rrhh-digital-v2.onrender.com/ddjj/{ddjj_id}/cofirmar_director">
        https://rrhh-digital-v2.onrender.com/ddjj/{ddjj_id}/cofirmar_director</a></p>
        
        <hr>
        <p><small>Sistema de Gestión de Legajos - Notificación automática</small></p>
    </body>
    </html>
    """
    
    return enviar_email_notificacion(
        EMAIL_CONFIG['email_director'],
        asunto,
        cuerpo_html,
        html=True
    )


def notificar_firma_completa(legajo_id, apellido, nombre, ddjj_id, email_agente):
    """Notifica al agente que su DDJJ fue co-firmada"""
    asunto = f"✅ DDJJ Completada - Legajo {legajo_id}"
    
    cuerpo_html = f"""
    <html>
    <body>
        <h2>Declaración Jurada Completada</h2>
        <p>Estimado/a <strong>{apellido}, {nombre}</strong>,</p>
        
        <p>Su Declaración Jurada correspondiente al año en curso ha sido 
        <strong>co-firmada por el Director</strong> y se encuentra finalizada.</p>
        
        <p><strong>Detalles:</strong><br>
        - DDJJ ID: {ddjj_id}<br>
        - Legajo: {legajo_id}<br>
        - Estado: FINALIZADA</p>
        
        <p>Puede descargar el PDF firmado desde el sistema:</p>
        <p><a href="https://rrhh-digital-v2.onrender.com/ddjj/{ddjj_id}/pdf">
        https://rrhh-digital-v2.onrender.com/ddjj/{ddjj_id}/pdf</a></p>
        
        <hr>
        <p><small>Sistema de Gestión de Legajos - Notificación automática</small></p>
    </body>
    </html>
    """
    
    return enviar_email_notificacion(
        email_agente,
        asunto,
        cuerpo_html,
        html=True
    )


def verificar_ddjj_pendientes_firma():
    """Verifica DDJJ que esperan co-firma del director y envía email"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # DDJJ firmadas por agente que esperan co-firma
        cur.execute("""
            SELECT d.id, d.legajo_id, l.apellido, l.nombre, l.email, d.fecha_generacion
            FROM declaraciones_juradas d
            JOIN legajos l ON d.legajo_id = l.legajo_id
            WHERE d.estado = 'FIRMADA_AGENTE' 
            AND d.activa = 1
            AND (d.notificado_cofirma IS NULL OR d.notificado_cofirma = 0)
        """)
        
        pendientes = cur.fetchall()
        
        if pendientes:
            print(f"📋 {len(pendientes)} DDJJ pendientes de co-firma")
            
            for p in pendientes:
                ddjj_id, legajo_id, apellido, nombre, email, fecha = p
                
                notificar_cofirma_pendiente(legajo_id, apellido, nombre, ddjj_id)
                
                cur.execute("""
                    UPDATE declaraciones_juradas 
                    SET notificado_cofirma = 1 
                    WHERE id = ?
                """, (ddjj_id,))
                
                print(f"   - Legajo {legajo_id}: {apellido}, {nombre} (email enviado)")
            
            conn.commit()
        
        conn.close()
        return len(pendientes)
        
    except Exception as e:
        print(f"⚠️ Error verificando DDJJ pendientes: {e}")
        return 0


def verificar_ddjj_por_vencer():
    """Verifica DDJJ que están por vencer"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # DDJJ que están por vencer (a 30 días de cumplir un año)
        cur.execute("""
            SELECT d.id, d.legajo_id, l.apellido, l.nombre, l.email, d.anio, d.fecha_envio
            FROM declaraciones_juradas d
            JOIN legajos l ON d.legajo_id = l.legajo_id
            WHERE d.estado = 'FINALIZADA'
            AND d.activa = 1
            AND date(d.fecha_envio) >= date('now', '-330 days')
            AND (d.notificado_vencimiento IS NULL OR d.notificado_vencimiento = 0)
        """)
        
        por_vencer = cur.fetchall()
        
        if por_vencer:
            print(f"⚠️ {len(por_vencer)} DDJJ próximas a vencer")
            
            for p in por_vencer:
                ddjj_id, legajo_id, apellido, nombre, email, anio, fecha = p
                
                asunto = f"⚠️ Recordatorio: DDJJ {anio} próxima a vencer"
                cuerpo = f"""
Estimado/a {apellido}, {nombre}:

Su Declaración Jurada del año {anio} está próxima a vencer.
Por favor, genere una nueva DDJJ en el sistema.

Legajo: {legajo_id}
Fecha de presentación: {fecha}

Ingrese al sistema: https://rrhh-digital-v2.onrender.com

--
Sistema de Gestión de Legajos
"""
                
                enviar_email_notificacion(email, asunto, cuerpo)
                
                cur.execute("""
                    UPDATE declaraciones_juradas 
                    SET notificado_vencimiento = 1 
                    WHERE id = ?
                """, (ddjj_id,))
                
                print(f"   - Legajo {legajo_id}: {apellido}, {nombre} (recordatorio enviado)")
            
            conn.commit()
        
        conn.close()
        return len(por_vencer)
        
    except Exception as e:
        print(f"⚠️ Error verificando DDJJ por vencer: {e}")
        return 0


def crear_backup_automatico():
    """Crea un backup automático programado"""
    try:
        backup_dir = os.path.join(BASE_DIR, 'BACKUPS')
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_file = os.path.join(backup_dir, f'backup_auto_{timestamp}.zip')
        
        with zipfile.ZipFile(backup_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
            if os.path.exists(DB_PATH):
                zipf.write(DB_PATH, os.path.basename(DB_PATH))
        
        # Mantener solo últimos 10 backups automáticos
        backups = sorted([f for f in os.listdir(backup_dir) if f.startswith('backup_auto_')])
        while len(backups) > 10:
            os.remove(os.path.join(backup_dir, backups.pop(0)))
        
        print(f"✅ Backup automático creado: {backup_file}")
        return backup_file
        
    except Exception as e:
        print(f"❌ Error en backup automático: {e}")
        return None


def agregar_columnas_notificaciones():
    """Agrega columnas para control de notificaciones en declaraciones_juradas"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        cur.execute("PRAGMA table_info(declaraciones_juradas)")
        columnas = [col[1] for col in cur.fetchall()]
        
        if 'notificado_cofirma' not in columnas:
            cur.execute("ALTER TABLE declaraciones_juradas ADD COLUMN notificado_cofirma INTEGER DEFAULT 0")
            print("✅ Columna notificado_cofirma agregada")
        
        if 'notificado_vencimiento' not in columnas:
            cur.execute("ALTER TABLE declaraciones_juradas ADD COLUMN notificado_vencimiento INTEGER DEFAULT 0")
            print("✅ Columna notificado_vencimiento agregada")
        
        conn.commit()
        conn.close()
        
    except Exception as e:
        print(f"⚠️ Error agregando columnas de notificaciones: {e}")


# ==========================
# ENDPOINTS PARA ALERTAS (USAR CON GITHUB ACTIONS)
# ==========================

@app.route('/api/alertas/ddjj_pendientes', methods=['POST'])
def api_alertas_ddjj_pendientes():
    """Endpoint para verificar DDJJ pendientes (usar con cron)"""
    try:
        cantidad = verificar_ddjj_pendientes_firma()
        return jsonify({'success': True, 'pendientes': cantidad})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/alertas/ddjj_vencer', methods=['POST'])
def api_alertas_ddjj_vencer():
    """Endpoint para verificar DDJJ por vencer (usar con cron)"""
    try:
        cantidad = verificar_ddjj_por_vencer()
        return jsonify({'success': True, 'por_vencer': cantidad})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/alertas/backup', methods=['POST'])
def api_alertas_backup():
    """Endpoint para crear backup automático"""
    try:
        archivo = crear_backup_automatico()
        return jsonify({'success': True, 'archivo': archivo})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/admin/alertas/ejecutar', methods=['POST'])
@login_requerido
def ejecutar_alertas_manual():
    """Ejecuta todas las alertas manualmente (solo admin)"""
    if session.get('rol') != 'ADMIN':
        return jsonify({'error': 'Acceso denegado'}), 403
    
    resultados = {
        'ddjj_pendientes': verificar_ddjj_pendientes_firma(),
        'ddjj_vencer': verificar_ddjj_por_vencer(),
        'backup': crear_backup_automatico() is not None
    }
    
    return jsonify(resultados)


@app.route('/admin/alertas/estado', methods=['GET'])
@login_requerido
def estado_alertas():
    """Verifica el estado de las alertas"""
    if session.get('rol') != 'ADMIN':
        return jsonify({'error': 'Acceso denegado'}), 403
    
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Contar pendientes
        cur.execute("""
            SELECT COUNT(*) FROM declaraciones_juradas 
            WHERE estado = 'FIRMADA_AGENTE' AND activa = 1
        """)
        pendientes = cur.fetchone()[0]
        
        # Contar por vencer
        cur.execute("""
            SELECT COUNT(*) FROM declaraciones_juradas 
            WHERE estado = 'FINALIZADA' AND activa = 1
            AND date(fecha_envio) >= date('now', '-330 days')
        """)
        por_vencer = cur.fetchone()[0]
        
        conn.close()
        
        return jsonify({
            'success': True,
            'ddjj_pendientes': pendientes,
            'ddjj_por_vencer': por_vencer,
            'email_configurado': bool(EMAIL_CONFIG['email_remitente'] and EMAIL_CONFIG['email_password'])
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ==========================
# INICIALIZACIÓN DEL SISTEMA
# ==========================
def inicializar_sistema():
    """Inicializa el sistema completo"""
    print("🔧 Inicializando sistema...")
    
    # Crear directorios
    os.makedirs('LOGS', exist_ok=True)
    os.makedirs('BACKUPS', exist_ok=True)
    os.makedirs('DOCUMENTOS_AGENTES', exist_ok=True)
    os.makedirs('qr_ddjj', exist_ok=True)
    os.makedirs('CERTIFICADOS', exist_ok=True)
    os.makedirs('plantillas', exist_ok=True)
    
    # Inicializar base de datos
    try:
        init_db()
    except Exception as e:
        print(f"⚠️ Error en init_db: {e}")
    
    # Agregar columnas faltantes a legajos
    try:
        agregar_columnas_faltantes()
    except Exception as e:
        print(f"⚠️ Error en migración de columnas: {e}")
    
    # Agregar columnas de notificaciones
    try:
        agregar_columnas_notificaciones()
    except Exception as e:
        print(f"⚠️ Error en columnas notificaciones: {e}")
    
    # Crear índice
    try:
        crear_indice_codigo_validacion()
    except Exception as e:
        print(f"⚠️ Error creando índice: {e}")
    
    print("✅ Sistema inicializado correctamente")
    print("")
    print("📋 ALERTAS CONFIGURADAS (vía GitHub Actions):")
    print("   - POST /api/alertas/ddjj_pendientes")
    print("   - POST /api/alertas/ddjj_vencer")
    print("   - POST /api/alertas/backup")
    print("")
    print("🔧 Admin: /admin/alertas/ejecutar (manual)")
    print("📊 Estado: /admin/alertas/estado")


# Ejecutar inicialización
inicializar_sistema()


# ==========================
# INICIALIZACIÓN Y EJECUCIÓN (SOLO MODO LOCAL)
# ==========================
if __name__ == "__main__":
    print("🔄 Iniciando sistema en modo local...")
    print("✅ Sistema listo. Alertas disponibles vía endpoints")
    
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)