# verificar_plantilla.py
from docx import Document
import re

print("🔍 VERIFICANDO PLANTILLA...")

plantilla = r'G:/Mi unidad/RRHH DIGITAL OFICIAL/BASE_DATOS/plantillas/PLANTILLA_DDJJ_OFICIAL.docx'

try:
    doc = Document(plantilla)
    print("✅ Plantilla abierta correctamente")
    
    campos = set()
    
    # Buscar en párrafos
    for p in doc.paragraphs:
        matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', p.text)
        for match in matches:
            campos.add(match.strip())
    
    # Buscar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', cell.text)
                for match in matches:
                    campos.add(match.strip())
    
    print("\n📋 CAMPOS ENCONTRADOS:")
    print("-" * 40)
    for campo in sorted(campos):
        print(f"  • {campo}")
    print("-" * 40)
    print(f"📊 Total: {len(campos)} campos")
    
except Exception as e:
    print(f"❌ Error: {e}")